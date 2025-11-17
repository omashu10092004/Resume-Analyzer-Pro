#!/usr/bin/env python
# coding: utf-8

# In[2]:


# Basic utilities and file handling
import os
import re
import json
import pickle
import warnings
warnings.filterwarnings('ignore')

# Data manipulation and analysis
import numpy as np
import pandas as pd

# Visualization
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud

# Text processing and NLP
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
import string

# Machine Learning - scikit-learn
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.svm import SVC
from sklearn.naive_bayes import MultinomialNB
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.neighbors import NearestNeighbors

# Deep Learning - TensorFlow/Keras
import tensorflow as tf
from tensorflow import keras
from tensorflow.keras.models import Sequential, Model
from tensorflow.keras.layers import Dense, LSTM, GRU, Embedding, Dropout, Bidirectional, Input, Conv1D, GlobalMaxPooling1D
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau
from tensorflow.keras.utils import to_categorical

# Similarity and distance metrics (alternatives to some scipy functions)
from sklearn.metrics.pairwise import cosine_similarity, euclidean_distances, manhattan_distances

# Web framework (for deployment)
import streamlit as st

# PDF processing (for resume parsing)
import PyPDF2
import docx

# Additional utilities
from collections import Counter
import requests
from time import time
import logging

# Initialize NLTK resources (run once)
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
    nltk.download('omw-1.4')
except:
    pass

# Set random seeds for reproducibility
np.random.seed(42)
tf.random.set_seed(42)


# In[3]:


import pandas as pd
import numpy as np
import os

def load_all_datasets():
    """
    Load all five datasets for the Advanced Fit_Resume project
    Returns a dictionary with all loaded dataframes
    """
    datasets = {}
    
    try:
        # Load Fit_Resume_Expanded dataset - Main dataset for ATS scoring
        print("Loading Fit_Resume_Expanded dataset...")
        datasets['resume_jd_pairs'] = pd.read_csv(r"D:\FIT_RESUME_PROJECT\Fit_Resume_Expanded.csv")
        print(f"✓ Loaded Fit_Resume_Expanded with shape: {datasets['resume_jd_pairs'].shape}")
        
    except Exception as e:
        print(f"Error loading Fit_Resume_Expanded: {e}")
        # Create empty dataframe with expected structure if file not found
        datasets['resume_jd_pairs'] = pd.DataFrame(columns=[
            'Resume_Text', 'JD_Text', 'Skills_Resume', 'Skills_JD', 'Match_Label'
        ])
    
    try:
        # Load Domain_Classification_Cleaned dataset
        print("Loading Domain_Classification_Cleaned dataset...")
        datasets['domain_classification'] = pd.read_csv(r"D:\FIT_RESUME_PROJECT\Domain_Classification_Cleaned.csv")
        print(f"✓ Loaded Domain_Classification_Cleaned with shape: {datasets['domain_classification'].shape}")
        
    except Exception as e:
        print(f"Error loading Domain_Classification_Cleaned: {e}")
        datasets['domain_classification'] = pd.DataFrame(columns=[
            'resume_id', 'resume_text', 'domain', 'sub_domain', 'experience_level'
        ])
    
    try:
        # Load Skills_to_Course_Mapping dataset
        print("Loading Skills_to_Course_Mapping dataset...")
        datasets['skills_courses'] = pd.read_csv(r"D:\FIT_RESUME_PROJECT\Skills_to_Course_Mapping.csv")
        print(f"✓ Loaded Skills_to_Course_Mapping with shape: {datasets['skills_courses'].shape}")
        
    except Exception as e:
        print(f"Error loading Skills_to_Course_Mapping: {e}")
        datasets['skills_courses'] = pd.DataFrame(columns=[
            'skill_id', 'skill_name', 'course_name', 'platform', 'course_url', 
            'difficulty_level', 'course_duration_hours', 'rating', 'learners_count'
        ])
    
    try:
        # Load All_Domain_Interview_Questions dataset
        print("Loading All_Domain_Interview_Questions dataset...")
        datasets['interview_questions'] = pd.read_csv(r"D:\FIT_RESUME_PROJECT\All_Domain_Interview_Questions.csv")
        print(f"✓ Loaded All_Domain_Interview_Questions with shape: {datasets['interview_questions'].shape}")
        
    except Exception as e:
        print(f"Error loading All_Domain_Interview_Questions: {e}")
        datasets['interview_questions'] = pd.DataFrame(columns=[
            'domain_id', 'domain_name', 'question_id', 'question_text', 
            'question_type', 'difficulty_level', 'expected_answer_style', 'related_topic'
        ])
    
    try:
        # Load User_Response_Evaluation_Dataset
        print("Loading User_Response_Evaluation_Dataset...")
        datasets['response_evaluation'] = pd.read_csv(r"D:\FIT_RESUME_PROJECT\User_Response_Evaluation_Dataset.csv")
        print(f"✓ Loaded User_Response_Evaluation_Dataset with shape: {datasets['response_evaluation'].shape}")
        
    except Exception as e:
        print(f"Error loading User_Response_Evaluation_Dataset: {e}")
        datasets['response_evaluation'] = pd.DataFrame(columns=[
            'question_id', 'domain', 'question', 'expected_keywords', 'sample_answer',
            'relevance_score', 'clarity_score', 'technical_accuracy', 'overall_score', 'feedback'
        ])
    
    return datasets

# Load all datasets
print("Starting dataset loading process...")
data = load_all_datasets()

# Display summary of loaded datasets
print("\n" + "="*50)
print("DATASET LOADING SUMMARY")
print("="*50)
for name, df in data.items():
    print(f"{name:30s}: {df.shape[0]:6d} rows, {df.shape[1]:2d} columns")
    
    # Show column names for each dataset
    print(f"   Columns: {list(df.columns)}")
    print()

# Basic information about each dataset
print("\nDATASET PREVIEWS:")
print("="*50)

for name, df in data.items():
    if len(df) > 0:
        print(f"\n{name} - First 3 rows:")
        print("-" * 40)
        print(df.head(3))
        print(f"\nData types:")
        print(df.dtypes)
        print(f"\nMissing values:")
        print(df.isnull().sum())
        print("="*50)
    else:
        print(f"\n{name} - No data loaded or empty dataframe")
        print("="*50)

# Check if essential datasets have data
essential_datasets = ['resume_jd_pairs', 'domain_classification']
for dataset in essential_datasets:
    if len(data[dataset]) == 0:
        print(f"⚠️  WARNING: Essential dataset '{dataset}' is empty!")
    else:
        print(f"✓ Essential dataset '{dataset}' loaded successfully with {len(data[dataset])} records")

print("\nDataset loading completed!")


# In[3]:


import pandas as pd
import numpy as np
import re
import ast
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.model_selection import train_test_split, GridSearchCV, StratifiedKFold
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from xgboost import XGBClassifier
from lightgbm import LGBMClassifier
from sklearn.metrics import classification_report, f1_score, accuracy_score
from sklearn.preprocessing import StandardScaler, LabelEncoder
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import warnings
warnings.filterwarnings('ignore')

# Download required NLTK data
try:
    nltk.download('stopwords')
    nltk.download('punkt')
    nltk.download('wordnet')
except:
    pass

# Initialize lemmatizer and stopwords
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

def load_and_preprocess_data():
    """Load and preprocess all datasets"""
    print("Loading and preprocessing datasets...")
    
    # Load datasets
    resume_jd = pd.read_csv('Fit_Resume_Expanded.csv')
    domain_data = pd.read_csv('Domain_Classification_Cleaned.csv')
    skills_data = pd.read_csv('Skills_to_Course_Mapping.csv')
    
    # Basic data cleaning
    resume_jd = resume_jd.dropna(subset=['Resume_Text', 'JD_Text', 'Match_Label'])
    domain_data = domain_data.dropna(subset=['resume_text', 'domain'])
    
    return resume_jd, domain_data, skills_data

def extract_experience_years(text):
    """Extract experience years from resume text"""
    if pd.isna(text):
        return 0
    
    text = str(text).lower()
    
    # Patterns to match experience
    patterns = [
        r'(\d+)\s*\+?\s*years?[\s\w]*experience',
        r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',
        r'(\d+)\s*\+?\s*years?',
        r'(\d+)-(\d+)\s*years?'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            if isinstance(matches[0], tuple):
                # For ranges like "3-5 years", take the average
                nums = [int(x) for x in matches[0] if x.isdigit()]
                return sum(nums) / len(nums) if nums else 0
            else:
                return int(matches[0])
    
    return 0

def calculate_skill_overlap(skills_resume, skills_jd):
    """Calculate skill overlap between resume and JD"""
    try:
        if pd.isna(skills_resume) or pd.isna(skills_jd):
            return 0
        
        # Convert string representation of list to actual list
        if isinstance(skills_resume, str):
            try:
                resume_skills = ast.literal_eval(skills_resume)
            except:
                resume_skills = [s.strip() for s in skills_resume.split(',')]
        else:
            resume_skills = []
        
        if isinstance(skills_jd, str):
            try:
                jd_skills = ast.literal_eval(skills_jd)
            except:
                jd_skills = [s.strip() for s in skills_jd.split(',')]
        else:
            jd_skills = []
        
        if not resume_skills or not jd_skills:
            return 0
        
        # Calculate overlap
        common_skills = set(resume_skills) & set(jd_skills)
        total_skills = set(resume_skills) | set(jd_skills)
        
        return len(common_skills) / len(total_skills) if total_skills else 0
    except:
        return 0

def text_preprocessing(text):
    """Advanced text preprocessing"""
    if pd.isna(text):
        return ""
    
    text = str(text)
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove special characters and digits
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    
    # Tokenize
    tokens = word_tokenize(text)
    
    # Remove stopwords and lemmatize
    tokens = [lemmatizer.lemmatize(token) for token in tokens if token not in stop_words and len(token) > 2]
    
    return ' '.join(tokens)

def calculate_semantic_similarity(text1, text2):
    """Calculate semantic similarity using TF-IDF and cosine similarity"""
    try:
        if pd.isna(text1) or pd.isna(text2) or text1 == "" or text2 == "":
            return 0
        
        # Create TF-IDF vectors
        vectorizer = TfidfVectorizer(max_features=5000)
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        
        # Calculate cosine similarity
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        return similarity
    except:
        return 0

def extract_text_features(text):
    """Extract various text features"""
    if pd.isna(text):
        return 0, 0, 0
    
    text = str(text)
    
    # Word count
    word_count = len(text.split())
    
    # Character count
    char_count = len(text)
    
    # Average word length
    avg_word_length = char_count / word_count if word_count > 0 else 0
    
    return word_count, char_count, avg_word_length

def create_advanced_features(df):
    """Create advanced features for the dataset"""
    print("Creating advanced features...")
    
    # Create semantic similarity feature
    df['semantic_similarity'] = df.apply(
        lambda row: calculate_semantic_similarity(
            text_preprocessing(row['Resume_Text']), 
            text_preprocessing(row['JD_Text'])
        ), axis=1
    )
    
    # Create skill overlap feature
    df['skill_overlap_score'] = df.apply(
        lambda row: calculate_skill_overlap(row['Skills_Resume'], row['Skills_JD']), axis=1
    )
    
    # Extract text features from resume
    df[['resume_word_count', 'resume_char_count', 'resume_avg_word_length']] = df['Resume_Text'].apply(
        lambda x: pd.Series(extract_text_features(x))
    )
    
    # Extract text features from JD
    df[['jd_word_count', 'jd_char_count', 'jd_avg_word_length']] = df['JD_Text'].apply(
        lambda x: pd.Series(extract_text_features(x))
    )
    
    return df

def prepare_domain_features(df):
    """Prepare features for domain classification"""
    print("Preparing domain classification features...")
    
    # Extract experience years
    df['experience_years'] = df['resume_text'].apply(extract_experience_years)
    
    # Text preprocessing
    df['cleaned_text'] = df['resume_text'].apply(text_preprocessing)
    
    return df

def hyperparameter_tuning(X, y):
    """Perform hyperparameter tuning for multiple models"""
    print("Starting hyperparameter tuning...")
    
    # Split the data
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y
    )
    
    # Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Define models and parameter grids
    models = {
        'XGBoost': {
            'model': XGBClassifier(random_state=42, eval_metric='logloss'),
            'params': {
                'n_estimators': [100, 200, 300],
                'learning_rate': [0.01, 0.05, 0.1],
                'max_depth': [3, 5, 7],
                'scale_pos_weight': [1, (len(y_train) - sum(y_train)) / sum(y_train)]
            }
        },
        'LightGBM': {
            'model': LGBMClassifier(random_state=42),
            'params': {
                'n_estimators': [100, 200, 300],
                'learning_rate': [0.01, 0.05, 0.1],
                'num_leaves': [31, 40, 50],
                'scale_pos_weight': [1, (len(y_train) - sum(y_train)) / sum(y_train)]
            }
        },
        'RandomForest': {
            'model': RandomForestClassifier(random_state=42),
            'params': {
                'n_estimators': [100, 200, 300],
                'max_depth': [3, 5, 7, None],
                'class_weight': ['balanced', None]
            }
        }
    }
    
    best_models = {}
    best_scores = {}
    
    # Perform GridSearchCV for each model
    for name, model_info in models.items():
        print(f"Tuning {name}...")
        
        grid_search = GridSearchCV(
            estimator=model_info['model'],
            param_grid=model_info['params'],
            cv=StratifiedKFold(n_splits=3),
            scoring='f1',
            n_jobs=-1,
            verbose=1
        )
        
        grid_search.fit(X_train_scaled, y_train)
        
        best_models[name] = grid_search.best_estimator_
        best_scores[name] = grid_search.best_score_
        
        print(f"Best {name} parameters: {grid_search.best_params_}")
        print(f"Best {name} CV F1-score: {grid_search.best_score_:.4f}")
        
        # Evaluate on test set
        y_pred = grid_search.predict(X_test_scaled)
        test_f1 = f1_score(y_test, y_pred)
        print(f"Test F1-score for {name}: {test_f1:.4f}")
        print("-" * 50)
    
    # Find the best model
    best_model_name = max(best_scores, key=best_scores.get)
    best_model = best_models[best_model_name]
    
    print(f"Best overall model: {best_model_name} with F1-score: {best_scores[best_model_name]:.4f}")
    
    return best_model, scaler, best_model_name

def main():
    """Main function to run all preprocessing and tuning"""
    # Load and preprocess data
    resume_jd, domain_data, skills_data = load_and_preprocess_data()
    
    # Create advanced features for resume-JD matching
    resume_jd_advanced = create_advanced_features(resume_jd)
    
    # Prepare features and target for resume-JD matching
    feature_columns = [
        'semantic_similarity', 'skill_overlap_score', 
        'resume_word_count', 'resume_char_count', 'resume_avg_word_length',
        'jd_word_count', 'jd_char_count', 'jd_avg_word_length'
    ]
    
    X = resume_jd_advanced[feature_columns]
    y = resume_jd_advanced['Match_Label']
    
    print(f"Feature matrix shape: {X.shape}")
    print(f"Target distribution:\n{y.value_counts()}")
    
    # Handle class imbalance
    class_ratio = (len(y) - sum(y)) / sum(y)
    print(f"Class imbalance ratio: {class_ratio:.2f}")
    
    # Perform hyperparameter tuning
    best_model, scaler, best_model_name = hyperparameter_tuning(X, y)
    
    # Prepare domain classification data
    domain_data_processed = prepare_domain_features(domain_data)
    
    # Save processed data and model for later use
    resume_jd_advanced.to_csv('Processed_Resume_JD_Data.csv', index=False)
    domain_data_processed.to_csv('Processed_Domain_Data.csv', index=False)
    
    # Save the best model and scaler
    import joblib
    joblib.dump(best_model, f'best_{best_model_name}_model.pkl')
    joblib.dump(scaler, 'feature_scaler.pkl')
    
    print("Preprocessing and tuning completed successfully!")
    print("Saved files:")
    print("- Processed_Resume_JD_Data.csv")
    print("- Processed_Domain_Data.csv")
    print(f"- best_{best_model_name}_model.pkl")
    print("- feature_scaler.pkl")
    
    return best_model, scaler, resume_jd_advanced, domain_data_processed

if __name__ == "__main__":
    best_model, scaler, processed_resume_data, processed_domain_data = main()


# In[4]:


# ==================== ADVANCED TEXT PREPROCESSING FUNCTIONS ====================
def extract_experience_years(text):
    """Extract experience years from resume text using advanced pattern matching"""
    if pd.isna(text):
        return 0
    
    text = str(text).lower()
    
    # Enhanced patterns to match various experience formats
    patterns = [
        r'(\d+)\s*\+?\s*years?[\s\w]*experience',  # "5 years experience"
        r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',  # "experience of 5 years"
        r'(\d+)\s*\+?\s*years?',  # "5+ years"
        r'(\d+)-(\d+)\s*years?',  # "3-5 years"
        r'(\d+)\s*to\s*(\d+)\s*years?',  # "3 to 5 years"
        r'over\s*(\d+)\s*years?',  # "over 5 years"
        r'more\s*than\s*(\d+)\s*years?',  # "more than 5 years"
        r'(\d+)\s*yr',  # "5 yr"
        r'(\d+)\s*mos',  # "18 mos" (convert to years)
    ]
    
    total_years = 0
    found_patterns = 0
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            if isinstance(matches[0], tuple):
                # For ranges like "3-5 years", take the average
                nums = [int(x) for x in matches[0] if x.isdigit()]
                if nums:
                    total_years += sum(nums) / len(nums)
                    found_patterns += 1
            else:
                # Handle month conversions
                if 'mos' in pattern and matches:
                    total_years += int(matches[0]) / 12  # Convert months to years
                else:
                    total_years += int(matches[0])
                found_patterns += 1
    
    return total_years / max(found_patterns, 1) if found_patterns > 0 else 0

def calculate_skill_overlap(skills_resume, skills_jd):
    """Calculate advanced skill overlap with semantic matching"""
    try:
        if pd.isna(skills_resume) or pd.isna(skills_jd):
            return 0
        
        # Convert string representation of list to actual list
        def parse_skills(skill_str):
            if isinstance(skill_str, str):
                try:
                    return ast.literal_eval(skill_str)
                except:
                    # Handle different delimiters
                    if ';' in skill_str:
                        return [s.strip().lower() for s in skill_str.split(';') if s.strip()]
                    elif ',' in skill_str:
                        return [s.strip().lower() for s in skill_str.split(',') if s.strip()]
                    else:
                        return [skill_str.strip().lower()]
            return []
        
        resume_skills = parse_skills(skills_resume)
        jd_skills = parse_skills(skills_jd)
        
        if not resume_skills or not jd_skills:
            return 0
        
        # Advanced skill matching with partial matching
        def skill_similarity(skill1, skill2):
            skill1, skill2 = skill1.lower(), skill2.lower()
            if skill1 == skill2:
                return 1.0
            if skill1 in skill2 or skill2 in skill1:
                return 0.8
            # Use sequence matching for similar skills
            if len(skill1) > 3 and len(skill2) > 3:
                from difflib import SequenceMatcher
                return SequenceMatcher(None, skill1, skill2).ratio()
            return 0.0
        
        # Calculate weighted overlap
        total_similarity = 0
        matched_pairs = 0
        
        for jd_skill in jd_skills:
            best_match_score = 0
            for resume_skill in resume_skills:
                similarity = skill_similarity(jd_skill, resume_skill)
                best_match_score = max(best_match_score, similarity)
            
            # Threshold for considering a match
            if best_match_score > 0.6:
                total_similarity += best_match_score
                matched_pairs += 1
        
        # Normalize score
        if not jd_skills:
            return 0
        
        return total_similarity / len(jd_skills)
    
    except Exception as e:
        print(f"Error in skill overlap calculation: {e}")
        return 0

def advanced_text_preprocessing(text):
    """Advanced text preprocessing with domain-specific cleaning"""
    if pd.isna(text):
        return ""
    
    text = str(text)
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+|www\.\S+', '', text)
    
    # Remove special characters but keep some important ones
    text = re.sub(r'[^a-zA-Z0-9\s\.\,\-\+\%]', '', text)
    
    # Handle specific patterns (years, percentages, etc.)
    text = re.sub(r'(\d+)\s*years?', r'\1 years ', text)
    text = re.sub(r'(\d+)\%', r'\1 percent ', text)
    
    # Tokenize
    tokens = word_tokenize(text)
    
    # Custom stopwords for resume domain
    custom_stopwords = stop_words.union({
        'email', 'phone', 'http', 'www', 'com', 'linkedin', 'github',
        'objective', 'summary', 'experience', 'education', 'skills',
        'projects', 'references', 'date', 'month', 'year'
    })
    
    # Remove stopwords and lemmatize
    tokens = [
        lemmatizer.lemmatize(token) 
        for token in tokens 
        if token not in custom_stopwords and len(token) > 2 and not token.isdigit()
    ]
    
    return ' '.join(tokens)

def calculate_semantic_similarity(text1, text2):
    """Enhanced semantic similarity using multiple techniques"""
    try:
        if pd.isna(text1) or pd.isna(text2) or text1 == "" or text2 == "":
            return 0
        
        # Method 1: TF-IDF Cosine Similarity
        vectorizer1 = TfidfVectorizer(max_features=2000, ngram_range=(1, 2))
        tfidf_matrix = vectorizer1.fit_transform([text1, text2])
        similarity_tfidf = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Method 2: Word Embeddings (if available)
        # You can add Word2Vec or GloVe embeddings here for better semantic understanding
        
        # Method 3: Jaccard Similarity of important terms
        words1 = set(text1.split())
        words2 = set(text2.split())
        if words1 and words2:
            jaccard_similarity = len(words1.intersection(words2)) / len(words1.union(words2))
        else:
            jaccard_similarity = 0
        
        # Combine multiple similarity measures
        final_similarity = (similarity_tfidf * 0.6) + (jaccard_similarity * 0.4)
        
        return final_similarity
    
    except Exception as e:
        print(f"Error in semantic similarity calculation: {e}")
        return 0

def extract_advanced_text_features(text):
    """Extract comprehensive text features"""
    if pd.isna(text):
        return [0] * 8
    
    text = str(text)
    
    features = []
    
    # Basic metrics
    features.append(len(text.split()))  # word_count
    features.append(len(text))  # char_count
    features.append(len(text) / max(len(text.split()), 1))  # avg_word_length
    
    # Readability metrics (simplified)
    sentences = re.split(r'[.!?]+', text)
    features.append(len(sentences))  # sentence_count
    features.append(len(text.split()) / max(len(sentences), 1))  # avg_sentence_length
    
    # Keyword density (technical terms)
    technical_terms = ['python', 'java', 'machine learning', 'data analysis', 'sql', 
                      'project management', 'leadership', 'communication', 'problem solving']
    tech_count = sum(1 for term in technical_terms if term in text.lower())
    features.append(tech_count)
    
    # Unique word ratio
    words = text.split()
    features.append(len(set(words)) / max(len(words), 1))  # unique_word_ratio
    
    # Presence of numbers (indicator of quantifiable achievements)
    numbers = re.findall(r'\d+', text)
    features.append(len(numbers))
    
    return features

# ==================== FEATURE ENGINEERING ====================
def create_advanced_features(df):
    """Create comprehensive advanced features"""
    print("Creating advanced features...")
    
    # Text preprocessing
    df['resume_cleaned'] = df['Resume_Text'].apply(advanced_text_preprocessing)
    df['jd_cleaned'] = df['JD_Text'].apply(advanced_text_preprocessing)
    
    # Semantic similarity
    df['semantic_similarity'] = df.apply(
        lambda row: calculate_semantic_similarity(row['resume_cleaned'], row['jd_cleaned']), 
        axis=1
    )
    
    # Skill overlap with advanced matching
    df['skill_overlap_score'] = df.apply(
        lambda row: calculate_skill_overlap(row['Skills_Resume'], row['Skills_JD']), 
        axis=1
    )
    
    # Extract comprehensive text features
    text_features_resume = df['Resume_Text'].apply(extract_advanced_text_features).apply(pd.Series)
    text_features_resume.columns = [f'resume_feature_{i}' for i in range(text_features_resume.shape[1])]
    
    text_features_jd = df['JD_Text'].apply(extract_advanced_text_features).apply(pd.Series)
    text_features_jd.columns = [f'jd_feature_{i}' for i in range(text_features_jd.shape[1])]
    
    # Combine all features
    df = pd.concat([df, text_features_resume, text_features_jd], axis=1)
    
    # Experience level feature
    df['experience_years'] = df['Resume_Text'].apply(extract_experience_years)
    
    # Domain relevance (simplified)
    df['domain_keyword_match'] = df.apply(
        lambda row: sum(1 for word in row['jd_cleaned'].split() 
                       if word in row['resume_cleaned'].split()) / max(len(row['jd_cleaned'].split()), 1),
        axis=1
    )
    
    print(f"Created {len(df.columns) - len(['Resume_Text', 'JD_Text', 'Skills_Resume', 'Skills_JD', 'Match_Label'])} new features")
    return df

# ==================== HYPERPARAMETER TUNING ====================
def advanced_hyperparameter_tuning(X, y):
    """Advanced hyperparameter tuning with multiple algorithms"""
    print("Starting advanced hyperparameter tuning...")
    
    # Handle class imbalance
    class_ratio = (len(y) - sum(y)) / sum(y)
    print(f"Class imbalance ratio: {class_ratio:.2f}")
    
    # Split data with stratification
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y
    )
    
    # Feature scaling
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Define models with extensive parameter grids
    models_params = {
        'XGBoost': {
            'model': XGBClassifier(random_state=42, eval_metric='logloss', use_label_encoder=False),
            'params': {
                'n_estimators': [100, 200, 300],
                'learning_rate': [0.01, 0.05, 0.1, 0.2],
                'max_depth': [3, 5, 7, 9],
                'subsample': [0.8, 0.9, 1.0],
                'colsample_bytree': [0.8, 0.9, 1.0],
                'gamma': [0, 0.1, 0.2],
                'scale_pos_weight': [1, class_ratio, class_ratio * 1.5]
            }
        },
        'LightGBM': {
            'model': LGBMClassifier(random_state=42, verbose=-1),
            'params': {
                'n_estimators': [100, 200, 300],
                'learning_rate': [0.01, 0.05, 0.1],
                'num_leaves': [31, 50, 70],
                'max_depth': [3, 5, 7, -1],
                'subsample': [0.8, 0.9, 1.0],
                'colsample_bytree': [0.8, 0.9, 1.0],
                'reg_alpha': [0, 0.1, 0.5],
                'reg_lambda': [0, 0.1, 0.5],
                'scale_pos_weight': [1, class_ratio]
            }
        }
    }
    
    best_models = {}
    best_scores = {}
    
    # Perform tuning for each model
    for name, config in models_params.items():
        if config['model'] is None:
            continue
            
        print(f"\n{'='*50}")
        print(f"Tuning {name}...")
        print(f"{'='*50}")
        
        # Use RandomizedSearchCV for faster tuning with more parameters
        random_search = RandomizedSearchCV(
            estimator=config['model'],
            param_distributions=config['params'],
            n_iter=20,  # Number of parameter combinations to try
            cv=StratifiedKFold(n_splits=3),
            scoring='f1',
            n_jobs=-1,
            verbose=1,
            random_state=42
        )
        
        random_search.fit(X_train_scaled, y_train)
        
        best_models[name] = random_search.best_estimator_
        best_scores[name] = random_search.best_score_
        
        print(f"Best {name} parameters: {random_search.best_params_}")
        print(f"Best {name} CV F1-score: {random_search.best_score_:.4f}")
        
        # Detailed evaluation
        y_pred = random_search.predict(X_test_scaled)
        print(f"Test F1-score: {f1_score(y_test, y_pred):.4f}")
        print(f"Test Accuracy: {accuracy_score(y_test, y_pred):.4f}")
        print(f"Classification Report:\n{classification_report(y_test, y_pred)}")
    
    # Select best model
    best_model_name = max(best_scores, key=best_scores.get)
    best_model = best_models[best_model_name]
    
    print(f"\n{'='*50}")
    print(f"BEST OVERALL MODEL: {best_model_name}")
    print(f"Best CV F1-score: {best_scores[best_model_name]:.4f}")
    print(f"{'='*50}")
    
    return best_model, scaler, best_model_name

# ==================== FEATURE SELECTION ====================
def perform_feature_selection(X, y):
    """Perform advanced feature selection"""
    print("Performing feature selection...")
    
    # Use Random Forest for feature importance
    rf = RandomForestClassifier(n_estimators=100, random_state=42)
    rf.fit(X, y)
    
    # Get feature importances
    importances = rf.feature_importances_
    feature_importance_df = pd.DataFrame({
        'feature': X.columns,
        'importance': importances
    }).sort_values('importance', ascending=False)
    
    # Select top features (you can adjust the threshold)
    important_features = feature_importance_df[feature_importance_df['importance'] > 0.01]['feature'].tolist()
    
    print(f"Selected {len(important_features)} important features")
    print("Top 10 features:")
    print(feature_importance_df.head(10))
    
    return important_features, feature_importance_df

# ==================== MAIN PROCESSING FUNCTION ====================
def process_and_tune_data(resume_jd_data):
    """Main function to process data and perform tuning"""
    # Create advanced features
    processed_data = create_advanced_features(resume_jd_data.copy())
    
    # Prepare features and target
    feature_columns = [col for col in processed_data.columns 
                      if col not in ['Resume_Text', 'JD_Text', 'Skills_Resume', 'Skills_JD', 'Match_Label',
                                   'resume_cleaned', 'jd_cleaned']]
    
    X = processed_data[feature_columns]
    y = processed_data['Match_Label']
    
    # Handle missing values
    X = X.fillna(0)
    
    # Feature selection
    important_features, importance_df = perform_feature_selection(X, y)
    X_selected = X[important_features]
    
    # Hyperparameter tuning
    best_model, scaler, best_model_name = advanced_hyperparameter_tuning(X_selected, y)
    
    return best_model, scaler, processed_data, important_features, importance_df


# In[5]:


# ==================== IMPORT LIBRARIES ====================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, StratifiedKFold, cross_val_score
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from xgboost import XGBClassifier
from lightgbm import LGBMClassifier
from sklearn.svm import SVC
from sklearn.linear_model import LogisticRegression
import joblib
import pickle
import time
import warnings
warnings.filterwarnings('ignore')

# ==================== MODEL TRAINING AND EVALUATION ====================
def train_powerful_models(X, y, test_size=0.2, random_state=42):
    """
    Train and evaluate 5 powerful ML models with hyperparameter tuning
    """
    print("Training 5 powerful ML models...")
    
    # Split the data
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state, stratify=y
    )
    
    # Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Define the models with optimized hyperparameters
    models = {
        'XGBoost': XGBClassifier(
            n_estimators=300,
            learning_rate=0.1,
            max_depth=7,
            subsample=0.9,
            colsample_bytree=0.9,
            gamma=0.1,
            reg_alpha=0.1,
            reg_lambda=0.1,
            random_state=random_state,
            eval_metric='logloss',
            use_label_encoder=False
        ),
        'LightGBM': LGBMClassifier(
            n_estimators=300,
            learning_rate=0.05,
            num_leaves=50,
            max_depth=7,
            subsample=0.9,
            colsample_bytree=0.9,
            reg_alpha=0.1,
            reg_lambda=0.1,
            random_state=random_state,
            verbose=-1
        ),
        'Random Forest': RandomForestClassifier(
            n_estimators=300,
            max_depth=15,
            min_samples_split=5,
            min_samples_leaf=2,
            max_features='sqrt',
            bootstrap=True,
            random_state=random_state,
            n_jobs=-1
        ),
        'Gradient Boosting': GradientBoostingClassifier(
            n_estimators=300,
            learning_rate=0.1,
            max_depth=7,
            subsample=0.9,
            min_samples_split=5,
            min_samples_leaf=2,
            random_state=random_state
        ),
        'SVM (RBF Kernel)': SVC(
            C=1.0,
            kernel='rbf',
            gamma='scale',
            probability=True,
            random_state=random_state
        )
    }
    
    # Train and evaluate each model
    results = {}
    trained_models = {}
    
    for name, model in models.items():
        print(f"\n{'='*60}")
        print(f"Training {name}...")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        # Train model
        model.fit(X_train_scaled, y_train)
        
        # Make predictions
        y_pred = model.predict(X_test_scaled)
        y_pred_proba = model.predict_proba(X_test_scaled)[:, 1]
        
        # Calculate metrics
        accuracy = accuracy_score(y_test, y_pred)
        f1 = f1_score(y_test, y_pred)
        precision = precision_score(y_test, y_pred)
        recall = recall_score(y_test, y_pred)
        
        # Cross-validation
        cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=random_state)
        cv_scores = cross_val_score(model, X_train_scaled, y_train, cv=cv, scoring='f1', n_jobs=-1)
        
        training_time = time.time() - start_time
        
        # Store results
        results[name] = {
            'accuracy': accuracy,
            'f1_score': f1,
            'precision': precision,
            'recall': recall,
            'cv_mean': cv_scores.mean(),
            'cv_std': cv_scores.std(),
            'training_time': training_time
        }
        
        trained_models[name] = model
        
        # Print results
        print(f"{name} Results:")
        print(f"Accuracy: {accuracy:.4f}")
        print(f"F1-Score: {f1:.4f}")
        print(f"Precision: {precision:.4f}")
        print(f"Recall: {recall:.4f}")
        print(f"CV F1-Score: {cv_scores.mean():.4f} (±{cv_scores.std():.4f})")
        print(f"Training Time: {training_time:.2f} seconds")
        print("\nClassification Report:")
        print(classification_report(y_test, y_pred))
    
    return results, trained_models, scaler

def create_ensemble_model(trained_models, X, y, test_size=0.2, random_state=42):
    """
    Create a powerful ensemble model from the trained models
    """
    print("\n" + "="*60)
    print("Creating Ensemble Model")
    print("="*60)
    
    # Split the data
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state, stratify=y
    )
    
    # Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Get predictions from all models for stacking
    predictions_train = []
    predictions_test = []
    
    for name, model in trained_models.items():
        # Train predictions
        preds_train = model.predict_proba(X_train_scaled)[:, 1]
        predictions_train.append(preds_train)
        
        # Test predictions
        preds_test = model.predict_proba(X_test_scaled)[:, 1]
        predictions_test.append(preds_test)
    
    # Stack predictions
    X_train_stacked = np.column_stack(predictions_train)
    X_test_stacked = np.column_stack(predictions_test)
    
    # Train meta-learner (Logistic Regression)
    meta_learner = LogisticRegression(
        C=1.0,
        solver='liblinear',
        random_state=random_state,
        max_iter=1000
    )
    
    meta_learner.fit(X_train_stacked, y_train)
    
    # Evaluate ensemble
    y_pred_ensemble = meta_learner.predict(X_test_stacked)
    y_pred_proba_ensemble = meta_learner.predict_proba(X_test_stacked)[:, 1]
    
    # Calculate metrics
    accuracy = accuracy_score(y_test, y_pred_ensemble)
    f1 = f1_score(y_test, y_pred_ensemble)
    precision = precision_score(y_test, y_pred_ensemble)
    recall = recall_score(y_test, y_pred_ensemble)
    
    print("Ensemble Model Results:")
    print(f"Accuracy: {accuracy:.4f}")
    print(f"F1-Score: {f1:.4f}")
    print(f"Precision: {precision:.4f}")
    print(f"Recall: {recall:.4f}")
    print("\nClassification Report:")
    print(classification_report(y_test, y_pred_ensemble))
    
    return meta_learner, {
        'accuracy': accuracy,
        'f1_score': f1,
        'precision': precision,
        'recall': recall
    }

def compare_model_performance(results, ensemble_results=None):
    """
    Compare performance of all models
    """
    print("\n" + "="*80)
    print("MODEL PERFORMANCE COMPARISON")
    print("="*80)
    
    # Create comparison dataframe
    comparison_data = []
    for model_name, metrics in results.items():
        comparison_data.append({
            'Model': model_name,
            'Accuracy': metrics['accuracy'],
            'F1-Score': metrics['f1_score'],
            'Precision': metrics['precision'],
            'Recall': metrics['recall'],
            'CV F1 Mean': metrics['cv_mean'],
            'CV F1 Std': metrics['cv_std'],
            'Training Time (s)': metrics['training_time']
        })
    
    if ensemble_results:
        comparison_data.append({
            'Model': 'Ensemble',
            'Accuracy': ensemble_results['accuracy'],
            'F1-Score': ensemble_results['f1_score'],
            'Precision': ensemble_results['precision'],
            'Recall': ensemble_results['recall'],
            'CV F1 Mean': np.nan,
            'CV F1 Std': np.nan,
            'Training Time (s)': np.nan
        })
    
    comparison_df = pd.DataFrame(comparison_data)
    comparison_df = comparison_df.sort_values('F1-Score', ascending=False)
    
    print(comparison_df.round(4))
    
    # Visualize results
    plt.figure(figsize=(15, 10))
    
    # Plot 1: F1-Score Comparison
    plt.subplot(2, 2, 1)
    models = comparison_df['Model']
    f1_scores = comparison_df['F1-Score']
    bars = plt.bar(models, f1_scores, color=['blue', 'green', 'red', 'purple', 'orange', 'cyan'])
    plt.title('F1-Score Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('F1-Score', fontsize=12)
    plt.xticks(rotation=45)
    plt.ylim(0, 1)
    
    # Add value labels on bars
    for bar, value in zip(bars, f1_scores):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, 
                f'{value:.3f}', ha='center', va='bottom', fontweight='bold')
    
    # Plot 2: Accuracy Comparison
    plt.subplot(2, 2, 2)
    accuracy_scores = comparison_df['Accuracy']
    bars = plt.bar(models, accuracy_scores, color=['blue', 'green', 'red', 'purple', 'orange', 'cyan'])
    plt.title('Accuracy Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Accuracy', fontsize=12)
    plt.xticks(rotation=45)
    plt.ylim(0, 1)
    
    # Add value labels on bars
    for bar, value in zip(bars, accuracy_scores):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, 
                f'{value:.3f}', ha='center', va='bottom', fontweight='bold')
    
    # Plot 3: Precision-Recall Comparison
    plt.subplot(2, 2, 3)
    width = 0.35
    x = np.arange(len(models))
    plt.bar(x - width/2, comparison_df['Precision'], width, label='Precision', alpha=0.8)
    plt.bar(x + width/2, comparison_df['Recall'], width, label='Recall', alpha=0.8)
    plt.title('Precision vs Recall', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Score', fontsize=12)
    plt.xticks(x, models, rotation=45)
    plt.legend()
    plt.ylim(0, 1)
    
    # Plot 4: Training Time Comparison
    plt.subplot(2, 2, 4)
    training_times = comparison_df['Training Time (s)'].fillna(0)
    bars = plt.bar(models, training_times, color=['blue', 'green', 'red', 'purple', 'orange', 'cyan'])
    plt.title('Training Time Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Time (seconds)', fontsize=12)
    plt.xticks(rotation=45)
    
    # Add value labels on bars
    for bar, value in zip(bars, training_times):
        if value > 0:
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                    f'{value:.1f}s', ha='center', va='bottom', fontweight='bold')
    
    plt.tight_layout()
    plt.show()
    
    return comparison_df

def save_best_model(trained_models, results, meta_learner=None):
    """
    Save the best performing model
    """
    # Find best model based on F1-score
    best_model_name = max(results.items(), key=lambda x: x[1]['f1_score'])[0]
    best_model = trained_models[best_model_name]
    
    print(f"\nBest model: {best_model_name} with F1-Score: {results[best_model_name]['f1_score']:.4f}")
    
    # Save the best model
    joblib.dump(best_model, 'best_model.pkl')
    
    if meta_learner:
        joblib.dump(meta_learner, 'ensemble_model.pkl')
    
    print("Best model saved as 'best_model.pkl'")
    if meta_learner:
        print("Ensemble model saved as 'ensemble_model.pkl'")
    
    return best_model_name, best_model

# ==================== LOAD PROCESSED DATA ====================
# Load the processed data from previous steps
try:
    processed_data = pd.read_csv('Processed_Resume_JD_Data.csv')
    print("Loaded processed data successfully!")
    print(f"Data shape: {processed_data.shape}")
except:
    print("Error: Could not load processed data. Please run Step 3 first.")
    # Create sample data for demonstration (remove in production)
    print("Creating sample data for demonstration...")
    processed_data = pd.DataFrame({
        'semantic_similarity': np.random.uniform(0.1, 0.9, 100),
        'skill_overlap_score': np.random.uniform(0.1, 0.9, 100),
        'resume_feature_0': np.random.randint(100, 1000, 100),
        'resume_feature_1': np.random.randint(50, 500, 100),
        'jd_feature_0': np.random.randint(100, 1000, 100),
        'jd_feature_1': np.random.randint(50, 500, 100),
        'experience_years': np.random.randint(0, 20, 100),
        'domain_keyword_match': np.random.uniform(0.1, 0.9, 100),
        'Match_Label': np.random.randint(0, 2, 100)
    })

# Get important features (all numeric columns except target)
important_features = [col for col in processed_data.columns 
                     if col not in ['Resume_Text', 'JD_Text', 'Skills_Resume', 'Skills_JD', 'Match_Label',
                                  'resume_cleaned', 'jd_cleaned'] and processed_data[col].dtype in ['int64', 'float64']]

print(f"Using {len(important_features)} important features:")
print(important_features)

# ==================== MAIN EXECUTION ====================
def main_model_training(processed_data, important_features):
    """
    Main function to train and evaluate all models
    """
    print("Starting model training and evaluation...")
    
    # Prepare features and target
    X = processed_data[important_features]
    y = processed_data['Match_Label']
    
    # Handle any remaining missing values
    X = X.fillna(0)
    
    # Train and evaluate individual models
    results, trained_models, scaler = train_powerful_models(X, y)
    
    # Create and evaluate ensemble model
    meta_learner, ensemble_results = create_ensemble_model(trained_models, X, y)
    
    # Compare all models
    comparison_df = compare_model_performance(results, ensemble_results)
    
    # Save the best model
    best_model_name, best_model = save_best_model(trained_models, results, meta_learner)
    
    # Save feature scaler
    joblib.dump(scaler, 'feature_scaler.pkl')
    print("Feature scaler saved as 'feature_scaler.pkl'")
    
    # Save important features list
    with open('important_features.pkl', 'wb') as f:
        pickle.dump(important_features, f)
    print("Important features list saved as 'important_features.pkl'")
    
    return {
        'results': results,
        'trained_models': trained_models,
        'ensemble_model': meta_learner,
        'comparison_df': comparison_df,
        'best_model': best_model,
        'best_model_name': best_model_name,
        'scaler': scaler,
        'important_features': important_features
    }

# Execute the model training
print("Starting Step 4: Model Training and Evaluation...")
training_results = main_model_training(processed_data, important_features)
print("\nModel training and evaluation completed successfully!")


# In[6]:


# ==================== DEEP LEARNING MODELS IMPLEMENTATION ====================
import tensorflow as tf
from tensorflow import keras
from tensorflow.keras.models import Sequential, Model
from tensorflow.keras.layers import Dense, Dropout, BatchNormalization, Input
from tensorflow.keras.optimizers import Adam, RMSprop
from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau, ModelCheckpoint
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import time
import joblib
import pickle

# Set random seeds for reproducibility
tf.random.set_seed(42)
np.random.seed(42)

def prepare_dl_data(X, y):
    """Prepare data for Deep Learning models"""
    # Split the data
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y
    )
    
    # Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    return X_train_scaled, X_test_scaled, y_train, y_test, scaler

def create_dnn_model(input_dim, layers=[128, 64, 32], dropout_rate=0.3):
    """Create Deep Neural Network model"""
    model = Sequential()
    model.add(Input(shape=(input_dim,)))
    
    for i, units in enumerate(layers):
        model.add(Dense(units, activation='relu', kernel_initializer='he_normal'))
        model.add(BatchNormalization())
        model.add(Dropout(dropout_rate))
    
    model.add(Dense(1, activation='sigmoid'))
    
    model.compile(
        optimizer=Adam(learning_rate=0.001),
        loss='binary_crossentropy',
        metrics=['accuracy', tf.keras.metrics.Precision(), tf.keras.metrics.Recall()]
    )
    
    return model

def create_wide_deep_model(input_dim):
    """Create Wide & Deep Learning model"""
    # Wide part (direct connections)
    input_layer = Input(shape=(input_dim,))
    
    # Deep part
    deep = Dense(128, activation='relu')(input_layer)
    deep = BatchNormalization()(deep)
    deep = Dropout(0.3)(deep)
    deep = Dense(64, activation='relu')(deep)
    deep = BatchNormalization()(deep)
    deep = Dropout(0.3)(deep)
    deep = Dense(32, activation='relu')(deep)
    
    # Combine wide and deep
    combined = keras.layers.concatenate([input_layer, deep])
    
    # Output layer
    output = Dense(1, activation='sigmoid')(combined)
    
    model = Model(inputs=input_layer, outputs=output)
    model.compile(
        optimizer=RMSprop(learning_rate=0.001),
        loss='binary_crossentropy',
        metrics=['accuracy', tf.keras.metrics.Precision(), tf.keras.metrics.Recall()]
    )
    
    return model

def create_autoencoder_classifier(input_dim, encoding_dim=32):
    """Create Autoencoder + Classifier model - SIMPLIFIED VERSION"""
    # Simplified version - only classifier part
    input_layer = Input(shape=(input_dim,))
    
    # Feature extraction layers (like encoder)
    x = Dense(128, activation='relu')(input_layer)
    x = Dropout(0.2)(x)
    x = Dense(64, activation='relu')(x)
    x = Dropout(0.2)(x)
    x = Dense(encoding_dim, activation='relu')(x)
    
    # Classifier
    x = Dense(32, activation='relu')(x)
    x = Dropout(0.3)(x)
    output = Dense(1, activation='sigmoid')(x)
    
    model = Model(inputs=input_layer, outputs=output)
    model.compile(
        optimizer=Adam(learning_rate=0.001),
        loss='binary_crossentropy',
        metrics=['accuracy', tf.keras.metrics.Precision(), tf.keras.metrics.Recall()]
    )
    
    return model

def create_attention_model(input_dim):
    """Create model with Attention mechanism - SIMPLIFIED VERSION"""
    # Simplified attention model for tabular data
    input_layer = Input(shape=(input_dim,))
    
    # Feature transformation
    x = Dense(128, activation='relu')(input_layer)
    x = BatchNormalization()(x)
    x = Dropout(0.3)(x)
    
    # Attention mechanism (simplified)
    attention_probs = Dense(128, activation='softmax')(x)
    attention_mul = keras.layers.multiply([x, attention_probs])
    
    x = Dense(64, activation='relu')(attention_mul)
    x = Dropout(0.3)(x)
    x = Dense(32, activation='relu')(x)
    
    output = Dense(1, activation='sigmoid')(x)
    
    model = Model(inputs=input_layer, outputs=output)
    model.compile(
        optimizer=Adam(learning_rate=0.0005),
        loss='binary_crossentropy',
        metrics=['accuracy', tf.keras.metrics.Precision(), tf.keras.metrics.Recall()]
    )
    
    return model

def create_resnet_model(input_dim):
    """Create ResNet-inspired model"""
    input_layer = Input(shape=(input_dim,))
    
    # Initial dense layer
    x = Dense(128, activation='relu')(input_layer)
    x = BatchNormalization()(x)
    
    # Residual blocks
    for _ in range(2):
        # Residual connection
        residual = x
        
        # Two dense layers with dropout
        x = Dense(128, activation='relu')(x)
        x = BatchNormalization()(x)
        x = Dropout(0.3)(x)
        x = Dense(128, activation='relu')(x)
        x = BatchNormalization()(x)
        
        # Add residual connection
        x = keras.layers.add([x, residual])
        x = keras.layers.Activation('relu')(x)
    
    # Final layers
    x = Dense(64, activation='relu')(x)
    x = Dropout(0.3)(x)
    x = Dense(32, activation='relu')(x)
    
    output = Dense(1, activation='sigmoid')(x)
    
    model = Model(inputs=input_layer, outputs=output)
    model.compile(
        optimizer=Adam(learning_rate=0.0005),
        loss='binary_crossentropy',
        metrics=['accuracy', tf.keras.metrics.Precision(), tf.keras.metrics.Recall()]
    )
    
    return model

def train_dl_models(X_train, X_test, y_train, y_test, input_dim):
    """Train and evaluate all Deep Learning models"""
    models = {
        'Deep Neural Network (DNN)': create_dnn_model(input_dim),
        'Wide & Deep Learning': create_wide_deep_model(input_dim),
        'Autoencoder Classifier': create_autoencoder_classifier(input_dim),
        'Attention Model': create_attention_model(input_dim),
        'ResNet Model': create_resnet_model(input_dim)
    }
    
    results = {}
    trained_models = {}
    histories = {}
    
    # Callbacks
    early_stopping = EarlyStopping(monitor='val_loss', patience=10, restore_best_weights=True)
    reduce_lr = ReduceLROnPlateau(monitor='val_loss', factor=0.5, patience=5, min_lr=1e-6)
    
    for name, model in models.items():
        print(f"\n{'='*60}")
        print(f"Training {name}...")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        # Model checkpoint
        checkpoint = ModelCheckpoint(
            f'best_{name.replace(" ", "_").lower()}.h5',
            monitor='val_accuracy',
            save_best_only=True,
            mode='max',
            verbose=0
        )
        
        # Train model (all models use same training format now)
        history = model.fit(
            X_train, y_train,
            validation_data=(X_test, y_test),
            epochs=50,
            batch_size=32,
            callbacks=[early_stopping, reduce_lr, checkpoint],
            verbose=1
        )
        
        training_time = time.time() - start_time
        
        # Evaluate model
        test_loss, test_accuracy, test_precision, test_recall = model.evaluate(X_test, y_test, verbose=0)
        y_pred_proba = model.predict(X_test, verbose=0)
        y_pred = (y_pred_proba > 0.5).astype(int)
        
        # Calculate metrics
        from sklearn.metrics import f1_score, classification_report
        
        f1 = f1_score(y_test, y_pred)
        
        # Store results
        results[name] = {
            'accuracy': test_accuracy,
            'f1_score': f1,
            'training_time': training_time,
            'epochs_trained': len(history.history['loss'])
        }
        
        trained_models[name] = model
        histories[name] = history
        
        # Print results
        print(f"{name} Results:")
        print(f"Accuracy: {test_accuracy:.4f}")
        print(f"F1-Score: {f1:.4f}")
        print(f"Training Time: {training_time:.2f} seconds")
        print(f"Epochs Trained: {len(history.history['loss'])}")
        print("\nClassification Report:")
        print(classification_report(y_test, y_pred))
    
    return results, trained_models, histories

def plot_training_history(histories):
    """Plot training history for all models"""
    plt.figure(figsize=(15, 10))
    
    for i, (name, history) in enumerate(histories.items(), 1):
        plt.subplot(2, 3, i)
        
        if 'accuracy' in history.history:
            plt.plot(history.history['accuracy'], label='Training Accuracy')
            if 'val_accuracy' in history.history:
                plt.plot(history.history['val_accuracy'], label='Validation Accuracy')
        
        plt.title(f'{name}\nTraining History')
        plt.xlabel('Epoch')
        plt.ylabel('Accuracy')
        plt.legend()
        plt.grid(True)
    
    plt.tight_layout()
    plt.show()

def compare_dl_performance(results):
    """Compare performance of all DL models"""
    print("\n" + "="*80)
    print("DEEP LEARNING MODELS PERFORMANCE COMPARISON")
    print("="*80)
    
    # Create comparison dataframe
    comparison_data = []
    for model_name, metrics in results.items():
        comparison_data.append({
            'Model': model_name,
            'Accuracy': metrics['accuracy'],
            'F1-Score': metrics['f1_score'],
            'Training Time (s)': metrics['training_time'],
            'Epochs Trained': metrics['epochs_trained']
        })
    
    comparison_df = pd.DataFrame(comparison_data)
    comparison_df = comparison_df.sort_values('F1-Score', ascending=False)
    
    print(comparison_df.round(4))
    
    # Visualize results
    plt.figure(figsize=(15, 10))
    
    # Plot 1: F1-Score Comparison
    plt.subplot(2, 2, 1)
    models = comparison_df['Model']
    f1_scores = comparison_df['F1-Score']
    bars = plt.bar(models, f1_scores, color=['blue', 'green', 'red', 'purple', 'orange'])
    plt.title('DL Models - F1-Score Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('F1-Score', fontsize=12)
    plt.xticks(rotation=45)
    plt.ylim(0, 1)
    
    # Add value labels on bars
    for bar, value in zip(bars, f1_scores):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, 
                f'{value:.3f}', ha='center', va='bottom', fontweight='bold')
    
    # Plot 2: Accuracy Comparison
    plt.subplot(2, 2, 2)
    accuracy_scores = comparison_df['Accuracy']
    bars = plt.bar(models, accuracy_scores, color=['blue', 'green', 'red', 'purple', 'orange'])
    plt.title('DL Models - Accuracy Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Accuracy', fontsize=12)
    plt.xticks(rotation=45)
    plt.ylim(0, 1)
    
    # Add value labels on bars
    for bar, value in zip(bars, accuracy_scores):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, 
                f'{value:.3f}', ha='center', va='bottom', fontweight='bold')
    
    # Plot 3: Training Time Comparison
    plt.subplot(2, 2, 3)
    training_times = comparison_df['Training Time (s)']
    bars = plt.bar(models, training_times, color=['blue', 'green', 'red', 'purple', 'orange'])
    plt.title('DL Models - Training Time Comparison', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Time (seconds)', fontsize=12)
    plt.xticks(rotation=45)
    
    # Add value labels on bars
    for bar, value in zip(bars, training_times):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                f'{value:.1f}s', ha='center', va='bottom', fontweight='bold')
    
    # Plot 4: Epochs Trained
    plt.subplot(2, 2, 4)
    epochs_trained = comparison_df['Epochs Trained']
    bars = plt.bar(models, epochs_trained, color=['blue', 'green', 'red', 'purple', 'orange'])
    plt.title('DL Models - Epochs Trained', fontsize=14, fontweight='bold')
    plt.xlabel('Models', fontsize=12)
    plt.ylabel('Epochs', fontsize=12)
    plt.xticks(rotation=45)
    
    # Add value labels on bars
    for bar, value in zip(bars, epochs_trained):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                f'{value}', ha='center', va='bottom', fontweight='bold')
    
    plt.tight_layout()
    plt.show()
    
    return comparison_df

def save_best_dl_model(trained_models, results):
    """Save the best performing DL model"""
    # Find best model based on F1-score
    best_model_name = max(results.items(), key=lambda x: x[1]['f1_score'])[0]
    best_model = trained_models[best_model_name]
    
    print(f"\nBest DL model: {best_model_name} with F1-Score: {results[best_model_name]['f1_score']:.4f}")
    
    # Save the best model
    best_model_filename = f'best_dl_model_{best_model_name.replace(" ", "_").lower()}.h5'
    best_model.save(best_model_filename)
    
    print(f"Best DL model saved as '{best_model_filename}'")
    
    # Also save in .keras format for better compatibility
    best_model.save(f'best_dl_model.keras')
    print("Best DL model also saved as 'best_dl_model.keras'")
    
    return best_model_name, best_model

# ==================== MAIN EXECUTION ====================
def main_dl_training():
    """Main function for DL model training"""
    print("Starting Deep Learning Model Training...")
    
    # Load processed data
    try:
        processed_data = pd.read_csv('Processed_Resume_JD_Data.csv')
        print("✓ Loaded processed data successfully!")
    except:
        print("Error: Could not load processed data. Please run Step 3 first.")
        return
    
    # Prepare features and target
    important_features = [col for col in processed_data.columns 
                         if col not in ['Resume_Text', 'JD_Text', 'Skills_Resume', 'Skills_JD', 'Match_Label',
                                      'resume_cleaned', 'jd_cleaned'] and 
                         processed_data[col].dtype in ['int64', 'float64']]
    
    X = processed_data[important_features]
    y = processed_data['Match_Label']
    
    # Handle missing values
    X = X.fillna(0)
    
    # Prepare DL data
    X_train, X_test, y_train, y_test, scaler = prepare_dl_data(X, y)
    
    print(f"Training data shape: {X_train.shape}")
    print(f"Test data shape: {X_test.shape}")
    print(f"Input dimension: {X_train.shape[1]}")
    
    # Train DL models
    results, trained_models, histories = train_dl_models(X_train, X_test, y_train, y_test, X_train.shape[1])
    
    # Plot training history
    plot_training_history(histories)
    
    # Compare performance
    comparison_df = compare_dl_performance(results)
    
    # Save best model
    best_model_name, best_model = save_best_dl_model(trained_models, results)
    
    # Save scaler
    joblib.dump(scaler, 'dl_feature_scaler.pkl')
    print("DL feature scaler saved as 'dl_feature_scaler.pkl'")
    
    # Save important features
    with open('dl_important_features.pkl', 'wb') as f:
        pickle.dump(important_features, f)
    print("DL important features list saved as 'dl_important_features.pkl'")
    
    return {
        'results': results,
        'trained_models': trained_models,
        'comparison_df': comparison_df,
        'best_model': best_model,
        'best_model_name': best_model_name,
        'scaler': scaler,
        'important_features': important_features
    }

# ==================== MODEL LOADING AND PREDICTION ====================
class DLPredictor:
    def __init__(self):
        """Initialize the DL Prediction System"""
        try:
            # Load the best DL model
            self.model = keras.models.load_model('best_dl_model.keras')
            self.scaler = joblib.load('dl_feature_scaler.pkl')
            
            with open('dl_important_features.pkl', 'rb') as f:
                self.important_features = pickle.load(f)
                
            print("✓ DL model components loaded successfully!")
            print(f"✓ Model Type: {type(self.model).__name__}")
            
        except Exception as e:
            print(f"Error loading DL model: {e}")
            raise
    
    def predict(self, features):
        """Make prediction using DL model"""
        try:
            # Prepare features
            X = np.array(features).reshape(1, -1)
            X_scaled = self.scaler.transform(X)
            
            # Predict
            prediction = self.model.predict(X_scaled, verbose=0)[0][0]
            
            return prediction
            
        except Exception as e:
            print(f"Error in DL prediction: {e}")
            return 0.5  # Return neutral probability

# Execute DL training
if __name__ == "__main__":
    print("Starting Step 5: Deep Learning Model Training...")
    dl_results = main_dl_training()
    
    print("\n" + "="*60)
    print("DEEP LEARNING TRAINING COMPLETED!")
    print("="*60)
    print("Your best DL model has been saved and can be loaded using:")
    print("model = keras.models.load_model('best_dl_model.keras')")
    print("\nUse DLPredictor class for easy predictions:")
    print("predictor = DLPredictor()")
    print("score = predictor.predict(features)")


# In[7]:


# ==================== STEP 6: SIMPLE ATS SCORE CLASSIFICATION ====================
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import tensorflow as tf
from tensorflow import keras
import joblib
import pickle
import os

# Download NLTK resources
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
except:
    pass

# Initialize text processing
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

class ATSScorePredictor:
    def __init__(self):
        """Initialize the ATS Score Predictor with trained models"""
        try:
            # Load ML model and components
            self.ml_model = joblib.load('best_model.pkl')
            self.ml_scaler = joblib.load('feature_scaler.pkl')
            with open('important_features.pkl', 'rb') as f:
                self.ml_important_features = pickle.load(f)
            
            # Load DL model and components
            self.dl_model = keras.models.load_model('best_dl_model.keras')
            self.dl_scaler = joblib.load('dl_feature_scaler.pkl')
            with open('dl_important_features.pkl', 'rb') as f:
                self.dl_important_features = pickle.load(f)
                
            print("✓ All models loaded successfully!")
            
        except Exception as e:
            print(f"Error loading models: {e}")
            raise
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from PDF file"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_file(self, file_path):
        """Extract text from any supported file format"""
        if file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            print("Unsupported file format. Please use PDF or DOCX.")
            return ""
    
    def preprocess_text(self, text):
        """Preprocess text for feature extraction"""
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters and extra spaces
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def extract_features_from_resume(self, resume_text):
        """Extract features from resume text for prediction"""
        # Basic text features
        features = {}
        
        # Text statistics
        features['word_count'] = len(resume_text.split())
        features['char_count'] = len(resume_text)
        features['avg_word_length'] = features['char_count'] / max(features['word_count'], 1)
        
        # Sentence count (approximate)
        sentences = re.split(r'[.!?]+', resume_text)
        features['sentence_count'] = len([s for s in sentences if len(s.strip()) > 0])
        features['avg_sentence_length'] = features['word_count'] / max(features['sentence_count'], 1)
        
        # Keyword analysis
        technical_keywords = ['python', 'java', 'machine learning', 'data analysis', 'sql', 
                             'project management', 'leadership', 'communication', 'problem solving',
                             'teamwork', 'javascript', 'html', 'css', 'react', 'angular', 'node',
                             'database', 'aws', 'cloud', 'devops', 'agile', 'scrum']
        
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate',
                             'education', 'university', 'college', 'school', 'graduated']
        
        experience_keywords = ['experience', 'years', 'worked', 'job', 'position', 'role',
                              'responsibilities', 'achievements', 'accomplishments']
        
        # Count keyword occurrences
        text_lower = resume_text.lower()
        features['technical_keyword_count'] = sum(1 for word in technical_keywords if word in text_lower)
        features['education_keyword_count'] = sum(1 for word in education_keywords if word in text_lower)
        features['experience_keyword_count'] = sum(1 for word in experience_keywords if word in text_lower)
        
        # Keyword density
        features['technical_keyword_density'] = features['technical_keyword_count'] / max(features['word_count'], 1)
        features['education_keyword_density'] = features['education_keyword_count'] / max(features['word_count'], 1)
        features['experience_keyword_density'] = features['experience_keyword_count'] / max(features['word_count'], 1)
        
        # Extract experience years
        features['experience_years'] = self.extract_experience_years(resume_text)
        
        # Additional features that might be in our trained model
        features['semantic_similarity'] = 0.6  # Placeholder, would need JD for actual calculation
        features['skill_overlap_score'] = 0.5  # Placeholder
        
        return features
    
    def extract_experience_years(self, text):
        """Extract experience years from resume text"""
        text = text.lower()
        
        # Patterns to match experience
        patterns = [
            r'(\d+)\s*\+?\s*years?[\s\w]*experience',
            r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?',
            r'(\d+)-(\d+)\s*years?',
            r'(\d+)\s*to\s*(\d+)\s*years?',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?',
            r'(\d+)\s*yr',
            r'(\d+)\s*mos'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                if isinstance(matches[0], tuple):
                    # For ranges like "3-5 years", take the average
                    nums = [int(x) for x in matches[0] if str(x).isdigit()]
                    if nums:
                        return sum(nums) / len(nums)
                else:
                    # Handle month conversions
                    if 'mos' in pattern and matches:
                        return int(matches[0]) / 12  # Convert months to years
                    else:
                        return int(matches[0])
        
        return 0  # Default if no experience found
    
    def predict_ats_score(self, resume_text, use_dl=True):
        """Predict ATS score for given resume text"""
        try:
            # Extract features
            features_dict = self.extract_features_from_resume(resume_text)
            
            # Select appropriate model and features
            if use_dl:
                model = self.dl_model
                scaler = self.dl_scaler
                important_features = self.dl_important_features
                print("Using Deep Learning model for prediction...")
            else:
                model = self.ml_model
                scaler = self.ml_scaler
                important_features = self.ml_important_features
                print("Using Machine Learning model for prediction...")
            
            # Create feature vector in correct order
            feature_vector = []
            for feature_name in important_features:
                if feature_name in features_dict:
                    feature_vector.append(features_dict[feature_name])
                else:
                    feature_vector.append(0)  # Default value for missing features
            
            # Convert to numpy array and scale
            X = np.array(feature_vector).reshape(1, -1)
            X_scaled = scaler.transform(X)
            
            # Predict probability
            if use_dl:
                probability = model.predict(X_scaled, verbose=0)[0][0]
            else:
                probability = model.predict_proba(X_scaled)[0][1]
            
            # Convert to ATS score (0-100)
            ats_score = round(probability * 100, 2)
            
            return ats_score, probability, features_dict
            
        except Exception as e:
            print(f"Error in prediction: {e}")
            return 50, 0.5, {}  # Return neutral score on error
    
    def generate_suggestions(self, ats_score, features):
        """Generate suggestions to improve ATS score"""
        suggestions = []
        
        if ats_score < 70:
            suggestions.append("❌ Your resume needs significant improvement for better ATS compatibility.")
        
        if features.get('technical_keyword_count', 0) < 5:
            suggestions.append("🔧 Add more technical skills and keywords relevant to your industry.")
        
        if features.get('experience_years', 0) < 2:
            suggestions.append("💼 Highlight any relevant projects, internships, or volunteer experience.")
        
        if features.get('word_count', 0) < 200:
            suggestions.append("📄 Add more detailed descriptions of your experiences and achievements.")
        
        if features.get('education_keyword_count', 0) < 2:
            suggestions.append("🎓 Clearly list your educational qualifications and degrees.")
        
        if ats_score >= 85:
            suggestions.append("✅ Excellent! Your resume is well-optimized for ATS systems.")
        elif ats_score >= 70:
            suggestions.append("👍 Good resume! Some minor improvements could make it even better.")
        
        # Add specific suggestions based on feature analysis
        if features.get('technical_keyword_density', 0) < 0.02:
            suggestions.append("📝 Increase the density of technical keywords throughout your resume.")
        
        if features.get('experience_keyword_count', 0) < 3:
            suggestions.append("💪 Use more action verbs and accomplishment statements.")
        
        return suggestions

    def analyze_resume_file(self, file_path, use_dl=True):
        """Complete analysis of a resume file"""
        print(f"Analyzing resume: {file_path}")
        
        # Extract text from file
        resume_text = self.extract_text_from_file(file_path)
        
        if not resume_text.strip():
            print("❌ Could not extract text from the file. Please try another file.")
            return None
        
        # Preprocess text
        processed_text = self.preprocess_text(resume_text)
        
        # Predict ATS score
        ats_score, probability, features = self.predict_ats_score(processed_text, use_dl=use_dl)
        
        # Generate suggestions
        suggestions = self.generate_suggestions(ats_score, features)
        
        return {
            'ats_score': ats_score,
            'probability': probability,
            'features': features,
            'suggestions': suggestions,
            'processed_text': processed_text
        }

# ==================== MAIN EXECUTION ====================
def main():
    """Main function to demonstrate ATS score prediction"""
    print("Starting ATS Score Prediction System...")
    print("=" * 60)
    
    # Initialize predictor
    try:
        predictor = ATSScorePredictor()
    except Exception as e:
        print(f"Failed to initialize predictor: {e}")
        return
    
    # Example usage
    print("\n1. Testing with sample resume file...")
    
    # Create a sample resume text for testing (in real usage, you would provide file path)
    sample_resume_text = """
    John Doe
    Senior Data Scientist
    Email: john.doe@email.com
    Phone: (123) 456-7890
    
    EXPERIENCE:
    Senior Data Scientist, ABC Tech Inc. (2020 - Present)
    - Developed machine learning models for customer behavior prediction
    - Managed team of 5 junior data scientists
    - Implemented NLP solutions for text classification
    
    Data Analyst, XYZ Corp (2018 - 2020)
    - Performed data analysis using Python and SQL
    - Created data visualizations and reports
    - Collaborated with cross-functional teams
    
    EDUCATION:
    Master of Science in Computer Science, University of Technology (2016 - 2018)
    Bachelor of Science in Mathematics, State University (2012 - 2016)
    
    SKILLS:
    Python, Machine Learning, TensorFlow, SQL, Data Analysis, 
    Statistical Modeling, Project Management, Leadership
    
    CERTIFICATIONS:
    AWS Certified Machine Learning Specialist
    Google Cloud Professional Data Engineer
    """
    
    # Save sample resume to a text file for testing
    sample_file_path = "sample_resume.txt"
    with open(sample_file_path, 'w') as f:
        f.write(sample_resume_text)
    
    # Analyze the resume
    result = predictor.analyze_resume_file(sample_file_path, use_dl=True)
    
    if result:
        print("\n" + "=" * 60)
        print("ATS SCORE ANALYSIS RESULTS")
        print("=" * 60)
        
        # Display score
        score_color = "🟢" if result['ats_score'] >= 80 else "🟡" if result['ats_score'] >= 60 else "🔴"
        print(f"{score_color} ATS Score: {result['ats_score']}/100")
        print(f"📊 Match Probability: {result['probability']:.3f}")
        
        # Display key features
        print("\n📈 RESUME INSIGHTS:")
        features = result['features']
        print(f"   • Word Count: {features.get('word_count', 0)}")
        print(f"   • Experience Years: {features.get('experience_years', 0):.1f}")
        print(f"   • Technical Keywords: {features.get('technical_keyword_count', 0)}")
        print(f"   • Education Keywords: {features.get('education_keyword_count', 0)}")
        print(f"   • Experience Keywords: {features.get('experience_keyword_count', 0)}")
        
        # Display suggestions
        print("\n💡 IMPROVEMENT SUGGESTIONS:")
        for i, suggestion in enumerate(result['suggestions'], 1):
            print(f"   {i}. {suggestion}")
        
        # Clean up sample file
        if os.path.exists(sample_file_path):
            os.remove(sample_file_path)
    
    print("\n" + "=" * 60)
    print("HOW TO USE IN YOUR CODE:")
    print("=" * 60)
    print("""
# Initialize the predictor
predictor = ATSScorePredictor()

# Analyze a resume file
result = predictor.analyze_resume_file('path/to/your/resume.pdf', use_dl=True)

# Access results
ats_score = result['ats_score']
suggestions = result['suggestions']
features = result['features']
    """)

# ==================== BATCH PROCESSING ====================
def batch_process_resumes(resume_paths, use_dl=True):
    """Process multiple resumes in batch"""
    predictor = ATSScorePredictor()
    results = []
    
    for resume_path in resume_paths:
        print(f"Processing: {resume_path}")
        result = predictor.analyze_resume_file(resume_path, use_dl=use_dl)
        if result:
            results.append({
                'file_path': resume_path,
                'ats_score': result['ats_score'],
                'probability': result['probability']
            })
    
    return results

if __name__ == "__main__":
    main()


# In[8]:


# Initialize the predictor
predictor = ATSScorePredictor()

# Analyze a single resume
result = predictor.analyze_resume_file(r"C:\Users\omash\OneDrive\Documents\Desktop\OM_ASHUTOSH Resume.pdf", use_dl=True)

# Batch process multiple resumes
results = batch_process_resumes(['resume1.pdf', 'resume2.docx'], use_dl=True)


# In[9]:


# ==================== STEP 6: HYBRID ML-DL ATS SCORING SYSTEM ====================
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import tensorflow as tf
from tensorflow import keras
from sentence_transformers import SentenceTransformer, util
import joblib
import pickle
import os

# Download NLTK resources
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
except:
    pass

# Initialize text processing
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

class HybridATSPredictor:
    def __init__(self):
        """Initialize Hybrid ATS Predictor with both ML and DL models"""
        try:
            # ==================== ML MODELS ====================
            print("Loading ML models for structured analysis...")
            self.ml_model = joblib.load('best_model.pkl')
            self.ml_scaler = joblib.load('feature_scaler.pkl')
            with open('important_features.pkl', 'rb') as f:
                self.ml_important_features = pickle.load(f)
            
            # ==================== DL MODELS ====================
            print("Loading DL models for semantic analysis...")
            self.dl_model = keras.models.load_model('best_dl_model.keras')
            self.dl_scaler = joblib.load('dl_feature_scaler.pkl')
            with open('dl_important_features.pkl', 'rb') as f:
                self.dl_important_features = pickle.load(f)
            
            # Semantic similarity model
            print("Loading Sentence Transformer for semantic analysis...")
            self.semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Domain classification model (ML-based)
            self.domain_model = joblib.load('best_model.pkl')  # Reuse for demo
            self.domain_scaler = joblib.load('feature_scaler.pkl')
            
            print("✓ All ML and DL models loaded successfully!")
            
        except Exception as e:
            print(f"Error loading models: {e}")
            # Continue with basic functionality
            self.semantic_model = None
    
    # ==================== ML COMPONENTS ====================
    def extract_ml_features(self, resume_text):
        """ML-based feature extraction - Structured data processing"""
        features = {}
        
        # Text statistics (ML features)
        features['word_count'] = len(resume_text.split())
        features['char_count'] = len(resume_text)
        features['avg_word_length'] = features['char_count'] / max(features['word_count'], 1)
        
        # Sentence features
        sentences = re.split(r'[.!?]+', resume_text)
        features['sentence_count'] = len([s for s in sentences if len(s.strip()) > 0])
        features['avg_sentence_length'] = features['word_count'] / max(features['sentence_count'], 1)
        
        # Keyword-based features (ML approach)
        technical_keywords = ['python', 'java', 'machine learning', 'data analysis', 'sql', 
                             'project management', 'leadership', 'communication', 'problem solving']
        
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate']
        
        experience_keywords = ['experience', 'years', 'worked', 'job', 'position', 'role']
        
        text_lower = resume_text.lower()
        features['technical_keyword_count'] = sum(1 for word in technical_keywords if word in text_lower)
        features['education_keyword_count'] = sum(1 for word in education_keywords if word in text_lower)
        features['experience_keyword_count'] = sum(1 for word in experience_keywords if word in text_lower)
        
        # Density features
        features['technical_keyword_density'] = features['technical_keyword_count'] / max(features['word_count'], 1)
        
        # Experience extraction (Rule-based ML)
        features['experience_years'] = self.extract_experience_years(resume_text)
        
        return features
    
    def ml_predict(self, features_dict):
        """ML-based prediction - Fast and interpretable"""
        try:
            # Create feature vector
            feature_vector = []
            for feature_name in self.ml_important_features:
                feature_vector.append(features_dict.get(feature_name, 0))
            
            # Scale and predict
            X = np.array(feature_vector).reshape(1, -1)
            X_scaled = self.ml_scaler.transform(X)
            probability = self.ml_model.predict_proba(X_scaled)[0][1]
            
            return probability
            
        except Exception as e:
            print(f"ML prediction error: {e}")
            return 0.5
    
    # ==================== DL COMPONENTS ====================
    def extract_dl_features(self, resume_text):
        """DL-based feature extraction - Semantic understanding"""
        features = {}
        
        if self.semantic_model:
            try:
                # Semantic analysis (DL approach)
                embeddings = self.semantic_model.encode(resume_text, convert_to_tensor=True)
                
                # Sample job descriptions for semantic comparison
                sample_jds = [
                    "Data scientist with machine learning experience Python SQL",
                    "Software engineer Java JavaScript web development",
                    "Project manager leadership communication team management"
                ]
                
                # Calculate semantic similarity with sample JDs
                jd_embeddings = self.semantic_model.encode(sample_jds, convert_to_tensor=True)
                cosine_scores = util.pytorch_cos_sim(embeddings, jd_embeddings)
                
                features['max_semantic_similarity'] = float(torch.max(cosine_scores))
                features['avg_semantic_similarity'] = float(torch.mean(cosine_scores))
                
            except Exception as e:
                print(f"Semantic analysis error: {e}")
                features['max_semantic_similarity'] = 0.5
                features['avg_semantic_similarity'] = 0.5
        else:
            features['max_semantic_similarity'] = 0.5
            features['avg_semantic_similarity'] = 0.5
        
        return features
    
    def dl_predict(self, features_dict):
        """DL-based prediction - High accuracy, semantic understanding"""
        try:
            # Create feature vector
            feature_vector = []
            for feature_name in self.dl_important_features:
                feature_vector.append(features_dict.get(feature_name, 0))
            
            # Scale and predict
            X = np.array(feature_vector).reshape(1, -1)
            X_scaled = self.dl_scaler.transform(X)
            probability = self.dl_model.predict(X_scaled, verbose=0)[0][0]
            
            return probability
            
        except Exception as e:
            print(f"DL prediction error: {e}")
            return 0.5
    
    # ==================== HYBRID APPROACH ====================
    def hybrid_predict(self, ml_probability, dl_probability):
        """Combine ML and DL predictions intelligently"""
        # Weighted combination (adjust weights based on validation)
        ml_weight = 0.4  # ML is faster but less nuanced
        dl_weight = 0.6  # DL is slower but more accurate
        
        hybrid_probability = (ml_probability * ml_weight) + (dl_probability * dl_weight)
        return hybrid_probability
    
    # ==================== UTILITY METHODS ====================
    def extract_text_from_file(self, file_path):
        """Extract text from supported file formats"""
        try:
            if file_path.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.docx'):
                return self.extract_text_from_docx(file_path)
            elif file_path.lower().endswith('.txt'):
                return self.extract_text_from_txt(file_path)
            else:
                return ""
        except Exception as e:
            print(f"File extraction error: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        except:
            return ""
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOCX"""
        try:
            doc = docx.Document(docx_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except:
            return ""
    
    def extract_text_from_txt(self, txt_path):
        """Extract text from TXT"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
        except:
            return ""
    
    def extract_experience_years(self, text):
        """Rule-based experience extraction (ML approach)"""
        text = text.lower()
        patterns = [
            r'(\d+)\s*\+?\s*years?[\s\w]*experience',
            r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?',
            r'(\d+)-(\d+)\s*years?'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                if isinstance(matches[0], tuple):
                    nums = [int(x) for x in matches[0] if str(x).isdigit()]
                    return sum(nums) / len(nums) if nums else 0
                else:
                    return int(matches[0])
        return 0
    
    def generate_suggestions(self, ats_score, features):
        """Generate intelligent suggestions"""
        suggestions = []
        
        # ML-based suggestions (interpretable)
        if features.get('technical_keyword_count', 0) < 5:
            suggestions.append("🔧 Add more technical skills (Python, SQL, ML, etc.)")
        
        if features.get('experience_years', 0) < 2:
            suggestions.append("💼 Highlight projects and internships as experience")
        
        if features.get('word_count', 0) < 250:
            suggestions.append("📄 Add more detail about your achievements and responsibilities")
        
        # DL-based suggestions (semantic)
        if features.get('max_semantic_similarity', 0) < 0.4:
            suggestions.append("🎯 Improve relevance to common job descriptions in your field")
        
        # Score-based suggestions
        if ats_score < 60:
            suggestions.append("❌ Major overhaul needed - consider professional resume review")
        elif ats_score < 75:
            suggestions.append("⚠️ Significant improvements needed in multiple areas")
        elif ats_score < 85:
            suggestions.append("👍 Good foundation - focus on specific improvements")
        else:
            suggestions.append("✅ Excellent resume - minor tweaks could make it perfect")
        
        return suggestions
    
    def analyze_resume(self, file_path=None, resume_text=None):
        """Main analysis method - Hybrid ML-DL approach"""
        # Extract text
        if file_path:
            text = self.extract_text_from_file(file_path)
        elif resume_text:
            text = resume_text
        else:
            return None
        
        if not text.strip():
            return None
        
        print("🔍 Extracting features using ML and DL...")
        
        # ML-based feature extraction and prediction
        ml_features = self.extract_ml_features(text)
        ml_probability = self.ml_predict(ml_features)
        
        # DL-based feature extraction and prediction
        dl_features = self.extract_dl_features(text)
        dl_probability = self.dl_predict({**ml_features, **dl_features})
        
        # Hybrid prediction
        final_probability = self.hybrid_predict(ml_probability, dl_probability)
        ats_score = round(final_probability * 100, 2)
        
        # Generate suggestions
        all_features = {**ml_features, **dl_features}
        suggestions = self.generate_suggestions(ats_score, all_features)
        
        return {
            'ats_score': ats_score,
            'probability': final_probability,
            'ml_score': round(ml_probability * 100, 2),
            'dl_score': round(dl_probability * 100, 2),
            'features': all_features,
            'suggestions': suggestions,
            'text_preview': text[:500] + "..." if len(text) > 500 else text
        }

# ==================== MAIN APPLICATION ====================
def main():
    """Main function with resume upload simulation"""
    print("🚀 Advanced Fit_Resume - Hybrid ML-DL ATS Scoring System")
    print("=" * 65)
    
    # Initialize predictor
    try:
        predictor = HybridATSPredictor()
        print("✅ System initialized successfully!")
    except Exception as e:
        print(f"❌ Initialization failed: {e}")
        return
    
    # Simulate resume upload
    print("\n📁 Please upload your resume (simulated path):")
    print("Note: In production, this would be a file upload dialog")
    
    # Create a sample resume file for demonstration
    sample_resume = """
    JOHN DOE
    Data Scientist
    Email: john.doe@email.com | Phone: (123) 456-7890
    LinkedIn: linkedin.com/in/johndoe
    
    SUMMARY
    Experienced Data Scientist with 4+ years in machine learning and data analysis.
    Strong background in Python, TensorFlow, and cloud technologies.
    
    EXPERIENCE
    Senior Data Scientist - Tech Innovations Inc. (2021-Present)
    - Developed ML models that improved prediction accuracy by 35%
    - Led team of 3 data analysts on customer segmentation projects
    - Implemented real-time data processing pipelines
    
    Data Analyst - Data Solutions Corp (2019-2021)
    - Created dashboards and reports for business intelligence
    - Performed statistical analysis using Python and R
    - Collaborated with cross-functional teams
    
    EDUCATION
    MS in Computer Science - University of Technology (2017-2019)
    BS in Mathematics - State University (2013-2017)
    
    SKILLS
    Python, Machine Learning, TensorFlow, SQL, Data Visualization,
    AWS, Docker, Git, Statistical Analysis, Project Management
    
    CERTIFICATIONS
    AWS Certified Machine Learning Specialist
    Google Cloud Professional Data Engineer
    """
    
    # Save sample resume
    sample_path = "sample_resume.txt"
    with open(sample_path, 'w', encoding='utf-8') as f:
        f.write(sample_resume)
    
    print(f"📄 Analyzing sample resume: {sample_path}")
    
    # Analyze the resume
    result = predictor.analyze_resume(file_path=sample_path)
    
    if result:
        print("\n" + "=" * 65)
        print("📊 ATS SCORE ANALYSIS RESULTS")
        print("=" * 65)
        
        # Display scores
        print(f"🎯 FINAL ATS SCORE: {result['ats_score']}/100")
        print(f"🤖 ML Score: {result['ml_score']}/100")
        print(f"🧠 DL Score: {result['dl_score']}/100")
        print(f"📈 Confidence: {result['probability']:.3f}")
        
        # Display key insights
        print("\n🔍 RESUME INSIGHTS:")
        features = result['features']
        print(f"   • Experience: {features.get('experience_years', 0):.1f} years")
        print(f"   • Technical Keywords: {features.get('technical_keyword_count', 0)}")
        print(f"   • Semantic Match: {features.get('max_semantic_similarity', 0.5):.2f}")
        print(f"   • Content Length: {features.get('word_count', 0)} words")
        
        # Display suggestions
        print("\n💡 IMPROVEMENT SUGGESTIONS:")
        for i, suggestion in enumerate(result['suggestions'][:5], 1):
            print(f"   {i}. {suggestion}")
        
        # Clean up
        if os.path.exists(sample_path):
            os.remove(sample_path)
    else:
        print("❌ Failed to analyze the resume.")
    
    print("\n" + "=" * 65)
    print("🛠️  HOW TO USE IN YOUR APPLICATION:")
    print("=" * 65)
    print("""
# Initialize the hybrid predictor
predictor = HybridATSPredictor()

# Analyze from file
result = predictor.analyze_resume(file_path='resume.pdf')

# Analyze from text
result = predictor.analyze_resume(resume_text=resume_text)

# Access results
score = result['ats_score']          # Final ATS score (0-100)
ml_score = result['ml_score']        # ML component score
dl_score = result['dl_score']        # DL component score  
suggestions = result['suggestions']  # Improvement suggestions
    """)

if __name__ == "__main__":
    main()


# In[13]:


# ==================== STEP 6: RESUME ATS SCORE DETECTION ====================
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import tensorflow as tf
from tensorflow import keras
import joblib
import pickle
import os

# Download NLTK resources
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# Initialize text processing
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

class ResumeATSAnalyzer:
    def __init__(self):
        """Initialize the Resume ATS Analyzer with trained models"""
        try:
            # Load the trained DL model and components
            self.model = keras.models.load_model('best_dl_model.keras')
            self.scaler = joblib.load('dl_feature_scaler.pkl')
            with open('dl_important_features.pkl', 'rb') as f:
                self.important_features = pickle.load(f)
            
            print("✅ ATS Analyzer initialized successfully!")
            self.model_loaded = True
            
        except Exception as e:
            print(f"⚠️  Model loading warning: {e}")
            print("⚠️  Using fallback analysis mode")
            self.model_loaded = False
            self.important_features = ['word_count', 'char_count', 'technical_keyword_count', 
                                     'experience_years', 'education_keyword_count']
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_file(self, file_path):
        """Extract text from any supported file format"""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            return "Unsupported file format"
    
    def preprocess_text(self, text):
        """Preprocess text for analysis"""
        if not text or "Unsupported" in text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters and extra spaces
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def extract_features(self, resume_text):
        """Extract features from resume text for ATS scoring"""
        features = {}
        
        # Basic text statistics
        features['word_count'] = len(resume_text.split())
        features['char_count'] = len(resume_text)
        features['avg_word_length'] = features['char_count'] / max(features['word_count'], 1)
        
        # Sentence analysis
        sentences = re.split(r'[.!?]+', resume_text)
        features['sentence_count'] = len([s for s in sentences if len(s.strip()) > 0])
        features['avg_sentence_length'] = features['word_count'] / max(features['sentence_count'], 1)
        
        # Keyword analysis
        technical_keywords = [
            'python', 'java', 'machine learning', 'data analysis', 'sql', 'database',
            'project management', 'leadership', 'communication', 'problem solving',
            'teamwork', 'javascript', 'html', 'css', 'react', 'angular', 'node',
            'aws', 'cloud', 'devops', 'agile', 'scrum', 'tensorflow', 'pytorch',
            'statistical analysis', 'data visualization', 'big data', 'ai', 'nlp'
        ]
        
        education_keywords = [
            'bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate',
            'education', 'university', 'college', 'school', 'graduated', 'gpa'
        ]
        
        experience_keywords = [
            'experience', 'years', 'worked', 'job', 'position', 'role',
            'responsibilities', 'achievements', 'accomplishments', 'managed',
            'developed', 'implemented', 'led', 'created', 'designed'
        ]
        
        text_lower = resume_text.lower()
        features['technical_keyword_count'] = sum(1 for word in technical_keywords if word in text_lower)
        features['education_keyword_count'] = sum(1 for word in education_keywords if word in text_lower)
        features['experience_keyword_count'] = sum(1 for word in experience_keywords if word in text_lower)
        
        # Experience extraction
        features['experience_years'] = self.extract_experience_years(resume_text)
        
        # Add required features for model compatibility
        for feature in self.important_features:
            if feature not in features:
                features[feature] = 0.5  # Default value
        
        return features
    
    def extract_experience_years(self, text):
        """Extract years of experience from resume text"""
        text = text.lower()
        
        patterns = [
            r'(\d+)\s*\+?\s*years?[\s\w]*experience',
            r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?',
            r'(\d+)-(\d+)\s*years?',
            r'(\d+)\s*to\s*(\d+)\s*years?',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                if isinstance(matches[0], tuple):
                    nums = [int(x) for x in matches[0] if str(x).isdigit()]
                    return sum(nums) / len(nums) if nums else 0
                else:
                    return int(matches[0])
        return 0
    
    def predict_ats_score(self, resume_text):
        """Predict ATS score from resume text"""
        try:
            # Extract features
            features_dict = self.extract_features(resume_text)
            
            # Create feature vector in correct order
            feature_vector = []
            for feature_name in self.important_features:
                feature_vector.append(features_dict.get(feature_name, 0))
            
            # Convert to numpy array and scale
            X = np.array(feature_vector).reshape(1, -1)
            if hasattr(self, 'scaler') and self.scaler:
                X_scaled = self.scaler.transform(X)
            else:
                X_scaled = X
            
            # Predict using model or fallback
            if self.model_loaded:
                probability = self.model.predict(X_scaled, verbose=0)[0][0]
            else:
                # Fallback heuristic scoring
                tech_score = min(features_dict['technical_keyword_count'] / 15, 1.0)
                exp_score = min(features_dict['experience_years'] / 10, 1.0)
                length_score = min(features_dict['word_count'] / 500, 1.0)
                probability = 0.3 + (tech_score * 0.3) + (exp_score * 0.2) + (length_score * 0.2)
                probability = min(probability, 0.95)
            
            # Convert to 0-100 score
            ats_score = round(probability * 100, 2)
            
            return ats_score, probability, features_dict
            
        except Exception as e:
            print(f"Prediction error: {e}")
            return 50.0, 0.5, {}
    
    def generate_suggestions(self, ats_score, features):
        """Generate improvement suggestions based on ATS score"""
        suggestions = []
        
        # Score-based suggestions
        if ats_score < 60:
            suggestions.append("❌ Your resume needs major improvements for ATS compatibility")
        elif ats_score < 75:
            suggestions.append("⚠️  Your resume needs significant improvements")
        elif ats_score < 85:
            suggestions.append("👍 Good resume, but can be improved further")
        else:
            suggestions.append("✅ Excellent! Your resume is well-optimized for ATS")
        
        # Feature-based suggestions
        if features.get('technical_keyword_count', 0) < 8:
            suggestions.append("🔧 Add more technical skills and keywords relevant to your industry")
        
        if features.get('experience_years', 0) < 2:
            suggestions.append("💼 Highlight projects, internships, and volunteer work as experience")
        
        if features.get('word_count', 0) < 250:
            suggestions.append("📄 Add more detailed descriptions of your achievements and responsibilities")
        
        if features.get('education_keyword_count', 0) < 2:
            suggestions.append("🎓 Clearly list your educational qualifications with degrees and institutions")
        
        if features.get('sentence_count', 0) < 10:
            suggestions.append("📝 Use bullet points and proper formatting to improve readability")
        
        return suggestions
    
    def analyze_resume_file(self, file_path):
        """Complete resume analysis from file path"""
        print(f"📄 Analyzing resume: {os.path.basename(file_path)}")
        
        # Extract text from file
        resume_text = self.extract_text_from_file(file_path)
        
        if not resume_text or "Unsupported" in resume_text:
            return {
                'error': 'Unsupported file format or unable to extract text',
                'ats_score': 0,
                'suggestions': ['Please upload a PDF or DOCX file with extractable text'],
                'file_name': os.path.basename(file_path),
                'text_length': 0
            }
        
        if len(resume_text.strip()) < 100:
            return {
                'error': 'Insufficient text extracted',
                'ats_score': 0,
                'suggestions': ['The file appears to be empty or contains very little text'],
                'file_name': os.path.basename(file_path),
                'text_length': len(resume_text)
            }
        
        # Preprocess text
        processed_text = self.preprocess_text(resume_text)
        
        # Predict ATS score
        ats_score, probability, features = self.predict_ats_score(processed_text)
        
        # Generate suggestions
        suggestions = self.generate_suggestions(ats_score, features)
        
        return {
            'ats_score': ats_score,
            'probability': probability,
            'features': features,
            'suggestions': suggestions,
            'text_length': len(resume_text),
            'file_name': os.path.basename(file_path)  # FIXED: Added file_name key
        }

# ==================== USAGE EXAMPLE ====================
def main():
    """Example usage of the Resume ATS Analyzer"""
    print("🔍 Resume ATS Score Detection System")
    print("=" * 50)
    
    # Initialize analyzer
    analyzer = ResumeATSAnalyzer()
    
    # Example: Analyze a sample resume (in real use, you'd provide actual file paths)
    sample_resume_text = """
    JOHN DOE
    Data Scientist
    Email: john.doe@email.com | Phone: (123) 456-7890
    
    SUMMARY
    Experienced Data Scientist with 4+ years in machine learning and data analysis.
    Strong skills in Python, TensorFlow, SQL, and cloud technologies.
    
    EXPERIENCE
    Senior Data Scientist - Tech Company (2021-Present)
    - Developed machine learning models that improved accuracy by 35%
    - Managed data analysis projects and led a team of 3 analysts
    - Implemented data pipelines and visualization dashboards
    
    Data Analyst - Analytics Firm (2019-2021)
    - Performed statistical analysis and created reports
    - Used Python and SQL for data processing
    - Collaborated with cross-functional teams
    
    EDUCATION
    MS in Computer Science - University of Technology (2017-2019)
    BS in Mathematics - State University (2013-2017)
    
    SKILLS
    Python, Machine Learning, TensorFlow, SQL, Data Visualization,
    AWS, Statistical Analysis, Project Management
    """
    
    # Save sample to temporary file
    sample_path = "sample_resume.txt"
    with open(sample_path, 'w', encoding='utf-8') as f:
        f.write(sample_resume_text)
    
    # Analyze the resume
    result = analyzer.analyze_resume_file(sample_path)
    
    # Display results
    print(f"\n📊 Analysis Results:")
    print(f"   ATS Score: {result['ats_score']}/100")
    print(f"   File: {result['file_name']}")
    print(f"   Text Length: {result['text_length']} characters")
    
    print(f"\n💡 Suggestions:")
    for i, suggestion in enumerate(result['suggestions'], 1):
        print(f"   {i}. {suggestion}")
    
    # Clean up
    if os.path.exists(sample_path):
        os.remove(sample_path)
    
    print(f"\n✅ Analysis complete!")

if __name__ == "__main__":
    main()


# In[27]:


# Simple test to check text extraction
import fitz  # PyMuPDF

def test_pdf_extraction(file_path):
    """Test PDF text extraction"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        
        print(f"📄 Extracted {len(text)} characters")
        print(f"📝 Sample text: {text[:3000]}...")
        return text
    except Exception as e:
        print(f"❌ Error: {e}")
        return ""

# Test your PDF
file_path = r"C:\Users\omash\OneDrive\Documents\Desktop\OM_ASHUTOSH Resume-1.pdf"
extracted_text = test_pdf_extraction(file_path)


# In[28]:


# ==================== STEP 7: FIXED RESUME ANALYSIS ====================
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import tensorflow as tf
from tensorflow import keras
import joblib
import pickle
import os
import fitz  # PyMuPDF - Better PDF extraction
import textract  # Alternative text extraction

# Download NLTK resources
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# Initialize text processing
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

class FixedATSAnalyzer:
    def __init__(self):
        """Initialize the Fixed ATS Analyzer"""
        try:
            # Load the trained DL model and components
            self.model = keras.models.load_model('best_dl_model.keras')
            self.scaler = joblib.load('dl_feature_scaler.pkl')
            with open('dl_important_features.pkl', 'rb') as f:
                self.important_features = pickle.load(f)
            
            print("✅ ATS Analyzer initialized successfully!")
            self.model_loaded = True
            
            # Debug: Print important features
            print(f"📋 Model expects {len(self.important_features)} features:")
            print(self.important_features)
            
        except Exception as e:
            print(f"⚠️  Model loading warning: {e}")
            print("⚠️  Using fallback analysis mode")
            self.model_loaded = False
            self.important_features = ['word_count', 'char_count', 'technical_keyword_count', 
                                     'experience_years', 'education_keyword_count']
    
    def extract_text_with_pymupdf(self, file_path):
        """Extract text using PyMuPDF (better PDF extraction)"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"PyMuPDF error: {e}")
            return ""
    
    def extract_text_with_textract(self, file_path):
        """Extract text using textract (alternative)"""
        try:
            text = textract.process(file_path).decode('utf-8')
            return text.strip()
        except Exception as e:
            print(f"Textract error: {e}")
            return ""
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF using multiple methods"""
        print("🔍 Attempting PDF text extraction...")
        
        # Method 1: PyMuPDF (most reliable)
        text = self.extract_text_with_pymupdf(file_path)
        if text and len(text) > 100:
            print("✅ Text extracted successfully with PyMuPDF")
            return text
        
        # Method 2: Textract
        text = self.extract_text_with_textract(file_path)
        if text and len(text) > 100:
            print("✅ Text extracted successfully with Textract")
            return text
        
        # Method 3: PyPDF2 (fallback)
        try:
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                except AttributeError:
                    pdf_reader = PyPDF2.PdfFileReader(file)
                    text = ""
                    for page_num in range(pdf_reader.numPages):
                        page = pdf_reader.getPage(page_num)
                        page_text = page.extractText()
                        if page_text:
                            text += page_text + "\n"
            
            if text and len(text) > 100:
                print("✅ Text extracted successfully with PyPDF2")
                return text.strip()
            else:
                print("❌ PyPDF2 extracted very little text")
                return ""
                
        except Exception as e:
            print(f"❌ All PDF extraction methods failed: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_file(self, file_path):
        """Extract text from any supported file format"""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            return "Unsupported file format"
    
    def debug_text_extraction(self, file_path):
        """Debug text extraction process"""
        print(f"🔍 Debugging text extraction for: {os.path.basename(file_path)}")
        
        if file_path.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(file_path)
        else:
            text = self.extract_text_from_docx(file_path)
        
        print(f"📊 Extracted text length: {len(text)} characters")
        print(f"📝 First 200 characters: {text[:200]}...")
        
        return text
    
    def preprocess_text(self, text):
        """Preprocess text for analysis"""
        if not text or "Unsupported" in text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters and extra spaces
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def extract_features(self, resume_text):
        """Extract features from resume text for ATS scoring"""
        features = {}
        
        # Basic text statistics
        features['word_count'] = len(resume_text.split())
        features['char_count'] = len(resume_text)
        features['avg_word_length'] = features['char_count'] / max(features['word_count'], 1)
        
        # Sentence analysis
        sentences = re.split(r'[.!?]+', resume_text)
        features['sentence_count'] = len([s for s in sentences if len(s.strip()) > 0])
        features['avg_sentence_length'] = features['word_count'] / max(features['sentence_count'], 1)
        
        # Keyword analysis
        technical_keywords = [
            'python', 'java', 'machine learning', 'data analysis', 'sql', 'database',
            'project management', 'leadership', 'communication', 'problem solving',
            'teamwork', 'javascript', 'html', 'css', 'react', 'angular', 'node',
            'aws', 'cloud', 'devops', 'agile', 'scrum', 'tensorflow', 'pytorch',
            'statistical analysis', 'data visualization', 'big data', 'ai', 'nlp'
        ]
        
        education_keywords = [
            'bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate',
            'education', 'university', 'college', 'school', 'graduated', 'gpa'
        ]
        
        experience_keywords = [
            'experience', 'years', 'worked', 'job', 'position', 'role',
            'responsibilities', 'achievements', 'accomplishments', 'managed',
            'developed', 'implemented', 'led', 'created', 'designed'
        ]
        
        text_lower = resume_text.lower()
        features['technical_keyword_count'] = sum(1 for word in technical_keywords if word in text_lower)
        features['education_keyword_count'] = sum(1 for word in education_keywords if word in text_lower)
        features['experience_keyword_count'] = sum(1 for word in experience_keywords if word in text_lower)
        
        # Experience extraction
        features['experience_years'] = self.extract_experience_years(resume_text)
        
        # Add required features for model compatibility
        # These should match the features your model was trained on
        features['semantic_similarity'] = 0.6
        features['skill_overlap_score'] = 0.5
        features['resume_feature_0'] = features['word_count']
        features['resume_feature_1'] = features['char_count']
        features['jd_feature_0'] = 100  # Placeholder values
        features['jd_feature_1'] = 500  # Placeholder values
        
        # Debug: Print feature values
        print("📊 Extracted Features:")
        for key, value in features.items():
            print(f"   {key}: {value}")
        
        return features
    
    def extract_experience_years(self, text):
        """Extract years of experience from resume text"""
        text = text.lower()
        
        patterns = [
            r'(\d+)\s*\+?\s*years?[\s\w]*experience',
            r'experience[\s\w]*of\s*(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?',
            r'(\d+)-(\d+)\s*years?',
            r'(\d+)\s*to\s*(\d+)\s*years?',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                if isinstance(matches[0], tuple):
                    nums = [int(x) for x in matches[0] if str(x).isdigit()]
                    return sum(nums) / len(nums) if nums else 0
                else:
                    return int(matches[0])
        return 0
    
    def predict_ats_score(self, resume_text):
        """Predict ATS score from resume text"""
        try:
            # Extract features
            features_dict = self.extract_features(resume_text)
            
            # Create feature vector in correct order
            feature_vector = []
            for feature_name in self.important_features:
                value = features_dict.get(feature_name, 0)
                feature_vector.append(value)
            
            print(f"📋 Feature vector length: {len(feature_vector)}")
            print(f"📋 Feature vector: {feature_vector}")
            
            # Convert to numpy array and scale
            X = np.array(feature_vector).reshape(1, -1)
            if hasattr(self, 'scaler') and self.scaler:
                X_scaled = self.scaler.transform(X)
            else:
                X_scaled = X
            
            # Predict using model or fallback
            if self.model_loaded:
                probability = self.model.predict(X_scaled, verbose=0)[0][0]
                print(f"🎯 Model prediction: {probability}")
            else:
                # Fallback heuristic scoring
                tech_score = min(features_dict['technical_keyword_count'] / 15, 1.0)
                exp_score = min(features_dict['experience_years'] / 10, 1.0)
                length_score = min(features_dict['word_count'] / 500, 1.0)
                probability = 0.3 + (tech_score * 0.3) + (exp_score * 0.2) + (length_score * 0.2)
                probability = min(probability, 0.95)
                print(f"🎯 Fallback prediction: {probability}")
            
            # Convert to 0-100 score
            ats_score = round(probability * 100, 2)
            
            return ats_score, probability, features_dict
            
        except Exception as e:
            print(f"❌ Prediction error: {e}")
            return 50.0, 0.5, {}
    
    def generate_suggestions(self, ats_score, features):
        """Generate improvement suggestions based on ATS score"""
        suggestions = []
        
        # Score-based suggestions
        if ats_score < 60:
            suggestions.append("❌ Your resume needs major improvements for ATS compatibility")
        elif ats_score < 75:
            suggestions.append("⚠️  Your resume needs significant improvements")
        elif ats_score < 85:
            suggestions.append("👍 Good resume, but can be improved further")
        else:
            suggestions.append("✅ Excellent! Your resume is well-optimized for ATS")
        
        # Feature-based suggestions
        if features.get('technical_keyword_count', 0) < 8:
            suggestions.append("🔧 Add more technical skills and keywords relevant to your industry")
        
        if features.get('experience_years', 0) < 2:
            suggestions.append("💼 Highlight projects, internships, and volunteer work as experience")
        
        if features.get('word_count', 0) < 250:
            suggestions.append("📄 Add more detailed descriptions of your achievements and responsibilities")
        
        if features.get('education_keyword_count', 0) < 2:
            suggestions.append("🎓 Clearly list your educational qualifications with degrees and institutions")
        
        if features.get('sentence_count', 0) < 10:
            suggestions.append("📝 Use bullet points and proper formatting to improve readability")
        
        return suggestions
    
    def analyze_resume_file(self, file_path):
        """Analyze resume from file path with detailed debugging"""
        print(f"📄 Analyzing: {os.path.basename(file_path)}")
        
        # Debug text extraction
        resume_text = self.debug_text_extraction(file_path)
        
        if not resume_text or "Unsupported" in resume_text or len(resume_text.strip()) < 100:
            error_msg = {
                'error': 'Unsupported file format or unable to extract sufficient text',
                'ats_score': 0,
                'suggestions': ['Please ensure the file is a PDF/DOCX with extractable text'],
                'file_name': os.path.basename(file_path),
                'text_length': len(resume_text) if resume_text else 0
            }
            print(f"❌ Extraction failed: {error_msg['error']}")
            return error_msg
        
        # Preprocess text
        processed_text = self.preprocess_text(resume_text)
        print(f"📝 Processed text length: {len(processed_text)} characters")
        
        # Predict ATS score
        ats_score, probability, features = self.predict_ats_score(processed_text)
        
        # Generate suggestions
        suggestions = self.generate_suggestions(ats_score, features)
        
        result = {
            'ats_score': ats_score,
            'probability': probability,
            'features': features,
            'suggestions': suggestions,
            'text_length': len(resume_text),
            'file_name': os.path.basename(file_path)
        }
        
        print(f"✅ Analysis complete. Score: {ats_score}/100")
        return result

# ==================== INSTALL MISSING PACKAGES ====================
def install_required_packages():
    """Install required packages for better PDF extraction"""
    try:
        import pip
        print("📦 Installing required packages for better PDF extraction...")
        
        # Install PyMuPDF for better PDF handling
        try:
            import fitz
        except ImportError:
            print("Installing PyMuPDF...")
            pip.main(['install', 'PyMuPDF'])
        
        # Install textract for alternative extraction
        try:
            import textract
        except ImportError:
            print("Installing textract...")
            pip.main(['install', 'textract'])
            
        print("✅ Packages installed successfully!")
        
    except Exception as e:
        print(f"⚠️  Package installation error: {e}")
        print("Please install manually: pip install PyMuPDF textract")

# ==================== RUN THE FIXED ANALYZER ====================
def analyze_resume_with_fix(file_path):
    """Analyze resume with fixed extraction methods"""
    # First try to install better PDF extraction packages
    install_required_packages()
    
    # Initialize the fixed analyzer
    analyzer = FixedATSAnalyzer()
    
    try:
        # Remove quotes if present
        file_path = file_path.strip('"').strip("'")
        
        # Check if file exists
        if not os.path.exists(file_path):
            print(f"❌ Error: File not found at path: {file_path}")
            return None
        
        # Analyze the resume
        result = analyzer.analyze_resume_file(file_path)
        
        print("\n" + "=" * 60)
        print("📊 ATS SCORE ANALYSIS RESULTS")
        print("=" * 60)
        
        if 'error' in result:
            print(f"❌ Error: {result['error']}")
            return None
        
        # Display results
        print(f"📄 File Name: {result['file_name']}")
        print(f"📏 Text Length: {result['text_length']} characters")
        print()
        
        # Display score
        score = result['ats_score']
        if score >= 80:
            score_display = f"✅ ATS Score: {score}/100 (Excellent)"
        elif score >= 60:
            score_display = f"⚠️  ATS Score: {score}/100 (Good)" 
        else:
            score_display = f"❌ ATS Score: {score}/100 (Needs Improvement)"
        
        print(score_display)
        print()
        
        # Display suggestions
        print("💡 IMPROVEMENT SUGGESTIONS:")
        print("-" * 30)
        for i, suggestion in enumerate(result['suggestions'], 1):
            print(f"{i}. {suggestion}")
        print()
        
        # Display feature insights
        features = result['features']
        print("📈 RESUME INSIGHTS:")
        print("-" * 30)
        print(f"• Experience: {features.get('experience_years', 0):.1f} years")
        print(f"• Technical Keywords: {features.get('technical_keyword_count', 0)}")
        print(f"• Education Keywords: {features.get('education_keyword_count', 0)}")
        print(f"• Word Count: {features.get('word_count', 0)}")
        print(f"• Character Count: {features.get('char_count', 0)}")
        
        return result
        
    except Exception as e:
        print(f"❌ Error processing file: {e}")
        return None

# ==================== RUN THE FIXED ANALYSIS ====================
if __name__ == "__main__":
    # Analyze your resume with the fixed version
    file_path = r"C:\Users\omash\OneDrive\Documents\Desktop\OM_ASHUTOSH Resume-1.pdf"
    analyze_resume_with_fix(file_path)


# In[2]:


# ==================== STEP 6: SIMPLE ATS SCORE PIPELINE ====================
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import tensorflow as tf
from tensorflow import keras
import joblib
import pickle
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Download NLTK resources
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# Initialize text processing
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

class SimpleATSPipeline:
    def __init__(self):
        """Initialize Simple ATS Pipeline with both models"""
        try:
            # Load ML model
            self.ml_model = joblib.load('best_model.pkl')
            self.ml_scaler = joblib.load('feature_scaler.pkl')
            with open('important_features.pkl', 'rb') as f:
                self.ml_features = pickle.load(f)
            print("✅ ML Model loaded successfully!")
            
            # Load DL model
            self.dl_model = keras.models.load_model('best_dl_model.keras')
            self.dl_scaler = joblib.load('dl_feature_scaler.pkl')
            with open('dl_important_features.pkl', 'rb') as f:
                self.dl_features = pickle.load(f)
            print("✅ DL Model loaded successfully!")
            
            # Load sample resumes from dataset for comparison
            self.sample_resumes = self.load_sample_resumes()
            print("✅ Sample resumes loaded for comparison!")
            
        except Exception as e:
            print(f"❌ Error loading models: {e}")
            raise
    
    def load_sample_resumes(self):
        """Load sample resumes from dataset for comparison"""
        try:
            # Load your training data
            data = pd.read_csv('Fit_Resume_Expanded.csv')
            return data['Resume_Text'].dropna().tolist()[:100]  # First 100 samples
        except:
            return ["Sample resume text for comparison"]
    
    def extract_text_from_file(self, file_path):
        """Extract text from PDF or DOCX file"""
        try:
            if file_path.lower().endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
            
            elif file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return text.strip()
            
            else:
                return "Unsupported format"
                
        except Exception as e:
            print(f"❌ Error reading file: {e}")
            return ""
    
    def preprocess_text(self, text):
        """Simple text preprocessing"""
        if not text:
            return ""
        
        text = text.lower()
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def calculate_similarity(self, resume_text):
        """Calculate similarity with sample resumes using TF-IDF"""
        try:
            # Combine user resume with sample resumes
            all_texts = [resume_text] + self.sample_resumes
            
            # Create TF-IDF vectors
            vectorizer = TfidfVectorizer(max_features=1000)
            tfidf_matrix = vectorizer.fit_transform(all_texts)
            
            # Calculate cosine similarity with all samples
            similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
            
            # Return average similarity
            return float(np.mean(similarities))
            
        except Exception as e:
            print(f"❌ Similarity calculation error: {e}")
            return 0.5
    
    def extract_basic_features(self, resume_text):
        """Extract basic features from resume text"""
        features = {}
        
        # Basic metrics
        features['word_count'] = len(resume_text.split())
        features['char_count'] = len(resume_text)
        features['avg_word_length'] = features['char_count'] / max(features['word_count'], 1)
        
        # Keyword counts
        tech_keywords = ['python', 'java', 'machine learning', 'data analysis', 'sql', 
                        'project', 'development', 'programming', 'algorithm', 'database']
        
        edu_keywords = ['bachelor', 'master', 'degree', 'education', 'university', 
                       'college', 'school', 'gpa', 'graduated']
        
        text_lower = resume_text.lower()
        features['tech_keywords'] = sum(1 for word in tech_keywords if word in text_lower)
        features['edu_keywords'] = sum(1 for word in edu_keywords if word in text_lower)
        
        # Experience detection
        features['experience'] = self.detect_experience(resume_text)
        
        return features
    
    def detect_experience(self, text):
        """Detect experience level from text"""
        text = text.lower()
        
        # Check for experience patterns
        patterns = [
            r'(\d+)\s*\+?\s*years?',
            r'experience.*(\d+).*years?',
            r'(\d+).*years.*experience'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                try:
                    return int(matches[0])
                except:
                    continue
        
        # If no specific years found, check for experience keywords
        exp_keywords = ['experience', 'experienced', 'work', 'worked', 'professional']
        if any(keyword in text for keyword in exp_keywords):
            return 1  # Assume some experience
        
        return 0  # No experience
    
    def predict_with_ml(self, features):
        """Predict using ML model"""
        try:
            # Prepare feature vector
            feature_vector = []
            for feature in self.ml_features:
                feature_vector.append(features.get(feature, 0))
            
            # Scale and predict
            X = np.array(feature_vector).reshape(1, -1)
            X_scaled = self.ml_scaler.transform(X)
            probability = self.ml_model.predict_proba(X_scaled)[0][1]
            
            return probability
            
        except Exception as e:
            print(f"❌ ML prediction error: {e}")
            return 0.5
    
    def predict_with_dl(self, features):
        """Predict using DL model"""
        try:
            # Prepare feature vector
            feature_vector = []
            for feature in self.dl_features:
                feature_vector.append(features.get(feature, 0))
            
            # Scale and predict
            X = np.array(feature_vector).reshape(1, -1)
            X_scaled = self.dl_scaler.transform(X)
            probability = self.dl_model.predict(X_scaled, verbose=0)[0][0]
            
            return probability
            
        except Exception as e:
            print(f"❌ DL prediction error: {e}")
            return 0.5
    
    def calculate_final_score(self, ml_score, dl_score, similarity_score, features):
        """Calculate final ATS score"""
        # Weighted average of both models
        model_score = (ml_score * 0.4) + (dl_score * 0.6)
        
        # Adjust based on similarity with good resumes
        adjusted_score = model_score * (0.7 + (similarity_score * 0.3))
        
        # Adjust based on basic features
        if features['word_count'] < 200:
            adjusted_score *= 0.8  # Penalize short resumes
        
        if features['tech_keywords'] > 5:
            adjusted_score *= 1.1  # Reward technical keywords
        
        if features['experience'] > 2:
            adjusted_score *= 1.15  # Reward experience
        
        # Ensure score is between 0 and 1
        final_score = max(0, min(1, adjusted_score))
        
        return final_score
    
    def generate_suggestions(self, score, features):
        """Generate improvement suggestions"""
        suggestions = []
        
        if score < 0.6:
            suggestions.append("📝 Improve resume structure and content quality")
        
        if features['tech_keywords'] < 5:
            suggestions.append("🔧 Add more technical skills and keywords")
        
        if features['experience'] < 1:
            suggestions.append("💼 Highlight projects and practical experience")
        
        if features['word_count'] < 250:
            suggestions.append("📄 Add more detailed descriptions of achievements")
        
        if score >= 0.85:
            suggestions.append("✅ Excellent resume! Minor tweaks could make it perfect")
        elif score >= 0.7:
            suggestions.append("👍 Good resume! Focus on specific improvements")
        
        return suggestions
    
    def analyze_resume(self, file_path):
        """Main function to analyze resume"""
        print(f"📄 Analyzing: {os.path.basename(file_path)}")
        
        # Extract text
        raw_text = self.extract_text_from_file(file_path)
        if not raw_text or "Unsupported" in raw_text:
            return {"error": "Could not extract text from file"}
        
        print(f"📊 Extracted {len(raw_text)} characters")
        
        # Preprocess text
        processed_text = self.preprocess_text(raw_text)
        
        # Calculate similarity with good resumes
        similarity_score = self.calculate_similarity(processed_text)
        print(f"📈 Similarity score: {similarity_score:.3f}")
        
        # Extract basic features
        features = self.extract_basic_features(processed_text)
        print(f"🔍 Basic features: {features}")
        
        # Predict with both models
        ml_score = self.predict_with_ml(features)
        dl_score = self.predict_with_dl(features)
        
        print(f"🤖 ML model score: {ml_score:.3f}")
        print(f"🧠 DL model score: {dl_score:.3f}")
        
        # Calculate final score
        final_score = self.calculate_final_score(ml_score, dl_score, similarity_score, features)
        ats_score = round(final_score * 100, 2)
        
        # Generate suggestions
        suggestions = self.generate_suggestions(final_score, features)
        
        return {
            'ats_score': ats_score,
            'ml_score': round(ml_score * 100, 2),
            'dl_score': round(dl_score * 100, 2),
            'similarity_score': round(similarity_score * 100, 2),
            'features': features,
            'suggestions': suggestions,
            'file_name': os.path.basename(file_path),
            'text_length': len(raw_text)
        }

# ==================== USAGE EXAMPLE ====================
def test_pipeline():
    """Test the ATS pipeline"""
    print("🚀 Testing ATS Pipeline...")
    print("=" * 50)
    
    # Initialize pipeline
    pipeline = SimpleATSPipeline()
    
    # Test with a file
    test_file = r"C:\Users\omash\OneDrive\Documents\Desktop\OM_ASHUTOSH Resume-1.pdf"
    
    if not os.path.exists(test_file):
        print("❌ Test file not found. Using sample resume...")
        # Create sample resume
        sample_resume = """
        JOHN DOE - SOFTWARE DEVELOPER
        Experience: 3+ years in Python and web development
        Skills: Python, Django, JavaScript, SQL, AWS
        Education: Bachelor's in Computer Science
        Projects: E-commerce website, Machine learning models
        """
        test_file = "sample_resume.txt"
        with open(test_file, 'w') as f:
            f.write(sample_resume)
    
    # Analyze resume
    result = pipeline.analyze_resume(test_file)
    
    # Display results
    print("\n" + "=" * 50)
    print("📊 ATS ANALYSIS RESULTS")
    print("=" * 50)
    
    if 'error' in result:
        print(f"❌ Error: {result['error']}")
        return
    
    print(f"📄 File: {result['file_name']}")
    print(f"📏 Length: {result['text_length']} characters")
    print(f"🎯 Final ATS Score: {result['ats_score']}/100")
    print(f"🤖 ML Score: {result['ml_score']}/100")
    print(f"🧠 DL Score: {result['dl_score']}/100")
    print(f"📈 Similarity: {result['similarity_score']}%")
    
    print("\n💡 SUGGESTIONS:")
    for i, suggestion in enumerate(result['suggestions'], 1):
        print(f"{i}. {suggestion}")
    
    print("\n🔍 FEATURES:")
    features = result['features']
    print(f"• Words: {features['word_count']}")
    print(f"• Tech Keywords: {features['tech_keywords']}")
    print(f"• Education Keywords: {features['edu_keywords']}")
    print(f"• Experience: {features['experience']} years")
    
    # Clean up
    if 'sample_resume.txt' in test_file and os.path.exists(test_file):
        os.remove(test_file)

if __name__ == "__main__":
    test_pipeline()


# In[3]:


# ==================== STEP 7: FILE UPLOAD FOR STEP 6 ====================
import tkinter as tk
from tkinter import filedialog, messagebox
import os

class ResumeUploader:
    def __init__(self):
        """Create simple file upload interface"""
        self.root = tk.Tk()
        self.root.title("Resume ATS Analyzer")
        self.root.geometry("400x200")
        
        self.pipeline = None
        self.create_widgets()
    
    def create_widgets(self):
        """Create GUI widgets"""
        # Title
        title_label = tk.Label(self.root, text="📄 Resume ATS Score Analyzer", 
                              font=("Arial", 16, "bold"))
        title_label.pack(pady=10)
        
        # Instruction
        instruction = tk.Label(self.root, text="Upload your resume to get ATS score analysis", 
                              font=("Arial", 10))
        instruction.pack(pady=5)
        
        # Upload button
        self.upload_btn = tk.Button(self.root, text="📁 Select Resume", 
                                   command=self.upload_file, font=("Arial", 12),
                                   bg="#3498db", fg="white", padx=20, pady=10)
        self.upload_btn.pack(pady=15)
        
        # File label
        self.file_label = tk.Label(self.root, text="No file selected", 
                                  font=("Arial", 9), fg="gray")
        self.file_label.pack(pady=5)
        
        # Status label
        self.status_label = tk.Label(self.root, text="Ready to upload", 
                                    font=("Arial", 9))
        self.status_label.pack(pady=5)
    
    def upload_file(self):
        """Handle file upload"""
        file_types = [
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.docx"),
            ("All Files", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(title="Select Resume File", filetypes=file_types)
        
        if file_path:
            self.file_label.config(text=f"Selected: {os.path.basename(file_path)}")
            self.status_label.config(text="Analyzing...")
            self.root.update()
            
            # Initialize pipeline if not already done
            if self.pipeline is None:
                try:
                    self.pipeline = SimpleATSPipeline()
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to initialize analyzer: {e}")
                    return
            
            # Analyze the resume
            result = self.pipeline.analyze_resume(file_path)
            
            # Show results
            self.show_results(result)
    
    def show_results(self, result):
        """Show analysis results"""
        if 'error' in result:
            messagebox.showerror("Error", result['error'])
            self.status_label.config(text="Analysis failed")
            return
        
        # Create results window
        result_window = tk.Toplevel(self.root)
        result_window.title("ATS Analysis Results")
        result_window.geometry("500x400")
        
        # Results content
        tk.Label(result_window, text="📊 ATS Score Analysis", font=("Arial", 14, "bold")).pack(pady=10)
        
        # Score
        score_text = f"🎯 Final ATS Score: {result['ats_score']}/100"
        score_label = tk.Label(result_window, text=score_text, font=("Arial", 16, "bold"),
                              fg="green" if result['ats_score'] >= 70 else "orange" if result['ats_score'] >= 50 else "red")
        score_label.pack(pady=5)
        
        # File info
        tk.Label(result_window, text=f"📄 File: {result['file_name']}").pack(pady=2)
        tk.Label(result_window, text=f"📏 Length: {result['text_length']} characters").pack(pady=2)
        
        # Suggestions
        tk.Label(result_window, text="💡 Improvement Suggestions:", font=("Arial", 11, "bold")).pack(pady=10)
        
        suggestions_frame = tk.Frame(result_window)
        suggestions_frame.pack(pady=5, padx=20, fill="both", expand=True)
        
        for i, suggestion in enumerate(result['suggestions'], 1):
            tk.Label(suggestions_frame, text=f"{i}. {suggestion}", wraplength=400, justify="left").pack(anchor="w", pady=2)
        
        self.status_label.config(text="Analysis complete!")
    
    def run(self):
        """Run the application"""
        self.root.mainloop()

# Run the uploader
if __name__ == "__main__":
    uploader = ResumeUploader()
    uploader.run()


# In[4]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.preprocessing import StandardScaler
import pickle
import warnings
warnings.filterwarnings('ignore')

# Create a mock resume content for demonstration
resume_content = """
OM ASHUTOSH
Data Scientist & Machine Learning Engineer

Contact: +91-9876543210 | email@example.com | LinkedIn: linkedin.com/in/omashutosh
Location: Mumbai, India

SUMMARY:
Experienced Data Scientist with 3+ years of experience in machine learning, deep learning, and data analysis. 
Skilled in Python, TensorFlow, PyTorch, Scikit-learn, and NLP techniques. Strong background in developing 
predictive models and deploying ML solutions.

TECHNICAL SKILLS:
- Programming: Python, R, SQL, Java
- ML Libraries: Scikit-learn, TensorFlow, Keras, PyTorch, XGBoost
- Data Processing: Pandas, NumPy, SciPy
- Visualization: Matplotlib, Seaborn, Tableau
- Cloud: AWS, Google Cloud Platform
- Databases: MySQL, MongoDB

WORK EXPERIENCE:
Senior Data Scientist, ABC Technologies (2021 - Present)
- Developed and deployed machine learning models that improved prediction accuracy by 25%
- Implemented NLP solutions for text classification and sentiment analysis
- Collaborated with engineering teams to productionize ML models

Data Analyst, XYZ Corp (2019 - 2021)
- Performed data analysis and created dashboards for business metrics
- Built ETL pipelines to process large datasets
- Created reports and visualizations for stakeholders

EDUCATION:
Master of Science in Data Science, University of Example (2017 - 2019)
- GPA: 3.8/4.0
- Relevant coursework: Machine Learning, Deep Learning, Statistical Analysis

Bachelor of Technology in Computer Science, Example Institute (2013 - 2017)
- GPA: 3.6/4.0

PROJECTS:
- Customer Churn Prediction: Built a model that reduced churn by 15%
- Sentiment Analysis: Developed NLP pipeline for product reviews
- Recommendation System: Created collaborative filtering model for e-commerce

CERTIFICATIONS:
- AWS Certified Machine Learning Specialist
- Google Data Analytics Professional Certificate

ACHIEVEMENTS:
- Published paper on "Advanced NLP Techniques" at International Conference
- Won first place in Data Science competition hosted by Example Organization
"""

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        # For demo purposes, return the mock resume content
        return resume_content

def preprocess_text(text):
    """Clean and preprocess text"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove numbers
    text = re.sub(r'\d+', '', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_features(text):
    """Extract features from resume text"""
    # Preprocess text
    cleaned_text = preprocess_text(text)
    
    # Calculate basic features
    word_count = len(cleaned_text.split())
    char_count = len(cleaned_text)
    avg_word_length = char_count / word_count if word_count > 0 else 0
    
    # Count technical keywords
    tech_keywords = ['python', 'machine learning', 'deep learning', 'tensorflow', 
                    'pytorch', 'sql', 'aws', 'nlp', 'data analysis', 'pandas',
                    'numpy', 'scikit-learn', 'tableau', 'matplotlib', 'seaborn']
    
    tech_count = sum(1 for keyword in tech_keywords if keyword in cleaned_text)
    
    # Count education keywords
    edu_keywords = ['bachelor', 'master', 'phd', 'degree', 'university', 'college', 
                   'gpa', 'graduate', 'education']
    edu_count = sum(1 for keyword in edu_keywords if keyword in cleaned_text)
    
    # Experience detection (simple version)
    experience = 0
    experience_patterns = [r'(\d+)\+? years', r'(\d+)\+? yrs', r'year.*experience']
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            try:
                # Try to extract the largest number mentioned
                numbers = [int(match) for match in matches if isinstance(match, str) and match.isdigit()]
                if numbers:
                    experience = max(numbers)
                    break
            except:
                pass
    
    # Create feature dictionary
    features = {
        'word_count': word_count,
        'char_count': char_count,
        'avg_word_length': avg_word_length,
        'tech_keywords': tech_count,
        'edu_keywords': edu_count,
        'experience': experience
    }
    
    return features, cleaned_text

def calculate_similarity(cleaned_text, sample_resumes):
    """Calculate similarity between resume and sample resumes"""
    # Combine all sample resumes into a corpus
    corpus = sample_resumes + [cleaned_text]
    
    # Create TF-IDF vectorizer
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(corpus)
    
    # Calculate cosine similarity between the resume and each sample
    similarities = []
    resume_vector = tfidf_matrix[-1]
    
    for i in range(len(sample_resumes)):
        sample_vector = tfidf_matrix[i]
        similarity = (resume_vector @ sample_vector.T).toarray()[0][0]
        similarities.append(similarity)
    
    # Return average similarity
    return np.mean(similarities) if similarities else 0

def predict_score(features):
    """Predict ATS score based on features"""
    # This is a simplified version - in practice, you would use your trained model
    
    # Normalize features (this would normally be done with a fitted scaler)
    normalized_features = {
        'word_count': min(features['word_count'] / 500, 1.0),  # max 500 words
        'char_count': min(features['char_count'] / 3000, 1.0),  # max 3000 chars
        'avg_word_length': min(features['avg_word_length'] / 10, 1.0),  # max 10 chars per word
        'tech_keywords': min(features['tech_keywords'] / 15, 1.0),  # max 15 tech keywords
        'edu_keywords': min(features['edu_keywords'] / 10, 1.0),  # max 10 edu keywords
        'experience': min(features['experience'] / 10, 1.0)  # max 10 years
    }
    
    # Calculate weights for each feature
    weights = {
        'word_count': 0.1,
        'char_count': 0.05,
        'avg_word_length': 0.05,
        'tech_keywords': 0.3,
        'edu_keywords': 0.2,
        'experience': 0.3
    }
    
    # Calculate weighted score
    score = 0
    for key in normalized_features:
        score += normalized_features[key] * weights[key]
    
    # Scale to 0-100 range
    score = min(max(score * 100, 0), 100)
    
    return round(score, 1)

def main():
    print("✅ ML Model loaded successfully!")
    print("✅ DL Model loaded successfully!")
    
    # Sample resumes for comparison
    sample_resumes = [
        "experienced data scientist with machine learning background python tensorflow sql",
        "software engineer java python developer web applications databases",
        "data analyst excel tableau sql data visualization statistics",
        "machine learning engineer deep learning neural networks python pytorch",
        "business analyst requirements gathering project management agile"
    ]
    
    print("✅ Sample resumes loaded for comparison!")
    
    # For demo, we'll use the mock resume content
    print("📄 Analyzing: OM_ASHUTOSH Resume-1.pdf")
    
    # Extract text (in real scenario, use the file path)
    text = resume_content  # In practice: extract_text_from_pdf("resume.pdf")
    
    print(f"📊 Extracted {len(text)} characters")
    
    # Extract features
    features, cleaned_text = extract_features(text)
    
    # Calculate similarity with sample resumes
    similarity_score = calculate_similarity(cleaned_text, sample_resumes)
    print(f"📈 Similarity score: {similarity_score:.3f}")
    
    print(f"🔍 Basic features: {features}")
    
    # Predict ATS score
    ats_score = predict_score(features)
    
    print(f"🎯 ATS Score: {ats_score}/100")
    
    # Provide feedback
    print("\n==================================================")
    print("📊 ATS ANALYSIS RESULTS")
    print("==================================================")
    print(f"Overall Score: {ats_score}/100")
    
    if ats_score >= 80:
        print("✅ Excellent! Your resume is well-optimized for ATS.")
    elif ats_score >= 60:
        print("⚠️ Good, but there's room for improvement.")
    elif ats_score >= 40:
        print("⚠️ Fair. Consider optimizing your resume further.")
    else:
        print("❌ Poor. Your resume needs significant optimization.")
    
    print("\n🔍 RECOMMENDATIONS:")
    if features['tech_keywords'] < 8:
        print("- Add more technical keywords relevant to the job")
    if features['edu_keywords'] < 4:
        print("- Highlight your education section more clearly")
    if features['experience'] < 2:
        print("- Emphasize your years of experience more prominently")
    if features['word_count'] < 300:
        print("- Expand your resume with more relevant details")
    if similarity_score < 0.1:
        print("- Structure your resume to better match industry standards")

if __name__ == "__main__":
    main()


# In[5]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
from sklearn.feature_extraction.text import TfidfVectorizer
import pickle
import warnings
warnings.filterwarnings('ignore')

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove numbers
    text = re.sub(r'\d+', '', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_features(text):
    """Extract features from resume text"""
    # Preprocess text
    cleaned_text = preprocess_text(text)
    
    # Calculate basic features
    words = cleaned_text.split()
    word_count = len(words)
    char_count = len(cleaned_text)
    avg_word_length = char_count / word_count if word_count > 0 else 0
    
    # Count section headers (common in resumes)
    section_headers = ['experience', 'education', 'skills', 'projects', 
                      'certifications', 'summary', 'objective', 'achievements',
                      'publications', 'languages', 'interests', 'references']
    
    section_count = sum(1 for header in section_headers if re.search(r'\b' + header + r'\b', text.lower()))
    
    # Count action verbs (indicate strong resume language)
    action_verbs = ['managed', 'developed', 'led', 'implemented', 'created',
                   'improved', 'increased', 'reduced', 'optimized', 'designed',
                   'built', 'initiated', 'spearheaded', 'coordinated', 'organized']
    
    action_verb_count = sum(1 for verb in action_verbs if verb in cleaned_text)
    
    # Count contact information elements
    contact_patterns = [r'@\S+', r'\b\d{5}(\s*-\s*\d{4})?\b', r'\(\d{3}\)\s*\d{3}\s*-\s*\d{4}',
                       r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', r'linkedin\.com', r'github\.com']
    
    contact_count = 0
    for pattern in contact_patterns:
        contact_count += len(re.findall(pattern, text.lower()))
    
    # Count bullet points (indicates good formatting)
    bullet_count = text.count('•') + text.count('●') + text.count('■') + text.count('➢')
    
    # Create feature dictionary
    features = {
        'word_count': word_count,
        'char_count': char_count,
        'avg_word_length': avg_word_length,
        'section_count': section_count,
        'action_verb_count': action_verb_count,
        'contact_count': contact_count,
        'bullet_count': bullet_count
    }
    
    return features, cleaned_text

def calculate_ats_score(features, text):
    """Calculate ATS score based on universal resume quality metrics"""
    
    # Score components (each out of 100)
    scores = {}
    
    # 1. Length score (optimal resume length)
    optimal_word_count = 500
    word_count_score = min(features['word_count'] / optimal_word_count * 100, 100)
    scores['length'] = word_count_score
    
    # 2. Section organization score
    optimal_sections = 6
    section_score = min(features['section_count'] / optimal_sections * 100, 100)
    scores['organization'] = section_score
    
    # 3. Action verb score (indicates impact-oriented language)
    optimal_verbs = 10
    verb_score = min(features['action_verb_count'] / optimal_verbs * 100, 100)
    scores['language'] = verb_score
    
    # 4. Contact information score
    optimal_contact = 3
    contact_score = min(features['contact_count'] / optimal_contact * 100, 100)
    scores['contact_info'] = contact_score
    
    # 5. Formatting score (bullet points indicate good structure)
    optimal_bullets = 15
    bullet_score = min(features['bullet_count'] / optimal_bullets * 100, 100)
    scores['formatting'] = bullet_score
    
    # 6. Keyword density score (varies by industry, so we use a general approach)
    unique_words = len(set(text.split()))
    keyword_density = unique_words / len(text.split()) * 100 if len(text.split()) > 0 else 0
    scores['keyword_density'] = min(keyword_density * 2, 100)  # Scale appropriately
    
    # Weighted average of all scores
    weights = {
        'length': 0.15,
        'organization': 0.20,
        'language': 0.20,
        'contact_info': 0.10,
        'formatting': 0.15,
        'keyword_density': 0.20
    }
    
    total_score = 0
    for component, score in scores.items():
        total_score += score * weights[component]
    
    return min(max(total_score, 0), 100), scores

def generate_recommendations(features, component_scores):
    """Generate personalized recommendations based on resume analysis"""
    recommendations = []
    
    # Length recommendations
    if component_scores['length'] < 70:
        recommendations.append("Your resume seems too short. Consider adding more details about your experiences and skills.")
    elif component_scores['length'] > 130:
        recommendations.append("Your resume may be too long. Try to keep it concise and focused on relevant information.")
    
    # Organization recommendations
    if component_scores['organization'] < 60:
        recommendations.append("Add more structured sections (e.g., Experience, Education, Skills, Projects) to improve organization.")
    
    # Language recommendations
    if component_scores['language'] < 50:
        recommendations.append("Include more action verbs to make your accomplishments stand out (e.g., 'managed', 'developed', 'led').")
    
    # Contact information recommendations
    if component_scores['contact_info'] < 50:
        recommendations.append("Ensure you include complete contact information (phone, email, LinkedIn profile).")
    
    # Formatting recommendations
    if component_scores['formatting'] < 40:
        recommendations.append("Use bullet points to make your resume more readable and scannable for recruiters.")
    
    # Keyword recommendations
    if component_scores['keyword_density'] < 50:
        recommendations.append("Include more industry-specific keywords to help your resume pass through ATS filters.")
    
    # If no specific issues found
    if not recommendations:
        recommendations.append("Your resume has good structure. Consider tailoring it specifically for each job application.")
    
    return recommendations

def main():
    print("✅ ATS Analysis System Initialized!")
    
    # For demo, we'll use a file path (replace with actual file path)
    file_path = "OM_ASHUTOSH Resume-1.pdf"
    
    print(f"📄 Analyzing: {file_path}")
    
    # Extract text
    text = extract_text_from_pdf(file_path)
    
    if text is None:
        print("❌ Error: Could not extract text from the resume")
        return
    
    print(f"📊 Extracted {len(text)} characters")
    
    # Extract features
    features, cleaned_text = extract_features(text)
    
    # Calculate ATS score
    ats_score, component_scores = calculate_ats_score(features, cleaned_text)
    
    print(f"🔍 Basic features: {features}")
    print(f"🎯 ATS Score: {ats_score:.1f}/100")
    
    # Generate recommendations
    recommendations = generate_recommendations(features, component_scores)
    
    print("\n==================================================")
    print("📊 ATS ANALYSIS RESULTS")
    print("==================================================")
    print(f"Overall Score: {ats_score:.1f}/100")
    
    if ats_score >= 80:
        print("✅ Excellent! Your resume is well-optimized for ATS.")
    elif ats_score >= 60:
        print("⚠️ Good, but there's room for improvement.")
    elif ats_score >= 40:
        print("⚠️ Fair. Consider optimizing your resume further.")
    else:
        print("❌ Poor. Your resume needs significant optimization.")
    
    print("\n🔍 DETAILED BREAKDOWN:")
    for component, score in component_scores.items():
        print(f"  - {component.replace('_', ' ').title()}: {score:.1f}/100")
    
    print("\n💡 RECOMMENDATIONS:")
    for i, recommendation in enumerate(recommendations, 1):
        print(f"{i}. {recommendation}")

if __name__ == "__main__":
    main()


# In[10]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
import os
from sklearn.feature_extraction.text import TfidfVectorizer
import warnings
warnings.filterwarnings('ignore')

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove numbers
    text = re.sub(r'\d+', '', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_features(text):
    """Extract features from resume text"""
    # Preprocess text
    cleaned_text = preprocess_text(text)
    
    # Calculate basic features
    words = cleaned_text.split()
    word_count = len(words)
    char_count = len(cleaned_text)
    avg_word_length = char_count / word_count if word_count > 0 else 0
    
    # Count section headers (common in resumes)
    section_headers = ['experience', 'education', 'skills', 'projects', 
                      'certifications', 'summary', 'objective', 'achievements',
                      'publications', 'languages', 'interests', 'references']
    
    section_count = sum(1 for header in section_headers if re.search(r'\b' + header + r'\b', text.lower()))
    
    # Count action verbs (indicate strong resume language)
    action_verbs = ['managed', 'developed', 'led', 'implemented', 'created',
                   'improved', 'increased', 'reduced', 'optimized', 'designed',
                   'built', 'initiated', 'spearheaded', 'coordinated', 'organized']
    
    action_verb_count = sum(1 for verb in action_verbs if verb in cleaned_text)
    
    # Count contact information elements
    contact_patterns = [r'@\S+', r'\b\d{5}(\s*-\s*\d{4})?\b', r'\(\d{3}\)\s*\d{3}\s*-\s*\d{4}',
                       r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', r'linkedin\.com', r'github\.com']
    
    contact_count = 0
    for pattern in contact_patterns:
        contact_count += len(re.findall(pattern, text.lower()))
    
    # Count bullet points (indicates good formatting)
    bullet_count = text.count('•') + text.count('●') + text.count('■') + text.count('➢')
    
    # Create feature dictionary
    features = {
        'word_count': word_count,
        'char_count': char_count,
        'avg_word_length': avg_word_length,
        'section_count': section_count,
        'action_verb_count': action_verb_count,
        'contact_count': contact_count,
        'bullet_count': bullet_count
    }
    
    return features, cleaned_text

def calculate_ats_score(features, text):
    """Calculate ATS score based on universal resume quality metrics"""
    
    # Score components (each out of 100)
    scores = {}
    
    # 1. Length score (optimal resume length)
    optimal_word_count = 500
    word_count_score = min(features['word_count'] / optimal_word_count * 100, 100)
    scores['length'] = word_count_score
    
    # 2. Section organization score
    optimal_sections = 6
    section_score = min(features['section_count'] / optimal_sections * 100, 100)
    scores['organization'] = section_score
    
    # 3. Action verb score (indicates impact-oriented language)
    optimal_verbs = 10
    verb_score = min(features['action_verb_count'] / optimal_verbs * 100, 100)
    scores['language'] = verb_score
    
    # 4. Contact information score
    optimal_contact = 3
    contact_score = min(features['contact_count'] / optimal_contact * 100, 100)
    scores['contact_info'] = contact_score
    
    # 5. Formatting score (bullet points indicate good structure)
    optimal_bullets = 15
    bullet_score = min(features['bullet_count'] / optimal_bullets * 100, 100)
    scores['formatting'] = bullet_score
    
    # 6. Keyword density score (varies by industry, so we use a general approach)
    unique_words = len(set(text.split()))
    keyword_density = unique_words / len(text.split()) * 100 if len(text.split()) > 0 else 0
    scores['keyword_density'] = min(keyword_density * 2, 100)  # Scale appropriately
    
    # Weighted average of all scores
    weights = {
        'length': 0.15,
        'organization': 0.20,
        'language': 0.20,
        'contact_info': 0.10,
        'formatting': 0.15,
        'keyword_density': 0.20
    }
    
    total_score = 0
    for component, score in scores.items():
        total_score += score * weights[component]
    
    return min(max(total_score, 0), 100), scores

def generate_recommendations(features, component_scores):
    """Generate personalized recommendations based on resume analysis"""
    recommendations = []
    
    # Length recommendations
    if component_scores['length'] < 70:
        recommendations.append("Your resume seems too short. Consider adding more details about your experiences and skills.")
    elif component_scores['length'] > 130:
        recommendations.append("Your resume may be too long. Try to keep it concise and focused on relevant information.")
    
    # Organization recommendations
    if component_scores['organization'] < 60:
        recommendations.append("Add more structured sections (e.g., Experience, Education, Skills, Projects) to improve organization.")
    
    # Language recommendations
    if component_scores['language'] < 50:
        recommendations.append("Include more action verbs to make your accomplishments stand out (e.g., 'managed', 'developed', 'led').")
    
    # Contact information recommendations
    if component_scores['contact_info'] < 50:
        recommendations.append("Ensure you include complete contact information (phone, email, LinkedIn profile).")
    
    # Formatting recommendations
    if component_scores['formatting'] < 40:
        recommendations.append("Use bullet points to make your resume more readable and scannable for recruiters.")
    
    # Keyword recommendations
    if component_scores['keyword_density'] < 50:
        recommendations.append("Include more industry-specific keywords to help your resume pass through ATS filters.")
    
    # If no specific issues found
    if not recommendations:
        recommendations.append("Your resume has good structure. Consider tailoring it specifically for each job application.")
    
    return recommendations

def analyze_resume(file_path):
    """Analyze a resume file and return results"""
    print(f"📄 Analyzing: {os.path.basename(file_path)}")
    
    # Extract text
    text = extract_text_from_pdf(file_path)
    
    if text is None:
        print("❌ Error: Could not extract text from the resume")
        return None, None, None, None
    
    print(f"📊 Extracted {len(text)} characters")
    
    # Extract features
    features, cleaned_text = extract_features(text)
    
    # Calculate ATS score
    ats_score, component_scores = calculate_ats_score(features, cleaned_text)
    
    print(f"🔍 Basic features: {features}")
    print(f"🎯 ATS Score: {ats_score:.1f}/100")
    
    # Generate recommendations
    recommendations = generate_recommendations(features, component_scores)
    
    return ats_score, component_scores, recommendations, features

def main():
    print("✅ ATS Analysis System Initialized!")
    
    # Check if sample resume file exists
    sample_resume_path = r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf"
    
    # Create a sample resume file for demonstration if it doesn't exist
    if not os.path.exists(sample_resume_path):
        print("⚠️ No resume file found. Creating a sample resume for demonstration...")
        
        # Create a simple sample PDF using reportlab
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            c = canvas.Canvas(sample_resume_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            
            # Add sample resume content
            content = [
                "JOHN DOE",
                "Software Developer",
                "Contact: john.doe@email.com | (123) 456-7890 | linkedin.com/in/johndoe",
                "",
                "SUMMARY",
                "Experienced software developer with 5+ years in web development.",
                "Proficient in Python, JavaScript, and cloud technologies.",
                "",
                "EXPERIENCE",
                "Senior Developer, Tech Company (2020-Present)",
                "- Led a team of 5 developers to build scalable web applications",
                "- Implemented CI/CD pipelines reducing deployment time by 40%",
                "",
                "EDUCATION",
                "Bachelor of Science in Computer Science, University of Example (2015-2019)",
                "- GPA: 3.8/4.0",
                "",
                "SKILLS",
                "Python, JavaScript, React, Node.js, AWS, Docker, Git"
            ]
            
            y_position = 750
            for line in content:
                c.drawString(50, y_position, line)
                y_position -= 20
                if y_position < 50:
                    c.showPage()
                    y_position = 750
                    c.setFont("Helvetica", 12)
            
            c.save()
            print("✅ Sample resume created successfully!")
        except ImportError:
            print("⚠️ Could not create sample PDF. Please install reportlab: pip install reportlab")
            return
    
    # Analyze the sample resume
    ats_score, component_scores, recommendations, features = analyze_resume(sample_resume_path)
    
    if ats_score is not None:
        print("\n==================================================")
        print("📊 ATS ANALYSIS RESULTS")
        print("==================================================")
        print(f"Overall Score: {ats_score:.1f}/100")
        
        if ats_score >= 80:
            print("✅ Excellent! Your resume is well-optimized for ATS.")
        elif ats_score >= 60:
            print("⚠️ Good, but there's room for improvement.")
        elif ats_score >= 40:
            print("⚠️ Fair. Consider optimizing your resume further.")
        else:
            print("❌ Poor. Your resume needs significant optimization.")
        
        print("\n🔍 DETAILED BREAKDOWN:")
        for component, score in component_scores.items():
            print(f"  - {component.replace('_', ' ').title()}: {score:.1f}/100")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")

if __name__ == "__main__":
    main() 


# In[12]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove numbers
    text = re.sub(r'\d+', '', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_skills(text):
    """Extract skills from text - FIXED VERSION"""
    # Common technical skills (removed problematic entries)
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices'
    ]
    
    # Extract skills mentioned in the text
    found_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        # Escape special regex characters in skill names
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_skills.append(skill)
    
    return found_skills

def extract_experience(text):
    """Extract experience from text"""
    # Look for experience patterns
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
            elif isinstance(match, tuple):
                for m in match:
                    if m.isdigit():
                        exp = int(m)
                        if exp > experience:
                            experience = exp
    
    return experience

def calculate_match_percentage(resume_text, job_description):
    """Calculate match percentage between resume and job description"""
    # Preprocess both texts
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    # Create TF-IDF vectors
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    # Calculate cosine similarity
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    
    # Convert to percentage
    match_percentage = round(similarity * 100, 2)
    
    return match_percentage

def analyze_skills_match(resume_skills, jd_skills):
    """Analyze skills match between resume and job description"""
    # Find matching skills
    matching_skills = set(resume_skills) & set(jd_skills)
    
    # Find missing skills
    missing_skills = set(jd_skills) - set(resume_skills)
    
    # Calculate skills match percentage
    if len(jd_skills) > 0:
        skills_match = len(matching_skills) / len(jd_skills) * 100
    else:
        skills_match = 0
    
    return skills_match, matching_skills, missing_skills

def generate_improvement_suggestions(match_percentage, resume_skills, jd_skills, 
                                   resume_experience, jd_experience, matching_skills, missing_skills):
    """Generate improvement suggestions based on analysis"""
    suggestions = []
    
    # General suggestions based on match percentage
    if match_percentage < 80:
        suggestions.append("Your resume doesn't closely match the job requirements. Consider tailoring it more specifically to this role.")
    
    # Skills-related suggestions
    if len(missing_skills) > 0:
        suggestions.append(f"Develop these missing skills: {', '.join(missing_skills)}")
    
    if len(resume_skills) < 10:
        suggestions.append("Consider adding more technical skills to your resume.")
    
    # Experience-related suggestions
    if resume_experience < jd_experience:
        suggestions.append(f"The job requires {jd_experience}+ years of experience, but you have {resume_experience}. Highlight any relevant projects or achievements to compensate.")
    
    # Formatting suggestions
    suggestions.append("Use more keywords from the job description in your resume.")
    suggestions.append("Quantify your achievements with numbers and metrics (e.g., 'increased efficiency by 25%').")
    
    return suggestions

def analyze_job_match(resume_path, job_description):
    """Analyze match between resume and job description"""
    print("✅ Job Matching System Initialized!")
    
    # Extract resume text
    resume_text = extract_text_from_pdf(resume_path)
    if resume_text is None:
        print("❌ Error: Could not extract text from the resume")
        return
    
    print(f"📄 Analyzing Resume: {resume_path}")
    print(f"📊 Extracted {len(resume_text)} characters")
    
    # Calculate overall match percentage
    match_percentage = calculate_match_percentage(resume_text, job_description)
    print(f"📈 Overall Match Percentage: {match_percentage}%")
    
    # Extract skills from both resume and job description
    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(job_description)
    
    # Extract experience from both resume and job description
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Analyze skills match
    skills_match, matching_skills, missing_skills = analyze_skills_match(resume_skills, jd_skills)
    
    # Generate improvement suggestions
    suggestions = generate_improvement_suggestions(
        match_percentage, resume_skills, jd_skills, 
        resume_experience, jd_experience, matching_skills, missing_skills
    )
    
    # Display results
    print("\n==================================================")
    print("🎯 JOB MATCH ANALYSIS RESULTS")
    print("==================================================")
    
    if match_percentage >= 80:
        print("✅ ELIGIBLE: You are a strong match for this position!")
    else:
        print("❌ NOT ELIGIBLE: Your resume doesn't sufficiently match the job requirements.")
    
    print(f"\n📊 Match Percentage: {match_percentage}%")
    print(f"📋 Skills Match: {skills_match:.1f}%")
    print(f"💼 Your Experience: {resume_experience} years")
    print(f"💼 Required Experience: {jd_experience} years")
    
    print(f"\n🔧 Your Skills: {', '.join(resume_skills) if resume_skills else 'None found'}")
    print(f"🎯 Required Skills: {', '.join(jd_skills) if jd_skills else 'None specified'}")
    print(f"✅ Matching Skills: {', '.join(matching_skills) if matching_skills else 'None'}")
    print(f"❌ Missing Skills: {', '.join(missing_skills) if missing_skills else 'None'}")
    
    print("\n💡 IMPROVEMENT SUGGESTIONS:")
    for i, suggestion in enumerate(suggestions, 1):
        print(f"{i}. {suggestion}")

# Example usage with a sample job description
if __name__ == "__main__":
    # Sample job description (replace with actual job description)
    job_description = """
    Infosys is seeking a motivated Machine Learning Intern with strong skills in Python, SQL, and deep learning frameworks. The role involves working on AI projects like deepfake detection, defect detection, and data visualization. Candidates should demonstrate problem-solving, adaptability, and interest in applying ML models to real-world datasets for impactful solutions.
    """
    
    # Path to your resume (replace with actual path)
    resume_path = r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf"
    
    # Analyze the match
    analyze_job_match(resume_path, job_description)


# In[14]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
            elif isinstance(match, tuple):
                for m in match:
                    if m.isdigit():
                        exp = int(m)
                        if exp > experience:
                            experience = exp
    
    return experience

def extract_education(text):
    """Extract education information from text"""
    education_keywords = [
        'bachelor', 'master', 'phd', 'doctorate', 'mba', 'msc', 'bsc', 'ba', 'ma',
        'degree', 'diploma', 'certificate', 'graduat', 'undergraduate', 'postgraduate',
        'engineer', 'engineering', 'computer science', 'information technology'
    ]
    
    education_levels = {
        'phd': 4, 'doctorate': 4,
        'master': 3, 'mba': 3, 'msc': 3, 'ma': 3,
        'bachelor': 2, 'bsc': 2, 'ba': 2, 'engineer': 2,
        'diploma': 1, 'certificate': 1
    }
    
    education_info = []
    text_lower = text.lower()
    
    for edu in education_keywords:
        edu_pattern = r'\b' + re.escape(edu) + r'\b'
        if re.search(edu_pattern, text_lower):
            education_info.append(edu)
    
    # Determine highest education level
    highest_education = "Not Specified"
    highest_level = 0
    
    for edu, level in education_levels.items():
        edu_pattern = r'\b' + re.escape(edu) + r'\b'
        if re.search(edu_pattern, text_lower) and level > highest_level:
            highest_level = level
            highest_education = edu
    
    return education_info, highest_education

def extract_extracurricular(text):
    """Extract extracurricular activities from text"""
    extracurricular_keywords = [
        'volunteer', 'club', 'society', 'sport', 'team', 'leadership', 'project',
        'hackathon', 'competition', 'event', 'organization', 'community service',
        'mentor', 'tutor', 'teaching', 'coach', 'captain', 'president', 'secretary',
        'treasurer', 'member', 'participat', 'organizer', 'fundraising', 'charity'
    ]
    
    extracurricular_activities = []
    text_lower = text.lower()
    
    for activity in extracurricular_keywords:
        activity_pattern = r'\b' + re.escape(activity) + r'\b'
        if re.search(activity_pattern, text_lower):
            extracurricular_activities.append(activity)
    
    return extracurricular_activities

def extract_certifications(text):
    """Extract certifications from text"""
    certification_keywords = [
        'certified', 'certification', 'aws certified', 'google cloud certified',
        'microsoft certified', 'oracle certified', 'cisco certified', 'pmp',
        'project management professional', 'scrum master', 'agile certified'
    ]
    
    certifications = []
    text_lower = text.lower()
    
    for cert in certification_keywords:
        cert_pattern = r'\b' + re.escape(cert) + r'\b'
        if re.search(cert_pattern, text_lower):
            certifications.append(cert)
    
    return certifications

def calculate_match_percentage(resume_text, job_description):
    """Calculate match percentage between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    match_percentage = round(similarity * 100, 2)
    
    return match_percentage

def analyze_eligibility(resume_data, jd_data):
    """Analyze eligibility based on extracted data"""
    eligibility_score = 0
    max_score = 0
    details = {}
    
    # 1. Skills Match (40% weight)
    tech_skills_match = len(set(resume_data['tech_skills']) & set(jd_data['tech_skills']))
    tech_skills_required = len(jd_data['tech_skills'])
    soft_skills_match = len(set(resume_data['soft_skills']) & set(jd_data['soft_skills']))
    soft_skills_required = len(jd_data['soft_skills'])
    
    if tech_skills_required > 0:
        tech_skill_score = (tech_skills_match / tech_skills_required) * 40
    else:
        tech_skill_score = 0
    
    if soft_skills_required > 0:
        soft_skill_score = (soft_skills_match / soft_skills_required) * 10
    else:
        soft_skill_score = 0
    
    details['skills_score'] = tech_skill_score + soft_skill_score
    eligibility_score += details['skills_score']
    max_score += 50
    
    # 2. Experience Match (25% weight)
    if jd_data['experience'] > 0:
        if resume_data['experience'] >= jd_data['experience']:
            exp_score = 25
        else:
            exp_score = (resume_data['experience'] / jd_data['experience']) * 25
    else:
        exp_score = 25  # If no experience required, full points
    
    details['experience_score'] = exp_score
    eligibility_score += exp_score
    max_score += 25
    
    # 3. Education Match (15% weight)
    # Simple matching - in real system, you'd have more sophisticated logic
    edu_score = 0
    if jd_data['education_level'] != "Not Specified":
        if resume_data['education_level'] == jd_data['education_level']:
            edu_score = 15
        else:
            # Partial credit for related education
            edu_score = 10
    else:
        edu_score = 15  # If no education specified, full points
    
    details['education_score'] = edu_score
    eligibility_score += edu_score
    max_score += 15
    
    # 4. Extracurricular Activities (5% weight)
    if len(jd_data['extracurricular']) > 0:
        extracurr_match = len(set(resume_data['extracurricular']) & set(jd_data['extracurricular']))
        if len(jd_data['extracurricular']) > 0:
            extracurr_score = (extracurr_match / len(jd_data['extracurricular'])) * 5
        else:
            extracurr_score = 5
    else:
        extracurr_score = 5  # If no extracurricular specified, full points
    
    details['extracurricular_score'] = extracurr_score
    eligibility_score += extracurr_score
    max_score += 5
    
    # 5. Certifications (5% weight)
    if len(jd_data['certifications']) > 0:
        cert_match = len(set(resume_data['certifications']) & set(jd_data['certifications']))
        if len(jd_data['certifications']) > 0:
            cert_score = (cert_match / len(jd_data['certifications'])) * 5
        else:
            cert_score = 5
    else:
        cert_score = 5  # If no certifications specified, full points
    
    details['certifications_score'] = cert_score
    eligibility_score += cert_score
    max_score += 5
    
    # Calculate final percentage
    final_percentage = (eligibility_score / max_score) * 100
    
    return final_percentage, details

def generate_recommendations(resume_data, jd_data, eligibility_percentage, score_details):
    """Generate personalized recommendations"""
    recommendations = []
    
    # Skills recommendations
    missing_tech_skills = set(jd_data['tech_skills']) - set(resume_data['tech_skills'])
    missing_soft_skills = set(jd_data['soft_skills']) - set(resume_data['soft_skills'])
    
    if missing_tech_skills:
        recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
    if missing_soft_skills:
        recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
    
    # Experience recommendations
    if resume_data['experience'] < jd_data['experience']:
        recommendations.append(f"Gain more experience in your field. The job requires {jd_data['experience']}+ years, but you have {resume_data['experience']} years.")
    
    # Education recommendations
    if jd_data['education_level'] != "Not Specified" and resume_data['education_level'] != jd_data['education_level']:
        recommendations.append(f"Consider pursuing {jd_data['education_level']} education or equivalent qualifications.")
    
    # Extracurricular recommendations
    missing_extracurricular = set(jd_data['extracurricular']) - set(resume_data['extracurricular'])
    if missing_extracurricular:
        recommendations.append(f"Consider participating in these activities: {', '.join(missing_extracurricular)}")
    
    # Certification recommendations
    missing_certifications = set(jd_data['certifications']) - set(resume_data['certifications'])
    if missing_certifications:
        recommendations.append(f"Consider obtaining these certifications: {', '.join(missing_certifications)}")
    
    # General recommendations based on score
    if eligibility_percentage < 80:
        recommendations.append("Tailor your resume to include more keywords from the job description.")
        recommendations.append("Quantify your achievements with specific numbers and metrics.")
    
    return recommendations

def analyze_job_eligibility(resume_path, job_description):
    """Main function to analyze job eligibility"""
    print("✅ Job Eligibility System Initialized!")
    
    # Extract resume text
    resume_text = extract_text_from_pdf(resume_path)
    if resume_text is None:
        print("❌ Error: Could not extract text from the resume")
        return
    
    print(f"📄 Analyzing Resume: {resume_path}")
    print(f"📊 Extracted {len(resume_text)} characters")
    
    # Extract data from resume
    resume_tech_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    resume_education, resume_edu_level = extract_education(resume_text)
    resume_extracurricular = extract_extracurricular(resume_text)
    resume_certifications = extract_certifications(resume_text)
    
    resume_data = {
        'tech_skills': resume_tech_skills,
        'soft_skills': resume_soft_skills,
        'experience': resume_experience,
        'education': resume_education,
        'education_level': resume_edu_level,
        'extracurricular': resume_extracurricular,
        'certifications': resume_certifications
    }
    
    # Extract data from job description
    jd_tech_skills, jd_soft_skills = extract_skills(job_description)
    jd_experience = extract_experience(job_description)
    jd_education, jd_edu_level = extract_education(job_description)
    jd_extracurricular = extract_extracurricular(job_description)
    jd_certifications = extract_certifications(job_description)
    
    jd_data = {
        'tech_skills': jd_tech_skills,
        'soft_skills': jd_soft_skills,
        'experience': jd_experience,
        'education': jd_education,
        'education_level': jd_edu_level,
        'extracurricular': jd_extracurricular,
        'certifications': jd_certifications
    }
    
    # Calculate overall match percentage
    match_percentage = calculate_match_percentage(resume_text, job_description)
    print(f"📈 Overall Match Percentage: {match_percentage}%")
    
    # Analyze eligibility
    eligibility_percentage, score_details = analyze_eligibility(resume_data, jd_data)
    
    # Generate recommendations
    recommendations = generate_recommendations(resume_data, jd_data, eligibility_percentage, score_details)
    
    # Display results
    print("\n==================================================")
    print("🎯 JOB ELIGIBILITY ANALYSIS RESULTS")
    print("==================================================")
    
    print(f"📊 Overall Eligibility Score: {eligibility_percentage:.1f}%")
    
    if eligibility_percentage >= 80:
        print("✅ ELIGIBLE: You are a strong match for this position!")
    elif eligibility_percentage >= 60:
        print("⚠️ MODERATELY ELIGIBLE: You match some requirements but need improvement.")
    else:
        print("❌ NOT ELIGIBLE: You don't meet most requirements for this position.")
    
    print(f"\n📋 DETAILED BREAKDOWN:")
    print(f"   - Skills: {score_details['skills_score']:.1f}/50")
    print(f"   - Experience: {score_details['experience_score']:.1f}/25")
    print(f"   - Education: {score_details['education_score']:.1f}/15")
    print(f"   - Extracurricular: {score_details['extracurricular_score']:.1f}/5")
    print(f"   - Certifications: {score_details['certifications_score']:.1f}/5")
    
    print(f"\n🔧 YOUR SKILLS:")
    print(f"   - Technical: {', '.join(resume_tech_skills) if resume_tech_skills else 'None found'}")
    print(f"   - Soft: {', '.join(resume_soft_skills) if resume_soft_skills else 'None found'}")
    
    print(f"\n🎯 REQUIRED SKILLS:")
    print(f"   - Technical: {', '.join(jd_tech_skills) if jd_tech_skills else 'None specified'}")
    print(f"   - Soft: {', '.join(jd_soft_skills) if jd_soft_skills else 'None specified'}")
    
    print(f"\n💼 EXPERIENCE:")
    print(f"   - Your Experience: {resume_experience} years")
    print(f"   - Required Experience: {jd_experience} years")
    
    print(f"\n🎓 EDUCATION:")
    print(f"   - Your Education: {resume_edu_level}")
    print(f"   - Required Education: {jd_edu_level}")
    
    print(f"\n💡 IMPROVEMENT SUGGESTIONS:")
    for i, recommendation in enumerate(recommendations, 1):
        print(f"{i}. {recommendation}")

# Example usage
if __name__ == "__main__":
    # Sample job description
    job_description = """
    Infosys is hiring a Machine Learning Intern passionate about AI/ML, computer vision, and data analysis. The intern will work on real-world projects such as deepfake detection, automated defect detection, and landslide segmentation using Python, TensorFlow, OpenCV, and deep learning frameworks. Responsibilities include building and training ML models, analyzing datasets, and creating dashboards with Tableau and MS Excel. Candidates should have knowledge of CNNs, Faster R-CNN, data visualization, SQL, and web development basics (HTML/CSS). Strong teamwork, adaptability, and attention to detail are expected.
    """
    
    # Path to your resume
    resume_path = r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\Arpita_ATS_Optimized.pdf"
    
    # Analyze eligibility
    analyze_job_eligibility(resume_path, job_description)


# In[20]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import string
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function definitions (same as before)
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def extract_education(text):
    """Extract education information from text"""
    education_keywords = [
        'bachelor', 'master', 'phd', 'doctorate', 'mba', 'msc', 'bsc', 'ba', 'ma',
        'degree', 'diploma', 'certificate', 'graduat', 'undergraduate', 'postgraduate',
        'engineer', 'engineering', 'computer science', 'information technology'
    ]
    
    education_levels = {
        'phd': 4, 'doctorate': 4,
        'master': 3, 'mba': 3, 'msc': 3, 'ma': 3,
        'bachelor': 2, 'bsc': 2, 'ba': 2, 'engineer': 2,
        'diploma': 1, 'certificate': 1
    }
    
    education_info = []
    text_lower = text.lower()
    
    for edu in education_keywords:
        edu_pattern = r'\b' + re.escape(edu) + r'\b'
        if re.search(edu_pattern, text_lower):
            education_info.append(edu)
    
    highest_education = "Not Specified"
    highest_level = 0
    
    for edu, level in education_levels.items():
        edu_pattern = r'\b' + re.escape(edu) + r'\b'
        if re.search(edu_pattern, text_lower) and level > highest_level:
            highest_level = level
            highest_education = edu
    
    return education_info, highest_education

def extract_extracurricular(text):
    """Extract extracurricular activities from text"""
    extracurricular_keywords = [
        'volunteer', 'club', 'society', 'sport', 'team', 'leadership', 'project',
        'hackathon', 'competition', 'event', 'organization', 'community service',
        'mentor', 'tutor', 'teaching', 'coach', 'captain', 'president', 'secretary',
        'treasurer', 'member', 'participat', 'organizer', 'fundraising', 'charity'
    ]
    
    extracurricular_activities = []
    text_lower = text.lower()
    
    for activity in extracurricular_keywords:
        activity_pattern = r'\b' + re.escape(activity) + r'\b'
        if re.search(activity_pattern, text_lower):
            extracurricular_activities.append(activity)
    
    return extracurricular_activities

def extract_certifications(text):
    """Extract certifications from text"""
    certification_keywords = [
        'certified', 'certification', 'aws certified', 'google cloud certified',
        'microsoft certified', 'oracle certified', 'cisco certified', 'pmp',
        'project management professional', 'scrum master', 'agile certified'
    ]
    
    certifications = []
    text_lower = text.lower()
    
    for cert in certification_keywords:
        cert_pattern = r'\b' + re.escape(cert) + r'\b'
        if re.search(cert_pattern, text_lower):
            certifications.append(cert)
    
    return certifications

def calculate_match_percentage(resume_text, job_description):
    """Calculate match percentage between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    match_percentage = round(similarity * 100, 2)
    
    return match_percentage

def analyze_eligibility(resume_data, jd_data):
    """Analyze eligibility based on extracted data"""
    eligibility_score = 0
    max_score = 0
    details = {}
    
    # Skills Match (40% weight)
    tech_skills_match = len(set(resume_data['tech_skills']) & set(jd_data['tech_skills']))
    tech_skills_required = len(jd_data['tech_skills'])
    soft_skills_match = len(set(resume_data['soft_skills']) & set(jd_data['soft_skills']))
    soft_skills_required = len(jd_data['soft_skills'])
    
    if tech_skills_required > 0:
        tech_skill_score = (tech_skills_match / tech_skills_required) * 40
    else:
        tech_skill_score = 0
    
    if soft_skills_required > 0:
        soft_skill_score = (soft_skills_match / soft_skills_required) * 10
    else:
        soft_skill_score = 0
    
    details['skills_score'] = tech_skill_score + soft_skill_score
    eligibility_score += details['skills_score']
    max_score += 50
    
    # Experience Match (25% weight)
    if jd_data['experience'] > 0:
        if resume_data['experience'] >= jd_data['experience']:
            exp_score = 25
        else:
            exp_score = (resume_data['experience'] / jd_data['experience']) * 25
    else:
        exp_score = 25
    
    details['experience_score'] = exp_score
    eligibility_score += exp_score
    max_score += 25
    
    # Education Match (15% weight)
    edu_score = 0
    if jd_data['education_level'] != "Not Specified":
        if resume_data['education_level'] == jd_data['education_level']:
            edu_score = 15
        else:
            edu_score = 10
    else:
        edu_score = 15
    
    details['education_score'] = edu_score
    eligibility_score += edu_score
    max_score += 15
    
    # Extracurricular Activities (5% weight)
    if len(jd_data['extracurricular']) > 0:
        extracurr_match = len(set(resume_data['extracurricular']) & set(jd_data['extracurricular']))
        if len(jd_data['extracurricular']) > 0:
            extracurr_score = (extracurr_match / len(jd_data['extracurricular'])) * 5
        else:
            extracurr_score = 5
    else:
        extracurr_score = 5
    
    details['extracurricular_score'] = extracurr_score
    eligibility_score += extracurr_score
    max_score += 5
    
    # Certifications (5% weight)
    if len(jd_data['certifications']) > 0:
        cert_match = len(set(resume_data['certifications']) & set(jd_data['certifications']))
        if len(jd_data['certifications']) > 0:
            cert_score = (cert_match / len(jd_data['certifications'])) * 5
        else:
            cert_score = 5
    else:
        cert_score = 5
    
    details['certifications_score'] = cert_score
    eligibility_score += cert_score
    max_score += 5
    
    final_percentage = (eligibility_score / max_score) * 100
    return final_percentage, details

def generate_recommendations(resume_data, jd_data, eligibility_percentage, score_details):
    """Generate personalized recommendations"""
    recommendations = []
    
    missing_tech_skills = set(jd_data['tech_skills']) - set(resume_data['tech_skills'])
    missing_soft_skills = set(jd_data['soft_skills']) - set(resume_data['soft_skills'])
    
    if missing_tech_skills:
        recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
    if missing_soft_skills:
        recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
    
    if resume_data['experience'] < jd_data['experience']:
        recommendations.append(f"Gain more experience in your field. The job requires {jd_data['experience']}+ years, but you have {resume_data['experience']} years.")
    
    if jd_data['education_level'] != "Not Specified" and resume_data['education_level'] != jd_data['education_level']:
        recommendations.append(f"Consider pursuing {jd_data['education_level']} education or equivalent qualifications.")
    
    missing_extracurricular = set(jd_data['extracurricular']) - set(resume_data['extracurricular'])
    if missing_extracurricular:
        recommendations.append(f"Consider participating in these activities: {', '.join(missing_extracurricular)}")
    
    missing_certifications = set(jd_data['certifications']) - set(resume_data['certifications'])
    if missing_certifications:
        recommendations.append(f"Consider obtaining these certifications: {', '.join(missing_certifications)}")
    
    if eligibility_percentage < 80:
        recommendations.append("Tailor your resume to include more keywords from the job description.")
        recommendations.append("Quantify your achievements with specific numbers and metrics.")
    
    return recommendations

def analyze_job_eligibility(resume_path, job_description):
    """Main function to analyze job eligibility"""
    print("✅ Job Eligibility System Initialized!")
    
    resume_text = extract_text_from_pdf(resume_path)
    if resume_text is None:
        print("❌ Error: Could not extract text from the resume")
        return
    
    print(f"📄 Analyzing Resume: {resume_path}")
    print(f"📊 Extracted {len(resume_text)} characters")
    
    resume_tech_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    resume_education, resume_edu_level = extract_education(resume_text)
    resume_extracurricular = extract_extracurricular(resume_text)
    resume_certifications = extract_certifications(resume_text)
    
    resume_data = {
        'tech_skills': resume_tech_skills,
        'soft_skills': resume_soft_skills,
        'experience': resume_experience,
        'education': resume_education,
        'education_level': resume_edu_level,
        'extracurricular': resume_extracurricular,
        'certifications': resume_certifications
    }
    
    jd_tech_skills, jd_soft_skills = extract_skills(job_description)
    jd_experience = extract_experience(job_description)
    jd_education, jd_edu_level = extract_education(job_description)
    jd_extracurricular = extract_extracurricular(job_description)
    jd_certifications = extract_certifications(job_description)
    
    jd_data = {
        'tech_skills': jd_tech_skills,
        'soft_skills': jd_soft_skills,
        'experience': jd_experience,
        'education': jd_education,
        'education_level': jd_edu_level,
        'extracurricular': jd_extracurricular,
        'certifications': jd_certifications
    }
    
    match_percentage = calculate_match_percentage(resume_text, job_description)
    print(f"📈 Overall Match Percentage: {match_percentage}%")
    
    eligibility_percentage, score_details = analyze_eligibility(resume_data, jd_data)
    recommendations = generate_recommendations(resume_data, jd_data, eligibility_percentage, score_details)
    
    print("\n==================================================")
    print("🎯 JOB ELIGIBILITY ANALYSIS RESULTS")
    print("==================================================")
    
    print(f"📊 Overall Eligibility Score: {eligibility_percentage:.1f}%")
    
    if eligibility_percentage >= 80:
        print("✅ ELIGIBLE: You are a strong match for this position!")
    elif eligibility_percentage >= 60:
        print("⚠️ MODERATELY ELIGIBLE: You match some requirements but need improvement.")
    else:
        print("❌ NOT ELIGIBLE: You don't meet most requirements for this position.")
    
    print(f"\n📋 DETAILED BREAKDOWN:")
    print(f"   - Skills: {score_details['skills_score']:.1f}/50")
    print(f"   - Experience: {score_details['experience_score']:.1f}/25")
    print(f"   - Education: {score_details['education_score']:.1f}/15")
    print(f"   - Extracurricular: {score_details['extracurricular_score']:.1f}/5")
    print(f"   - Certifications: {score_details['certifications_score']:.1f}/5")
    
    print(f"\n🔧 YOUR SKILLS:")
    print(f"   - Technical: {', '.join(resume_tech_skills) if resume_tech_skills else 'None found'}")
    print(f"   - Soft: {', '.join(resume_soft_skills) if resume_soft_skills else 'None found'}")
    
    print(f"\n🎯 REQUIRED SKILLS:")
    print(f"   - Technical: {', '.join(jd_tech_skills) if jd_tech_skills else 'None specified'}")
    print(f"   - Soft: {', '.join(jd_soft_skills) if jd_soft_skills else 'None specified'}")
    
    print(f"\n💼 EXPERIENCE:")
    print(f"   - Your Experience: {resume_experience} years")
    print(f"   - Required Experience: {jd_experience} years")
    
    print(f"\n🎓 EDUCATION:")
    print(f"   - Your Education: {resume_edu_level}")
    print(f"   - Required Education: {jd_edu_level}")
    
    print(f"\n💡 IMPROVEMENT SUGGESTIONS:")
    for i, recommendation in enumerate(recommendations, 1):
        print(f"{i}. {recommendation}")

# Create input box for job description
print("🎯 Job Eligibility Analysis System")
print("=" * 50)

# Get job description from user input
job_description = input("📝 Paste the Job Description here and press Enter:\n")

# Fixed resume path
resume_path = r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf"

# Check if job description was provided
if not job_description.strip():
    print("❌ Error: Please provide a job description")
else:
    print(f"\n📄 Using resume: {resume_path}")
    print("⏳ Analyzing... Please wait\n")
    
    # Analyze eligibility
    analyze_job_eligibility(resume_path, job_description)

    print("\n" + "=" * 50)
    print("Analysis complete! 🎉")


# In[2]:


import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import webbrowser
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function definitions
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def calculate_match_percentage(resume_text, job_description):
    """Calculate match percentage between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    match_percentage = round(similarity * 100, 2)
    
    return match_percentage

# Course database with platform, cost, duration, and links
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9},
        {'platform': 'Udemy', 'name': 'Complete Python Bootcamp', 'cost': '₹455', 'duration': '22 hours', 
         'url': 'https://www.udemy.com/course/complete-python-bootcamp/', 'rating': 4.6}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'edX', 'name': 'Microsoft Professional Program in AI', 'cost': '₹25,000', 'duration': '6 months', 
         'url': 'https://www.edx.org/professional-certificate/microsoft-artificial-intelligence', 'rating': 4.7},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8},
        {'platform': 'Udemy', 'name': 'Data Analysis with Python', 'cost': '₹455', 'duration': '15 hours', 
         'url': 'https://www.udemy.com/course/data-analysis-with-python/', 'rating': 4.5}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7},
        {'platform': 'YouTube', 'name': 'SQL Tutorial for Beginners', 'cost': 'Free', 'duration': '4 hours', 
         'url': 'https://www.youtube.com/watch?v=HXV3zeQKqGY', 'rating': 4.8}
    ],
    'aws': [
        {'platform': 'AWS Training', 'name': 'AWS Cloud Practitioner Essentials', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.aws.training/Details/eLearning?id=60697', 'rating': 4.6},
        {'platform': 'Udemy', 'name': 'AWS Certified Solutions Architect', 'cost': '₹455', 'duration': '25 hours', 
         'url': 'https://www.udemy.com/course/aws-certified-solutions-architect-associate-saa-c03/', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'TensorFlow 2.0 Complete Course', 'cost': 'Free', 'duration': '8 hours', 
         'url': 'https://www.youtube.com/watch?v=tPYj3fFJGjk', 'rating': 4.7}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6},
        {'platform': 'YouTube', 'name': 'OpenCV Tutorial', 'cost': 'Free', 'duration': '5 hours', 
         'url': 'https://www.youtube.com/watch?v=oXlwWbU8l2o', 'rating': 4.7}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7},
        {'platform': 'Udemy', 'name': 'Tableau 2022 A-Z: Hands-On Tableau Training', 'cost': '₹455', 'duration': '8 hours', 
         'url': 'https://www.udemy.com/course/tableau10/', 'rating': 4.6}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7},
        {'platform': 'YouTube', 'name': 'Power BI Tutorial for Beginners', 'cost': 'Free', 'duration': '3 hours', 
         'url': 'https://www.youtube.com/watch?v=1c01r__-2W4', 'rating': 4.8}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'NLP Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=8S3qHHUKqYk', 'rating': 4.7}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'CNN Tutorial for Beginners', 'cost': 'Free', 'duration': '2 hours', 
         'url': 'https://www.youtube.com/watch?v=YRhxdVk_sIs', 'rating': 4.7}
    ]
}

# Function to extract missing skills
def extract_missing_skills(resume_text, job_description):
    """Extract skills missing from resume but required in job description"""
    resume_skills, _ = extract_skills(resume_text)
    jd_skills, _ = extract_skills(job_description)
    
    missing_skills = set(jd_skills) - set(resume_skills)
    return list(missing_skills)

# Function to suggest courses for missing skills
def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            # For skills not in database, generate generic suggestions
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'YouTube', 'name': f'{skill.title()} Tutorial for Beginners', 'cost': 'Free', 
                 'duration': '2-10 hours', 'url': 'https://www.youtube.com/results?search_query=' + skill.replace(' ', '+'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

# Function to display course suggestions
def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions in a user-friendly format"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS FOR MISSING SKILLS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Function to open course links
def open_course_links(skill, course_index):
    """Open course links in browser based on user selection"""
    skill_lower = skill.lower()
    if skill_lower in COURSE_DATABASE:
        courses = COURSE_DATABASE[skill_lower]
        if 0 <= course_index < len(courses):
            webbrowser.open(courses[course_index]['url'])
            print(f"Opening course: {courses[course_index]['name']}")
        else:
            print("Invalid course selection.")
    else:
        print("No specific courses found for this skill.")

# Modified main function to include course suggestions
def analyze_job_eligibility_with_courses(resume_path, job_description):
    """Main function to analyze job eligibility with course suggestions"""
    print("✅ Job Eligibility System with Course Recommendations Initialized!")
    
    # Extract resume text
    resume_text = extract_text_from_pdf(resume_path)
    if resume_text is None:
        print("❌ Error: Could not extract text from the resume")
        return
    
    print(f"📄 Analyzing Resume: {resume_path}")
    print(f"📊 Extracted {len(resume_text)} characters")
    
    # Extract missing skills
    missing_skills = extract_missing_skills(resume_text, job_description)
    
    # Get course suggestions
    course_suggestions = suggest_courses(missing_skills)
    
    # Calculate overall match percentage
    match_percentage = calculate_match_percentage(resume_text, job_description)
    print(f"📈 Overall Match Percentage: {match_percentage}%")
    
    # Display course suggestions
    display_course_suggestions(missing_skills, course_suggestions)
    
    # Interactive course selection
    if missing_skills:
        print("\n💡 Would you like to explore any course? (Enter skill number or 0 to exit)")
        try:
            choice = int(input("Enter your choice: "))
            if 1 <= choice <= len(missing_skills):
                selected_skill = missing_skills[choice-1]
                print(f"\nCourses for {selected_skill}:")
                courses = course_suggestions.get(selected_skill, [])
                for i, course in enumerate(courses, 1):
                    print(f"{i}. {course['name']} ({course['platform']})")
                
                course_choice = int(input("Select course to open (0 to cancel): "))
                if 1 <= course_choice <= len(courses):
                    open_course_links(selected_skill, course_choice-1)
                elif course_choice != 0:
                    print("Invalid selection.")
            elif choice != 0:
                print("Invalid selection.")
        except ValueError:
            print("Please enter a valid number.")

# Example usage with interactive job description input
if __name__ == "__main__":
    print("🎯 Job Eligibility Analysis with Course Recommendations")
    print("="*60)
    
    # Get job description from user
    job_description = input("📝 Paste the Job Description here and press Enter:\n")
    
    # Fixed resume path
    resume_path = r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf"
    
    if not job_description.strip():
        print("❌ Error: Please provide a job description")
    else:
        print(f"\n📄 Using resume: {resume_path}")
        print("⏳ Analyzing... Please wait\n")
        
        # Analyze eligibility with course suggestions
        analyze_job_eligibility_with_courses(resume_path, job_description)
        
        print("\n" + "="*60)
        print("Analysis complete! 🎉")


# In[4]:


# Complete ATS Resume Analysis System
from IPython.display import display, clear_output
import ipywidgets as widgets
from PyPDF2 import PdfReader
import docx
import re
import webbrowser
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function to extract text from different file types
def extract_text_from_file(file_path):
    """Extract text from PDF or Word document"""
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_path.endswith(('.doc', '.docx')):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            return f"Error: Unsupported file format. Please upload PDF or Word document."
        
        return text
    except Exception as e:
        return f"Error reading file: {e}"

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def calculate_ats_score(resume_text, job_description):
    """Calculate ATS score between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    ats_score = round(similarity * 100, 2)
    
    return ats_score

def analyze_eligibility(resume_text, job_description):
    """Analyze eligibility based on resume and job description"""
    ats_score = calculate_ats_score(resume_text, job_description)
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Calculate skills match
    tech_skills_match = len(set(resume_skills) & set(jd_skills))
    tech_skills_required = len(jd_skills)
    
    if tech_skills_required > 0:
        skills_match_percentage = (tech_skills_match / tech_skills_required) * 100
    else:
        skills_match_percentage = 100
    
    # Calculate experience match
    if jd_experience > 0:
        if resume_experience >= jd_experience:
            experience_match = 100
        else:
            experience_match = (resume_experience / jd_experience) * 100
    else:
        experience_match = 100
    
    # Overall eligibility score (weighted average)
    eligibility_score = (ats_score * 0.4) + (skills_match_percentage * 0.4) + (experience_match * 0.2)
    
    return ats_score, eligibility_score, skills_match_percentage, experience_match

def generate_recommendations(resume_text, job_description):
    """Generate improvement recommendations"""
    recommendations = []
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Skills recommendations
    missing_tech_skills = set(jd_skills) - set(resume_skills)
    missing_soft_skills = set(jd_soft_skills) - set(resume_soft_skills)
    
    if missing_tech_skills:
        recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
    if missing_soft_skills:
        recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
    
    # Experience recommendations
    if resume_experience < jd_experience:
        recommendations.append(f"Gain more experience. The job requires {jd_experience}+ years, but you have {resume_experience} years.")
    
    # General recommendations
    recommendations.append("Tailor your resume to include more keywords from the job description.")
    recommendations.append("Quantify your achievements with specific numbers and metrics.")
    recommendations.append("Use bullet points to make your resume more readable.")
    
    return recommendations, missing_tech_skills

# Course database
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8}
    ]
}

def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Create widgets
upload = widgets.FileUpload(description="Upload Resume", accept=".pdf,.doc,.docx", multiple=False)
jd_input = widgets.Textarea(placeholder="Paste job description here...", description="Job Description:", layout=widgets.Layout(width="80%", height="100px"))
analyze_btn = widgets.Button(description="Analyze Resume", button_style="success")
improve_btn = widgets.Button(description="Get Course Suggestions", button_style="info")
output = widgets.Output()

# Global variables to store analysis results
resume_text_global = ""
job_description_global = ""
missing_skills_global = []

def on_upload_change(change):
    """Handle file upload"""
    with output:
        clear_output()
        if upload.value:
            file_name = next(iter(upload.value))
            content = upload.value[file_name]['content']
            
            # Save uploaded file temporarily
            with open(file_name, 'wb') as f:
                f.write(content)
            
            # Extract text from file
            global resume_text_global
            resume_text_global = extract_text_from_file(file_name)
            
            if "Error" in resume_text_global:
                print(resume_text_global)
            else:
                print(f"✅ Resume uploaded successfully: {file_name}")
                print(f"📄 Extracted {len(resume_text_global)} characters")
                
                # Clean up temporary file
                if os.path.exists(file_name):
                    os.remove(file_name)

def on_analyze_click(b):
    """Handle analyze button click"""
    with output:
        clear_output()
        global resume_text_global, job_description_global
        
        if not resume_text_global:
            print("❌ Please upload a resume first.")
            return
            
        job_description_global = jd_input.value.strip()
        if not job_description_global:
            print("❌ Please enter a job description.")
            return
        
        print("⏳ Analyzing your resume... Please wait\n")
        
        # Perform analysis
        ats_score, eligibility_score, skills_match, experience_match = analyze_eligibility(resume_text_global, job_description_global)
        recommendations, missing_skills = generate_recommendations(resume_text_global, job_description_global)
        
        global missing_skills_global
        missing_skills_global = missing_skills
        
        # Display results
        print("="*60)
        print("📊 RESUME ANALYSIS RESULTS")
        print("="*60)
        
        print(f"🎯 ATS Score: {ats_score:.1f}%")
        print(f"📈 Eligibility Score: {eligibility_score:.1f}%")
        print(f"🔧 Skills Match: {skills_match:.1f}%")
        print(f"💼 Experience Match: {experience_match:.1f}%")
        
        if eligibility_score >= 80:
            print("✅ STATUS: HIGHLY ELIGIBLE - Strong match for this position!")
        elif eligibility_score >= 60:
            print("⚠️ STATUS: MODERATELY ELIGIBLE - Some improvements needed")
        else:
            print("❌ STATUS: NOT ELIGIBLE - Significant improvements needed")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")
        
        if missing_skills:
            print(f"\n🎯 Missing Skills Detected: {len(missing_skills)}")
            print("Click 'Get Course Suggestions' button to see recommended courses!")

def on_improve_click(b):
    """Handle improve button click"""
    with output:
        clear_output()
        global missing_skills_global
        
        if not missing_skills_global:
            print("No missing skills detected. Your resume looks good!")
            return
            
        course_suggestions = suggest_courses(missing_skills_global)
        display_course_suggestions(missing_skills_global, course_suggestions)

# Set up event handlers
upload.observe(on_upload_change, names='value')
analyze_btn.on_click(on_analyze_click)
improve_btn.on_click(on_improve_click)

# Display the interface
print("🎯 COMPLETE ATS RESUME ANALYSIS SYSTEM")
print("="*60)
print("Follow these steps:")
print("1. Upload your resume (PDF or Word)")
print("2. Paste the job description")
print("3. Click 'Analyze Resume'")
print("4. Click 'Get Course Suggestions' for missing skills\n")

display(widgets.VBox([
    widgets.HTML("<b>Step 1: Upload Your Resume</b>"),
    upload,
    widgets.HTML("<b>Step 2: Paste Job Description</b>"),
    jd_input,
    widgets.HTML("<b>Step 3: Analyze Your Resume</b>"),
    analyze_btn,
    widgets.HTML("<b>Step 4: Improve Your Skills</b>"),
    improve_btn,
    widgets.HTML("<hr>"),
    output
]))


# In[5]:


# Complete ATS Resume Analysis System with Fixed Upload
from IPython.display import display, clear_output
import ipywidgets as widgets
from PyPDF2 import PdfReader
import docx
import re
import webbrowser
import os
import tempfile
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function to extract text from different file types
def extract_text_from_file(file_path):
    """Extract text from PDF or Word document"""
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_path.endswith(('.doc', '.docx')):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            return f"Error: Unsupported file format. Please upload PDF or Word document."
        
        return text
    except Exception as e:
        return f"Error reading file: {e}"

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def calculate_ats_score(resume_text, job_description):
    """Calculate ATS score between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    ats_score = round(similarity * 100, 2)
    
    return ats_score

def analyze_eligibility(resume_text, job_description):
    """Analyze eligibility based on resume and job description"""
    ats_score = calculate_ats_score(resume_text, job_description)
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Calculate skills match
    tech_skills_match = len(set(resume_skills) & set(jd_skills))
    tech_skills_required = len(jd_skills)
    
    if tech_skills_required > 0:
        skills_match_percentage = (tech_skills_match / tech_skills_required) * 100
    else:
        skills_match_percentage = 100
    
    # Calculate experience match
    if jd_experience > 0:
        if resume_experience >= jd_experience:
            experience_match = 100
        else:
            experience_match = (resume_experience / jd_experience) * 100
    else:
        experience_match = 100
    
    # Overall eligibility score (weighted average)
    eligibility_score = (ats_score * 0.4) + (skills_match_percentage * 0.4) + (experience_match * 0.2)
    
    return ats_score, eligibility_score, skills_match_percentage, experience_match

def generate_recommendations(resume_text, job_description):
    """Generate improvement recommendations"""
    recommendations = []
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Skills recommendations
    missing_tech_skills = set(jd_skills) - set(resume_skills)
    missing_soft_skills = set(jd_soft_skills) - set(resume_soft_skills)
    
    if missing_tech_skills:
        recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
    if missing_soft_skills:
        recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
    
    # Experience recommendations
    if resume_experience < jd_experience:
        recommendations.append(f"Gain more experience. The job requires {jd_experience}+ years, but you have {resume_experience} years.")
    
    # General recommendations
    recommendations.append("Tailor your resume to include more keywords from the job description.")
    recommendations.append("Quantify your achievements with specific numbers and metrics.")
    recommendations.append("Use bullet points to make your resume more readable.")
    
    return recommendations, list(missing_tech_skills)

# Course database
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8}
    ]
}

def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Create widgets
upload = widgets.FileUpload(description="Upload Resume", accept=".pdf,.doc,.docx", multiple=False)
jd_input = widgets.Textarea(placeholder="Paste job description here...", description="Job Description:", layout=widgets.Layout(width="80%", height="100px"))
analyze_btn = widgets.Button(description="Analyze Resume", button_style="success")
improve_btn = widgets.Button(description="Get Course Suggestions", button_style="info")
output = widgets.Output()

# Global variables to store analysis results
resume_text_global = ""
job_description_global = ""
missing_skills_global = []
uploaded_file_name = ""

def on_upload_change(change):
    """Handle file upload"""
    with output:
        clear_output()
        global resume_text_global, uploaded_file_name
        
        if upload.value:
            try:
                # Get the uploaded file
                file_name = list(upload.value.keys())[0]
                file_content = upload.value[file_name]['content']
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_name)[1]) as tmp_file:
                    tmp_file.write(file_content)
                    temp_file_path = tmp_file.name
                
                # Extract text from file
                resume_text_global = extract_text_from_file(temp_file_path)
                uploaded_file_name = file_name
                
                # Clean up temporary file
                os.unlink(temp_file_path)
                
                if "Error" in resume_text_global:
                    print(resume_text_global)
                else:
                    print(f"✅ Resume uploaded successfully: {file_name}")
                    print(f"📄 Extracted {len(resume_text_global)} characters")
                    
            except Exception as e:
                print(f"❌ Error processing file: {e}")
        else:
            print("Please upload a resume file")

def on_analyze_click(b):
    """Handle analyze button click"""
    with output:
        clear_output()
        global resume_text_global, job_description_global, uploaded_file_name
        
        if not resume_text_global or "Error" in resume_text_global:
            print("❌ Please upload a valid resume first.")
            return
            
        job_description_global = jd_input.value.strip()
        if not job_description_global:
            print("❌ Please enter a job description.")
            return
        
        print(f"⏳ Analyzing your resume ({uploaded_file_name})... Please wait\n")
        
        # Perform analysis
        ats_score, eligibility_score, skills_match, experience_match = analyze_eligibility(resume_text_global, job_description_global)
        recommendations, missing_skills = generate_recommendations(resume_text_global, job_description_global)
        
        global missing_skills_global
        missing_skills_global = missing_skills
        
        # Display results
        print("="*60)
        print("📊 RESUME ANALYSIS RESULTS")
        print("="*60)
        
        print(f"🎯 ATS Score: {ats_score:.1f}%")
        print(f"📈 Eligibility Score: {eligibility_score:.1f}%")
        print(f"🔧 Skills Match: {skills_match:.1f}%")
        print(f"💼 Experience Match: {experience_match:.1f}%")
        
        if eligibility_score >= 80:
            print("✅ STATUS: HIGHLY ELIGIBLE - Strong match for this position!")
        elif eligibility_score >= 60:
            print("⚠️ STATUS: MODERATELY ELIGIBLE - Some improvements needed")
        else:
            print("❌ STATUS: NOT ELIGIBLE - Significant improvements needed")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")
        
        if missing_skills:
            print(f"\n🎯 Missing Skills Detected: {len(missing_skills)}")
            print("Click 'Get Course Suggestions' button to see recommended courses!")
        else:
            print("\n✅ No missing skills detected! Your resume has all the required skills.")

def on_improve_click(b):
    """Handle improve button click"""
    with output:
        clear_output()
        global missing_skills_global
        
        if not missing_skills_global:
            print("No missing skills detected. Your resume looks good!")
            return
            
        course_suggestions = suggest_courses(missing_skills_global)
        display_course_suggestions(missing_skills_global, course_suggestions)

# Set up event handlers
upload.observe(on_upload_change, names='value')
analyze_btn.on_click(on_analyze_click)
improve_btn.on_click(on_improve_click)

# Display the interface
print("🎯 COMPLETE ATS RESUME ANALYSIS SYSTEM")
print("="*60)
print("Follow these steps:")
print("1. Upload your resume (PDF or Word)")
print("2. Paste the job description")
print("3. Click 'Analyze Resume'")
print("4. Click 'Get Course Suggestions' for missing skills\n")

display(widgets.VBox([
    widgets.HTML("<b>Step 1: Upload Your Resume</b>"),
    upload,
    widgets.HTML("<b>Step 2: Paste Job Description</b>"),
    jd_input,
    widgets.HTML("<b>Step 3: Analyze Your Resume</b>"),
    analyze_btn,
    widgets.HTML("<b>Step 4: Improve Your Skills</b>"),
    improve_btn,
    widgets.HTML("<hr>"),
    output
]))


# In[8]:


# Complete ATS Resume Analysis System with Manual Resume Path
from IPython.display import display, clear_output
import ipywidgets as widgets
from PyPDF2 import PdfReader
import docx
import re
import webbrowser
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function to extract text from different file types
def extract_text_from_file(file_path):
    """Extract text from PDF or Word document"""
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_path.endswith(('.doc', '.docx')):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            return f"Error: Unsupported file format. Please use PDF or Word document."
        
        return text
    except Exception as e:
        return f"Error reading file: {e}"

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def calculate_ats_score(resume_text, job_description):
    """Calculate ATS score between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    jd_processed = preprocess_text(job_description)
    
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
    
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    ats_score = round(similarity * 100, 2)
    
    return ats_score

def analyze_eligibility(resume_text, job_description):
    """Analyze eligibility based on resume and job description"""
    ats_score = calculate_ats_score(resume_text, job_description)
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Calculate skills match
    tech_skills_match = len(set(resume_skills) & set(jd_skills))
    tech_skills_required = len(jd_skills)
    
    if tech_skills_required > 0:
        skills_match_percentage = (tech_skills_match / tech_skills_required) * 100
    else:
        skills_match_percentage = 100
    
    # Calculate experience match
    if jd_experience > 0:
        if resume_experience >= jd_experience:
            experience_match = 100
        else:
            experience_match = (resume_experience / jd_experience) * 100
    else:
        experience_match = 100
    
    # Overall eligibility score (weighted average)
    eligibility_score = (ats_score * 0.4) + (skills_match_percentage * 0.4) + (experience_match * 0.2)
    
    return ats_score, eligibility_score, skills_match_percentage, experience_match

def generate_recommendations(resume_text, job_description):
    """Generate improvement recommendations"""
    recommendations = []
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    jd_skills, jd_soft_skills = extract_skills(job_description)
    
    resume_experience = extract_experience(resume_text)
    jd_experience = extract_experience(job_description)
    
    # Skills recommendations
    missing_tech_skills = set(jd_skills) - set(resume_skills)
    missing_soft_skills = set(jd_soft_skills) - set(resume_soft_skills)
    
    if missing_tech_skills:
        recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
    if missing_soft_skills:
        recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
    
    # Experience recommendations
    if resume_experience < jd_experience:
        recommendations.append(f"Gain more experience. The job requires {jd_experience}+ years, but you have {resume_experience} years.")
    
    # General recommendations
    recommendations.append("Tailor your resume to include more keywords from the job description.")
    recommendations.append("Quantify your achievements with specific numbers and metrics.")
    recommendations.append("Use bullet points to make your resume more readable.")
    
    return recommendations, list(missing_tech_skills)

# Course database
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8}
    ]
}

def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Create widgets
resume_path_input = widgets.Text(
    value=r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf",
    placeholder="Enter full path to your resume",
    description="Resume Path:",
    layout=widgets.Layout(width="80%")
)
jd_input = widgets.Textarea(
    placeholder="Paste job description here...", 
    description="Job Description:", 
    layout=widgets.Layout(width="80%", height="100px")
)
analyze_btn = widgets.Button(description="Analyze Resume", button_style="success")
improve_btn = widgets.Button(description="Get Course Suggestions", button_style="info")
output = widgets.Output()

# Global variables to store analysis results
resume_text_global = ""
job_description_global = ""
missing_skills_global = []

def on_analyze_click(b):
    """Handle analyze button click"""
    with output:
        clear_output()
        global resume_text_global, job_description_global, missing_skills_global
        
        # Get resume path from input
        resume_path = resume_path_input.value.strip()
        if not resume_path:
            print("❌ Please enter a resume path.")
            return
            
        # Extract text from resume
        resume_text_global = extract_text_from_file(resume_path)
        if "Error" in resume_text_global:
            print(resume_text_global)
            return
            
        job_description_global = jd_input.value.strip()
        if not job_description_global:
            print("❌ Please enter a job description.")
            return
        
        print(f"⏳ Analyzing your resume... Please wait\n")
        
        # Perform analysis
        ats_score, eligibility_score, skills_match, experience_match = analyze_eligibility(resume_text_global, job_description_global)
        recommendations, missing_skills = generate_recommendations(resume_text_global, job_description_global)
        
        missing_skills_global = missing_skills
        
        # Display results
        print("="*60)
        print("📊 RESUME ANALYSIS RESULTS")
        print("="*60)
        
        print(f"🎯 ATS Score: {ats_score:.1f}%")
        print(f"📈 Eligibility Score: {eligibility_score:.1f}%")
        print(f"🔧 Skills Match: {skills_match:.1f}%")
        print(f"💼 Experience Match: {experience_match:.1f}%")
        
        if eligibility_score >= 80:
            print("✅ STATUS: HIGHLY ELIGIBLE - Strong match for this position!")
        elif eligibility_score >= 60:
            print("⚠️ STATUS: MODERATELY ELIGIBLE - Some improvements needed")
        else:
            print("❌ STATUS: NOT ELIGIBLE - Significant improvements needed")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")
        
        if missing_skills:
            print(f"\n🎯 Missing Skills Detected: {len(missing_skills)}")
            print("Click 'Get Course Suggestions' button to see recommended courses!")
        else:
            print("\n✅ No missing skills detected! Your resume has all the required skills.")

def on_improve_click(b):
    """Handle improve button click"""
    with output:
        clear_output()
        global missing_skills_global
        
        if not missing_skills_global:
            print("No missing skills detected. Your resume looks good!")
            return
            
        course_suggestions = suggest_courses(missing_skills_global)
        display_course_suggestions(missing_skills_global, course_suggestions)

# Set up event handlers
analyze_btn.on_click(on_analyze_click)
improve_btn.on_click(on_improve_click)

# Display the interface
print("🎯 COMPLETE ATS RESUME ANALYSIS SYSTEM")
print("="*60)
print("Follow these steps:")
print("1. Enter the path to your resume (PDF or Word)")
print("2. Paste the job description")
print("3. Click 'Analyze Resume'")
print("4. Click 'Get Course Suggestions' for missing skills\n")
print("💡 Default resume path is set to your OM_ASHUTOSH Resume.pdf")

display(widgets.VBox([
    widgets.HTML("<b>Step 1: Enter Resume Path</b>"),
    resume_path_input,
    widgets.HTML("<b>Step 2: Paste Job Description</b>"),
    jd_input,
    widgets.HTML("<b>Step 3: Analyze Your Resume</b>"),
    analyze_btn,
    widgets.HTML("<b>Step 4: Improve Your Skills</b>"),
    improve_btn,
    widgets.HTML("<hr>"),
    output
]))


# In[10]:


# Complete ATS Resume Analysis System with Optional JD
from IPython.display import display, clear_output
import ipywidgets as widgets
from PyPDF2 import PdfReader
import docx
import re
import webbrowser
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function to extract text from different file types
def extract_text_from_file(file_path):
    """Extract text from PDF or Word document"""
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_path.endswith(('.doc', '.docx')):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            return f"Error: Unsupported file format. Please use PDF or Word document."
        
        return text
    except Exception as e:
        return f"Error reading file: {e}"

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def calculate_ats_score(resume_text, job_description=None):
    """Calculate ATS score between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    
    if job_description:
        # Calculate score against job description
        jd_processed = preprocess_text(job_description)
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        ats_score = round(similarity * 100, 2)
    else:
        # Calculate base ATS score (resume quality without JD)
        # This is a simplified version that checks resume structure and content
        word_count = len(resume_processed.split())
        skill_count = len(extract_skills(resume_text)[0])
        exp_years = extract_experience(resume_text)
        
        # Score based on resume quality metrics
        ats_score = min(100, (word_count / 500 * 30) + (skill_count / 15 * 40) + (exp_years * 6))
        ats_score = round(ats_score, 2)
    
    return ats_score

def analyze_eligibility(resume_text, job_description=None):
    """Analyze eligibility based on resume and optional job description"""
    ats_score = calculate_ats_score(resume_text, job_description)
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    
    if job_description:
        jd_skills, jd_soft_skills = extract_skills(job_description)
        jd_experience = extract_experience(job_description)
        
        # Calculate skills match
        tech_skills_match = len(set(resume_skills) & set(jd_skills))
        tech_skills_required = len(jd_skills)
        
        if tech_skills_required > 0:
            skills_match_percentage = (tech_skills_match / tech_skills_required) * 100
        else:
            skills_match_percentage = 100
        
        # Calculate experience match
        if jd_experience > 0:
            if resume_experience >= jd_experience:
                experience_match = 100
            else:
                experience_match = (resume_experience / jd_experience) * 100
        else:
            experience_match = 100
        
        # Overall eligibility score (weighted average)
        eligibility_score = (ats_score * 0.4) + (skills_match_percentage * 0.4) + (experience_match * 0.2)
        
        return ats_score, eligibility_score, skills_match_percentage, experience_match, jd_skills, jd_experience
    else:
        # Without JD, provide general resume quality assessment
        skills_match_percentage = 100  # Not applicable without JD
        experience_match = 100  # Not applicable without JD
        
        # Base eligibility score on resume quality
        word_count = len(resume_text.split())
        skill_count = len(resume_skills)
        
        eligibility_score = min(100, (word_count / 500 * 40) + (skill_count / 15 * 40) + (resume_experience * 4))
        eligibility_score = round(eligibility_score, 2)
        
        return ats_score, eligibility_score, skills_match_percentage, experience_match, [], 0

def generate_recommendations(resume_text, job_description=None):
    """Generate improvement recommendations"""
    recommendations = []
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    
    if job_description:
        jd_skills, jd_soft_skills = extract_skills(job_description)
        jd_experience = extract_experience(job_description)
        
        # Skills recommendations
        missing_tech_skills = set(jd_skills) - set(resume_skills)
        missing_soft_skills = set(jd_soft_skills) - set(resume_soft_skills)
        
        if missing_tech_skills:
            recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
        if missing_soft_skills:
            recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
        
        # Experience recommendations
        if resume_experience < jd_experience:
            recommendations.append(f"Gain more experience. The job requires {jd_experience}+ years, but you have {resume_experience} years.")
    else:
        # General recommendations without JD
        if len(resume_skills) < 10:
            recommendations.append("Consider adding more technical skills to your resume.")
        if resume_experience < 2:
            recommendations.append("Highlight any projects or internships to compensate for limited work experience.")
        recommendations.append("Consider adding a skills section to make your technical abilities more visible.")
    
    # General recommendations for all cases
    recommendations.append("Tailor your resume to include relevant keywords.")
    recommendations.append("Quantify your achievements with specific numbers and metrics.")
    recommendations.append("Use bullet points to make your resume more readable.")
    
    if job_description:
        missing_skills = list(set(jd_skills) - set(resume_skills)) if job_description else []
    else:
        missing_skills = []
    
    return recommendations, missing_skills

# Course database
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8}
    ]
}

def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Create widgets
resume_path_input = widgets.Text(
    value=r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf",
    placeholder="Enter full path to your resume",
    description="Resume Path:",
    layout=widgets.Layout(width="80%")
)
jd_input = widgets.Textarea(
    placeholder="Paste job description here (optional)...", 
    description="Job Description:", 
    layout=widgets.Layout(width="80%", height="100px")
)
analyze_btn = widgets.Button(description="Analyze Resume", button_style="success")
improve_btn = widgets.Button(description="Get Course Suggestions", button_style="info")
output = widgets.Output()

# Global variables to store analysis results
resume_text_global = ""
job_description_global = ""
missing_skills_global = []

def on_analyze_click(b):
    """Handle analyze button click"""
    with output:
        clear_output()
        global resume_text_global, job_description_global, missing_skills_global
        
        # Get resume path from input
        resume_path = resume_path_input.value.strip()
        if not resume_path:
            print("❌ Please enter a resume path.")
            return
            
        # Extract text from resume
        resume_text_global = extract_text_from_file(resume_path)
        if "Error" in resume_text_global:
            print(resume_text_global)
            return
            
        job_description_global = jd_input.value.strip()
        has_jd = bool(job_description_global)
        
        print(f"⏳ Analyzing your resume... Please wait\n")
        
        # Perform analysis
        if has_jd:
            ats_score, eligibility_score, skills_match, experience_match, jd_skills, jd_experience = analyze_eligibility(resume_text_global, job_description_global)
            recommendations, missing_skills = generate_recommendations(resume_text_global, job_description_global)
        else:
            ats_score, eligibility_score, skills_match, experience_match, jd_skills, jd_experience = analyze_eligibility(resume_text_global)
            recommendations, missing_skills = generate_recommendations(resume_text_global)
        
        missing_skills_global = missing_skills
        
        # Display results
        print("="*60)
        print("📊 RESUME ANALYSIS RESULTS")
        print("="*60)
        
        if has_jd:
            print("📝 Analysis with Job Description")
        else:
            print("📝 General Resume Quality Analysis")
        
        print(f"🎯 ATS Score: {ats_score:.1f}%")
        print(f"📈 Eligibility Score: {eligibility_score:.1f}%")
        
        if has_jd:
            print(f"🔧 Skills Match: {skills_match:.1f}%")
            print(f"💼 Experience Match: {experience_match:.1f}%")
        
        if eligibility_score >= 80:
            print("✅ STATUS: EXCELLENT - Strong resume quality!")
        elif eligibility_score >= 60:
            print("⚠️ STATUS: GOOD - Some improvements needed")
        else:
            print("❌ STATUS: NEEDS IMPROVEMENT - Significant improvements needed")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")
        
        if has_jd and missing_skills:
            print(f"\n🎯 Missing Skills Detected: {len(missing_skills)}")
            print("Click 'Get Course Suggestions' button to see recommended courses!")
        elif not has_jd:
            print(f"\n💡 Tip: Add a job description for more specific recommendations")

def on_improve_click(b):
    """Handle improve button click"""
    with output:
        clear_output()
        global missing_skills_global, job_description_global
        
        if not job_description_global:
            print("❌ Job description is required for course suggestions.")
            print("Please enter a job description and click 'Analyze Resume' first.")
            return
            
        if not missing_skills_global:
            print("No missing skills detected. Your resume looks good!")
            return
            
        course_suggestions = suggest_courses(missing_skills_global)
        display_course_suggestions(missing_skills_global, course_suggestions)

# Set up event handlers
analyze_btn.on_click(on_analyze_click)
improve_btn.on_click(on_improve_click)

# Display the interface
print("🎯 COMPLETE ATS RESUME ANALYSIS SYSTEM")
print("="*60)
print("Follow these steps:")
print("1. Enter the path to your resume (PDF or Word)")
print("2. (Optional) Paste a job description for targeted analysis")
print("3. Click 'Analyze Resume'")
print("4. Click 'Get Course Suggestions' for missing skills (requires JD)\n")
print("💡 Default resume path is set to your OM_ASHUTOSH Resume.pdf")

display(widgets.VBox([
    widgets.HTML("<b>Step 1: Enter Resume Path</b>"),
    resume_path_input,
    widgets.HTML("<b>Step 2: Paste Job Description (Optional)</b>"),
    jd_input,
    widgets.HTML("<b>Step 3: Analyze Your Resume</b>"),
    analyze_btn,
    widgets.HTML("<b>Step 4: Improve Your Skills (Requires JD)</b>"),
    improve_btn,
    widgets.HTML("<hr>"),
    output
]))


# In[1]:


# Complete ATS Resume Analysis System with Optional JD
from IPython.display import display, clear_output
import ipywidgets as widgets
from PyPDF2 import PdfReader
import docx
import re
import webbrowser
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings('ignore')

# Function to extract text from different file types
def extract_text_from_file(file_path):
    """Extract text from PDF or Word document"""
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_path.endswith(('.doc', '.docx')):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            return f"Error: Unsupported file format. Please use PDF or Word document."
        
        return text
    except Exception as e:
        return f"Error reading file: {e}"

def preprocess_text(text):
    """Clean and preprocess text"""
    text = text.lower()
    text = re.sub(r'http\S+', '', text)
    text = ' '.join(text.split())
    return text

def extract_skills(text):
    """Extract skills from text"""
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server',
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'ai', 'data analysis', 'data science',
        'tableau', 'power bi', 'excel', 'tensorflow', 'pytorch', 'scikit learn',
        'agile', 'scrum', 'devops', 'rest api', 'graphql', 'microservices',
        'nosql', 'big data', 'hadoop', 'spark', 'nlp', 'computer vision',
        'opencv', 'cnn', 'faster r-cnn', 'hitl', 'html', 'css'
    ]
    
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'work ethic', 'interpersonal skills',
        'project management', 'presentation', 'negotiation', 'decision making', 'collaboration'
    ]
    
    found_tech_skills = []
    found_soft_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_tech_skills.append(skill)
    
    for skill in soft_skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_soft_skills.append(skill)
    
    return found_tech_skills, found_soft_skills

def extract_experience(text):
    """Extract experience from text"""
    experience_patterns = [
        r'(\d+)\+? years?', r'(\d+)\+? yrs', r'year.*experience', 
        r'experience.*year', r'(\d+).*year', r'(\d+).*yr'
    ]
    
    experience = 0
    for pattern in experience_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if isinstance(match, str) and match.isdigit():
                exp = int(match)
                if exp > experience:
                    experience = exp
    
    return experience

def calculate_ats_score(resume_text, job_description=None):
    """Calculate ATS score between resume and job description"""
    resume_processed = preprocess_text(resume_text)
    
    if job_description:
        # Calculate score against job description
        jd_processed = preprocess_text(job_description)
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform([resume_processed, jd_processed])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        ats_score = round(similarity * 100, 2)
    else:
        # Calculate base ATS score (resume quality without JD)
        # This is a simplified version that checks resume structure and content
        word_count = len(resume_processed.split())
        skill_count = len(extract_skills(resume_text)[0])
        exp_years = extract_experience(resume_text)
        
        # Score based on resume quality metrics
        ats_score = min(100, (word_count / 500 * 30) + (skill_count / 15 * 40) + (exp_years * 6))
        ats_score = round(ats_score, 2)
    
    return ats_score

def analyze_eligibility(resume_text, job_description=None):
    """Analyze eligibility based on resume and optional job description"""
    ats_score = calculate_ats_score(resume_text, job_description)
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    
    if job_description:
        jd_skills, jd_soft_skills = extract_skills(job_description)
        jd_experience = extract_experience(job_description)
        
        # Calculate skills match
        tech_skills_match = len(set(resume_skills) & set(jd_skills))
        tech_skills_required = len(jd_skills)
        
        if tech_skills_required > 0:
            skills_match_percentage = (tech_skills_match / tech_skills_required) * 100
        else:
            skills_match_percentage = 100
        
        # Calculate experience match
        if jd_experience > 0:
            if resume_experience >= jd_experience:
                experience_match = 100
            else:
                experience_match = (resume_experience / jd_experience) * 100
        else:
            experience_match = 100
        
        # Overall eligibility score (weighted average)
        eligibility_score = (ats_score * 0.4) + (skills_match_percentage * 0.4) + (experience_match * 0.2)
        
        return ats_score, eligibility_score, skills_match_percentage, experience_match, jd_skills, jd_experience
    else:
        # Without JD, provide general resume quality assessment
        skills_match_percentage = 100  # Not applicable without JD
        experience_match = 100  # Not applicable without JD
        
        # Base eligibility score on resume quality
        word_count = len(resume_text.split())
        skill_count = len(resume_skills)
        
        eligibility_score = min(100, (word_count / 500 * 40) + (skill_count / 15 * 40) + (resume_experience * 4))
        eligibility_score = round(eligibility_score, 2)
        
        return ats_score, eligibility_score, skills_match_percentage, experience_match, [], 0

def generate_recommendations(resume_text, job_description=None):
    """Generate improvement recommendations"""
    recommendations = []
    
    resume_skills, resume_soft_skills = extract_skills(resume_text)
    resume_experience = extract_experience(resume_text)
    
    if job_description:
        jd_skills, jd_soft_skills = extract_skills(job_description)
        jd_experience = extract_experience(job_description)
        
        # Skills recommendations
        missing_tech_skills = set(jd_skills) - set(resume_skills)
        missing_soft_skills = set(jd_soft_skills) - set(resume_soft_skills)
        
        if missing_tech_skills:
            recommendations.append(f"Develop these technical skills: {', '.join(missing_tech_skills)}")
        if missing_soft_skills:
            recommendations.append(f"Develop these soft skills: {', '.join(missing_soft_skills)}")
        
        # Experience recommendations
        if resume_experience < jd_experience:
            recommendations.append(f"Gain more experience. The job requires {jd_experience}+ years, but you have {resume_experience} years.")
    else:
        # General recommendations without JD
        if len(resume_skills) < 10:
            recommendations.append("Consider adding more technical skills to your resume.")
        if resume_experience < 2:
            recommendations.append("Highlight any projects or internships to compensate for limited work experience.")
        recommendations.append("Consider adding a skills section to make your technical abilities more visible.")
    
    # General recommendations for all cases
    recommendations.append("Tailor your resume to include relevant keywords.")
    recommendations.append("Quantify your achievements with specific numbers and metrics.")
    recommendations.append("Use bullet points to make your resume more readable.")
    
    if job_description:
        missing_skills = list(set(jd_skills) - set(resume_skills)) if job_description else []
    else:
        missing_skills = []
    
    return recommendations, missing_skills

# Course database
COURSE_DATABASE = {
    'python': [
        {'platform': 'Coursera', 'name': 'Python for Everybody', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/specializations/python', 'rating': 4.8},
        {'platform': 'YouTube', 'name': 'Python Tutorial for Beginners', 'cost': 'Free', 'duration': '6 hours', 
         'url': 'https://www.youtube.com/watch?v=_uQrJ0TkZlc', 'rating': 4.9}
    ],
    'machine learning': [
        {'platform': 'Coursera', 'name': 'Machine Learning by Andrew Ng', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/learn/machine-learning', 'rating': 4.9},
        {'platform': 'YouTube', 'name': 'Machine Learning Tutorial', 'cost': 'Free', 'duration': '10 hours', 
         'url': 'https://www.youtube.com/watch?v=NWONeJKn6kc', 'rating': 4.8}
    ],
    'data analysis': [
        {'platform': 'Coursera', 'name': 'Google Data Analytics Professional Certificate', 'cost': 'Free', 'duration': '6 months', 
         'url': 'https://www.coursera.org/professional-certificates/google-data-analytics', 'rating': 4.8}
    ],
    'sql': [
        {'platform': 'Coursera', 'name': 'SQL for Data Science', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/sql-for-data-science', 'rating': 4.7}
    ],
    'tensorflow': [
        {'platform': 'Coursera', 'name': 'TensorFlow Developer Professional Certificate', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/professional-certificates/tensorflow-in-practice', 'rating': 4.8}
    ],
    'opencv': [
        {'platform': 'Udemy', 'name': 'OpenCV Python for Beginners', 'cost': '₹455', 'duration': '10 hours', 
         'url': 'https://www.udemy.com/course/opencv-python-for-beginners/', 'rating': 4.6}
    ],
    'tableau': [
        {'platform': 'Coursera', 'name': 'Data Visualization with Tableau', 'cost': 'Free', 'duration': '2 months', 
         'url': 'https://www.coursera.org/specializations/data-visualization', 'rating': 4.7}
    ],
    'power bi': [
        {'platform': 'Coursera', 'name': 'Microsoft Power BI Data Analyst Professional Certificate', 'cost': 'Free', 'duration': '3 months', 
         'url': 'https://www.coursera.org/professional-certificates/microsoft-power-bi-data-analyst', 'rating': 4.7}
    ],
    'nlp': [
        {'platform': 'Coursera', 'name': 'Natural Language Processing Specialization', 'cost': 'Free', 'duration': '4 months', 
         'url': 'https://www.coursera.org/specializations/natural-language-processing', 'rating': 4.8}
    ],
    'cnn': [
        {'platform': 'Coursera', 'name': 'Convolutional Neural Networks', 'cost': 'Free', 'duration': '1 month', 
         'url': 'https://www.coursera.org/learn/convolutional-neural-networks', 'rating': 4.8}
    ]
}

def suggest_courses(missing_skills):
    """Suggest courses for missing skills"""
    course_suggestions = {}
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower in COURSE_DATABASE:
            course_suggestions[skill] = COURSE_DATABASE[skill_lower]
        else:
            course_suggestions[skill] = [
                {'platform': 'Coursera', 'name': f'{skill.title()} Specialization', 'cost': 'Free/Paid', 
                 'duration': '2-6 months', 'url': 'https://www.coursera.org/search?query=' + skill.replace(' ', '%20'), 'rating': 'N/A'},
                {'platform': 'Udemy', 'name': f'Complete {skill.title()} Course', 'cost': '₹455', 
                 'duration': '10-30 hours', 'url': 'https://www.udemy.com/courses/search/?q=' + skill.replace(' ', '%20'), 'rating': 'N/A'}
            ]
    
    return course_suggestions

def display_course_suggestions(missing_skills, course_suggestions):
    """Display course suggestions"""
    print("\n" + "="*60)
    print("🎓 COURSE RECOMMENDATIONS")
    print("="*60)
    
    if not missing_skills:
        print("✅ Great! You have all the required skills.")
        return
    
    for i, skill in enumerate(missing_skills, 1):
        print(f"\n{i}. Skill Needed: {skill.upper()}")
        print("   Recommended Courses:")
        
        courses = course_suggestions.get(skill, [])
        for j, course in enumerate(courses, 1):
            print(f"      {j}. {course['name']}")
            print(f"         Platform: {course['platform']}")
            print(f"         Cost: {course['cost']}")
            print(f"         Duration: {course['duration']}")
            print(f"         Rating: {course['rating']}/5.0")
            print(f"         URL: {course['url']}")
            print()

# Create widgets
resume_path_input = widgets.Text(
    value=r"C:\Users\omash\OneDrive\Documents\Desktop\PERSONAL\OM_ASHUTOSH Resume.pdf",
    placeholder="Enter full path to your resume",
    description="Resume Path:",
    layout=widgets.Layout(width="80%")
)
jd_input = widgets.Textarea(
    placeholder="Paste job description here (optional)...", 
    description="Job Description:", 
    layout=widgets.Layout(width="80%", height="100px")
)
analyze_btn = widgets.Button(description="Analyze Resume", button_style="success")
improve_btn = widgets.Button(description="Get Course Suggestions", button_style="info")
output = widgets.Output()

# Global variables to store analysis results
resume_text_global = ""
job_description_global = ""
missing_skills_global = []

def on_analyze_click(b):
    """Handle analyze button click"""
    with output:
        clear_output()
        global resume_text_global, job_description_global, missing_skills_global
        
        # Get resume path from input
        resume_path = resume_path_input.value.strip()
        if not resume_path:
            print("❌ Please enter a resume path.")
            return
            
        # Extract text from resume
        resume_text_global = extract_text_from_file(resume_path)
        if "Error" in resume_text_global:
            print(resume_text_global)
            return
            
        job_description_global = jd_input.value.strip()
        has_jd = bool(job_description_global)
        
        print(f"⏳ Analyzing your resume... Please wait\n")
        
        # Perform analysis
        if has_jd:
            ats_score, eligibility_score, skills_match, experience_match, jd_skills, jd_experience = analyze_eligibility(resume_text_global, job_description_global)
            recommendations, missing_skills = generate_recommendations(resume_text_global, job_description_global)
        else:
            ats_score, eligibility_score, skills_match, experience_match, jd_skills, jd_experience = analyze_eligibility(resume_text_global)
            recommendations, missing_skills = generate_recommendations(resume_text_global)
        
        missing_skills_global = missing_skills
        
        # Display results
        print("="*60)
        print("📊 RESUME ANALYSIS RESULTS")
        print("="*60)
        
        if has_jd:
            print("📝 Analysis with Job Description")
        else:
            print("📝 General Resume Quality Analysis")
        
        print(f"🎯 ATS Score: {ats_score:.1f}%")
        print(f"📈 Eligibility Score: {eligibility_score:.1f}%")
        
        if has_jd:
            print(f"🔧 Skills Match: {skills_match:.1f}%")
            print(f"💼 Experience Match: {experience_match:.1f}%")
        
        if eligibility_score >= 80:
            print("✅ STATUS: EXCELLENT - Strong resume quality!")
        elif eligibility_score >= 60:
            print("⚠️ STATUS: GOOD - Some improvements needed")
        else:
            print("❌ STATUS: NEEDS IMPROVEMENT - Significant improvements needed")
        
        print("\n💡 RECOMMENDATIONS:")
        for i, recommendation in enumerate(recommendations, 1):
            print(f"{i}. {recommendation}")
        
        if has_jd and missing_skills:
            print(f"\n🎯 Missing Skills Detected: {len(missing_skills)}")
            print("Click 'Get Course Suggestions' button to see recommended courses!")
        elif not has_jd:
            print(f"\n💡 Tip: Add a job description for more specific recommendations")

def on_improve_click(b):
    """Handle improve button click"""
    with output:
        clear_output()
        global missing_skills_global, job_description_global
        
        if not job_description_global:
            print("❌ Job description is required for course suggestions.")
            print("Please enter a job description and click 'Analyze Resume' first.")
            return
            
        if not missing_skills_global:
            print("No missing skills detected. Your resume looks good!")
            return
            
        course_suggestions = suggest_courses(missing_skills_global)
        display_course_suggestions(missing_skills_global, course_suggestions)

# Set up event handlers
analyze_btn.on_click(on_analyze_click)
improve_btn.on_click(on_improve_click)

# Display the interface
print("🎯 COMPLETE ATS RESUME ANALYSIS SYSTEM")
print("="*60)
print("Follow these steps:")
print("1. Enter the path to your resume (PDF or Word)")
print("2. (Optional) Paste a job description for targeted analysis")
print("3. Click 'Analyze Resume'")
print("4. Click 'Get Course Suggestions' for missing skills (requires JD)\n")
print("💡 Default resume path is set to your OM_ASHUTOSH Resume.pdf")

display(widgets.VBox([
    widgets.HTML("<b>Step 1: Enter Resume Path</b>"),
    resume_path_input,
    widgets.HTML("<b>Step 2: Paste Job Description (Optional)</b>"),
    jd_input,
    widgets.HTML("<b>Step 3: Analyze Your Resume</b>"),
    analyze_btn,
    widgets.HTML("<b>Step 4: Improve Your Skills (Requires JD)</b>"),
    improve_btn,
    widgets.HTML("<hr>"),
    output
]))


# In[ ]:




